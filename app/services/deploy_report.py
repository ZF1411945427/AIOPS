"""部署报告格式化模块(纯函数, 从 deploy_service.py 拆分)。

包含报告生成/格式化的纯函数, 不依赖部署执行引擎内部状态:
  - _generate_fallback_report : AI 不可用时的结构化摘要报告
  - _report_to_markdown       : 报告 -> Markdown
  - _report_to_html           : 报告 -> HTML(可打印/PDF)
  - _report_to_docx           : 报告 -> Word .docx bytes
拆自 app/services/deploy_service.py(原 4064 行单体), 保持函数签名与行为不变。
"""

def _generate_fallback_report(plan, step_summary, assets, preflight_detail, test_detail, deploy_info=None) -> dict:
    """当 AI 不可用时，生成结构化摘要报告。"""
    total = len(step_summary)
    succeeded = sum(1 for s in step_summary if s["status"] == "succeeded")
    failed = sum(1 for s in step_summary if s["status"] == "failed")
    skipped = sum(1 for s in step_summary if s["status"] == "skipped")
    status_text = "成功" if plan.status == "succeeded" else "失败" if plan.status == "failed" else plan.status
    summary = (
        f"部署计划「{plan.name}」执行{status_text}。"
        f"共 {total} 步，成功 {succeeded} 步，失败 {failed} 步，跳过 {skipped} 步。"
        f"目标资产 {len(assets)} 台。"
    )
    steps_md = "| 序号 | 步骤 | 耗时 | 状态 | 重试 | 诊断 |\n|------|------|------|------|------|------|\n"
    for s in step_summary:
        diag = (s.get("diagnosis", "") or "")[:50]
        steps_md += f"| {s['order']} | {s['description'][:40]} | {s.get('duration', '-')} | {s['status']} | {s.get('retry_count', 0)} | {diag} |\n"

    env_text = "目标资产:\n"
    for a in assets:
        env_text += f"- {a.get('name', '')} ({a.get('ip', '')})\n"

    issues = []
    for s in step_summary:
        if s["status"] == "failed" and s.get("diagnosis"):
            issues.append({"severity": "high", "description": f"步骤 {s['order']} {s['description']} 失败", "resolution": s["diagnosis"][:200], "status": "resolved"})

    return {
        "title": f"部署报告 - {plan.name}",
        "executive_summary": summary,
        "deployment_architecture": f"目标资产 {len(assets)} 台，共 {total} 步部署步骤" + (f"\n{deploy_info.get('architecture', '')}" if deploy_info else ""),
        "start_stop_commands": "、\n".join(deploy_info.get("start_stop_commands", [])) if deploy_info else "见执行步骤",
        "deploy_paths": "、\n".join(deploy_info.get("deploy_paths", [])) if deploy_info else "无",
        "service_ports": "、\n".join(deploy_info.get("service_ports", [])) if deploy_info else "无",
        "access_methods": "、\n".join(deploy_info.get("access_methods", [])) if deploy_info else "无",
        "login_info": "、\n".join([f"{l['method']}" for l in (deploy_info.get("login_info", []) if deploy_info else [])]) or "无",
        "environment": {"os": "见环境探查", "notes": env_text},
        "timeline": f"总耗时: {sum(int(s.get('duration', '0').rstrip('s')) for s in step_summary if s.get('duration', '-') != '-')}s",
        "steps_table": steps_md,
        "key_observations": [f"共 {total} 步，{succeeded} 成功，{failed} 失败，{skipped} 跳过"],
        "verification": f"预检{'通过' if preflight_detail else '未执行'}，部署后验证{'通过' if test_detail else '未执行'}",
        "test_results": f"验证项 {len(test_detail)} 项" if test_detail else "未执行验证",
        "issues": issues or [{"severity": "low", "description": "无异常", "resolution": "-", "status": "resolved"}],
        "risk_assessment": "基于当前部署结果，风险可控",
        "recommendations": ["监控部署后的服务运行状态", "检查日志输出是否正常", "建议定期执行健康检查"],
        "overall_assessment": f"{plan.status} - {summary}",
    }


def _report_to_markdown(report: dict) -> str:
    """将部署报告转换为可直接交付的 Markdown 文档。"""
    title = report.get("title", "部署报告")
    plan_name = report.get("plan_name", "")
    status = report.get("status", "")
    deployed_at = report.get("deployed_at", "")
    deploy_count = report.get("deploy_count", 0)
    total_steps = report.get("total_steps", 0)
    succeeded_steps = report.get("succeeded_steps", 0)
    failed_steps = report.get("failed_steps", 0)
    skipped_steps = report.get("skipped_steps", 0)
    total_assets = report.get("total_assets", 0)
    preflight_passed = report.get("preflight_passed", False)
    verification_passed = report.get("verification_passed", False)
    ai_decisions = report.get("ai_decisions", 0)

    status_icon = "✅" if status == "succeeded" else "❌" if status == "failed" else "⚠️"
    preflight_icon = "✅" if preflight_passed else "❌"
    verify_icon = "✅" if verification_passed else "❌"

    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append("> **文档类型**: 部署报告 (Deployment Report)")
    lines.append(f"> **生成时间**: {deployed_at}")
    lines.append(f"> **状态**: {status_icon} **{status.upper()}**")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 📋 执行摘要")
    lines.append("")
    lines.append(report.get("executive_summary", "无"))
    lines.append("")
    lines.append("### 关键指标")
    lines.append("")
    lines.append("| 指标 | 值 |")
    lines.append("|------|----|")
    lines.append(f"| 部署计划 | {plan_name} |")
    lines.append(f"| 部署次数 | 第 {deploy_count} 次 |")
    lines.append(f"| 执行时间 | {deployed_at} |")
    lines.append(f"| 目标资产 | {total_assets} 台 |")
    lines.append(f"| 总步骤数 | {total_steps} |")
    lines.append(f"| 成功步骤 | {succeeded_steps} |")
    lines.append(f"| 失败步骤 | {failed_steps} |")
    lines.append(f"| 跳过步骤 | {skipped_steps} |")
    lines.append(f"| 预检结果 | {preflight_icon} {'全部通过' if preflight_passed else '有失败项'} |")
    lines.append(f"| 部署验证 | {verify_icon} {'全部通过' if verification_passed else '有失败项'} |")
    lines.append(f"| AI 决策次数 | {ai_decisions} |")
    lines.append(f"| 总体评估 | {status_icon} {report.get('overall_assessment', status)} |")
    lines.append("")

    env = report.get("environment", {})
    if isinstance(env, dict):
        lines.append("---")
        lines.append("")
        lines.append("## 🖥️ 环境信息")
        lines.append("")
        lines.append("| 项目 | 详情 |")
        lines.append("|------|------|")
        for k, v in env.items():
            lines.append(f"| {k} | {str(v)[:200]} |")
        lines.append("")
    elif isinstance(env, str):
        lines.append("---")
        lines.append("")
        lines.append("## 🖥️ 环境信息")
        lines.append("")
        lines.append(env)
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## ⏱️ 时间线")
    lines.append("")
    lines.append(report.get("timeline", "无"))
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 📊 步骤执行结果")
    lines.append("")
    lines.append(report.get("steps_table", "无"))
    lines.append("")

    if report.get("key_observations"):
        lines.append("---")
        lines.append("")
        lines.append("## 🔍 关键观察")
        lines.append("")
        for obs in report.get("key_observations", []):
            lines.append(f"- {obs}")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## ✅ 部署验证")
    lines.append("")
    lines.append(report.get("verification", "无"))
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 🧪 测试记录")
    lines.append("")
    lines.append(report.get("test_results", "无"))
    lines.append("")

    issues = report.get("issues", [])
    if issues:
        lines.append("---")
        lines.append("")
        lines.append("## 🐛 问题与处理")
        lines.append("")
        lines.append("| 严重程度 | 问题描述 | 处理方式 | 状态 |")
        lines.append("|---------|---------|---------|------|")
        for issue in issues:
            sev = issue.get("severity", "low")
            sev_icon = "🔴" if sev == "high" else "🟡" if sev == "medium" else "🟢"
            lines.append(f"| {sev_icon} {sev} | {issue.get('description', '')[:100]} | {issue.get('resolution', '')[:100]} | {issue.get('status', '')} |")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## ⚠️ 风险评估")
    lines.append("")
    lines.append(report.get("risk_assessment", "无"))
    lines.append("")

    recommendations = report.get("recommendations", [])
    if recommendations:
        lines.append("---")
        lines.append("")
        lines.append("## 💡 改进建议")
        lines.append("")
        for i, rec in enumerate(recommendations, 1):
            lines.append(f"{i}. {rec}")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 📝 总体评估")
    lines.append("")
    lines.append(report.get("overall_assessment", "无"))
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(f"*报告由 AIOps 自动部署系统自动生成 | {deployed_at}*")

    return "\n".join(lines)


def _report_to_html(report: dict) -> str:
    """将部署报告转换为可直接打印/保存为 PDF 的专业 HTML 文档。"""
    md = _report_to_markdown(report)
    import markdown as _md
    body_html = _md.markdown(md, extensions=["tables", "fenced_code", "codehilite"])

    status = report.get("status", "unknown")
    status_color = "#22c55e" if status == "succeeded" else "#ef4444" if status == "failed" else "#f59e0b"

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{report.get('title', '部署报告')}</title>
<style>
  @page {{ size: A4; margin: 2cm; }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: 'PingFang SC', 'Microsoft YaHei', 'Helvetica Neue', Arial, sans-serif; color: #1e293b; line-height: 1.7; padding: 40px; background: #f8fafc; }}
  .report-container {{ max-width: 900px; margin: 0 auto; background: #fff; padding: 50px 60px; box-shadow: 0 4px 24px rgba(0,0,0,0.06); border-radius: 12px; }}
  h1 {{ font-size: 28px; color: #0f172a; border-bottom: 3px solid {status_color}; padding-bottom: 16px; margin-bottom: 24px; }}
  h2 {{ font-size: 20px; color: #0f172a; margin-top: 32px; margin-bottom: 16px; padding-left: 10px; border-left: 4px solid {status_color}; }}
  h3 {{ font-size: 16px; color: #334155; margin-top: 20px; margin-bottom: 10px; }}
  p {{ margin-bottom: 12px; color: #475569; }}
  table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }}
  th {{ background: #f1f5f9; color: #0f172a; font-weight: 600; padding: 10px 12px; border: 1px solid #e2e8f0; text-align: left; }}
  td {{ padding: 8px 12px; border: 1px solid #e2e8f0; color: #475569; }}
  tr:nth-child(even) td {{ background: #f8fafc; }}
  blockquote {{ background: #f1f5f9; border-left: 4px solid {status_color}; padding: 12px 18px; margin: 16px 0; border-radius: 0 8px 8px 0; color: #475569; }}
  code {{ background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: 13px; font-family: 'JetBrains Mono', 'Fira Code', monospace; }}
  pre {{ background: #0f172a; color: #e2e8f0; padding: 16px; border-radius: 8px; overflow-x: auto; font-size: 13px; line-height: 1.5; margin: 12px 0; }}
  ul, ol {{ padding-left: 24px; margin-bottom: 12px; }}
  li {{ margin-bottom: 6px; color: #475569; }}
  .status-badge {{ display: inline-block; padding: 4px 16px; border-radius: 20px; font-size: 14px; font-weight: 600; color: #fff; background: {status_color}; }}
  .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; font-size: 12px; color: #94a3b8; text-align: center; }}
  hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 24px 0; }}
  @media print {{
    body {{ background: #fff; padding: 0; }}
    .report-container {{ box-shadow: none; padding: 40px; }}
    .no-print {{ display: none !important; }}
  }}
  .no-print {{ text-align: center; margin-bottom: 20px; }}
  .no-print button {{ padding: 10px 24px; background: #6366f1; color: #fff; border: none; border-radius: 8px; font-size: 14px; cursor: pointer; margin: 0 8px; }}
  .no-print button:hover {{ background: #4f46e5; }}
  .no-print button.secondary {{ background: #e2e8f0; color: #475569; }}
  .no-print button.secondary:hover {{ background: #cbd5e1; }}
</style>
</head>
<body>
<div class="no-print">
  <button onclick="window.print()">🖨️ 打印 / 导出 PDF</button>
  <button class="secondary" onclick="location.href='/deploy/api/plans/{report.get("plan_id", 0)}/report/download?format=md'">📥 下载 Markdown</button>
</div>
<div class="report-container">
{body_html}
<div class="footer">
  <p>本报告由 AIOps 智能运维平台自动生成 | {report.get("deployed_at", "")}</p>
  <p style="margin-top:4px;">AIOps Deployment Report · Confidential</p>
</div>
</div>
</body>
</html>"""


def _report_to_docx(report: dict) -> bytes:
    """将部署报告生成为专业 Word .docx 文档，返回 bytes。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── 全局样式设置 ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            hs.font.size = Pt(22)
        elif level == 2:
            hs.font.size = Pt(16)
        else:
            hs.font.size = Pt(13)

    title = report.get("title", "部署报告")
    status = report.get("status", "unknown")
    deployed_at = report.get("deployed_at", "")

    # ── 封面区域 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("部署报告")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{title}")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    status_text = "成功" if status == "succeeded" else "失败" if status == "failed" else status
    status_color = "22c55e" if status == "succeeded" else "ef4444" if status == "failed" else "f59e0b"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n状态: {status_text.upper()}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(int(status_color[:2], 16), int(status_color[2:4], 16), int(status_color[4:], 16))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n生成时间: {deployed_at}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_page_break()

    # ── 辅助函数 ──
    def _add_kv_table(doc, headers, rows):
        t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            for p in t.rows[0].cells[i].paragraphs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(10)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                t.rows[ri + 1].cells[ci].text = str(val)
                for p in t.rows[ri + 1].cells[ci].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
        doc.add_paragraph()

    # ── 1. 执行摘要 ──
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(report.get("executive_summary", "无"))

    # KPI 指标
    doc.add_heading("关键指标", level=2)
    _add_kv_table(doc, ["指标", "值"], [
        ["部署计划", report.get("plan_name", "")],
        ["部署次数", f"第 {report.get('deploy_count', 0)} 次"],
        ["执行时间", deployed_at],
        ["目标资产", f"{report.get('total_assets', 0)} 台"],
        ["总步骤数", str(report.get("total_steps", 0))],
        ["成功步骤", str(report.get("succeeded_steps", 0))],
        ["失败步骤", str(report.get("failed_steps", 0))],
        ["跳过步骤", str(report.get("skipped_steps", 0))],
        ["预检结果", "通过" if report.get("preflight_passed") else "未通过"],
        ["部署验证", "通过" if report.get("verification_passed") else "有失败项"],
        ["AI 决策次数", str(report.get("ai_decisions", 0))],
    ])

    # ── 1b. 部署架构 ──
    arch = report.get("deployment_architecture", "")
    if arch:
        doc.add_heading("部署架构", level=1)
        doc.add_paragraph(arch)

    # ── 1c. 启停服务命令 ──
    ssc = report.get("start_stop_commands", "")
    if ssc:
        doc.add_heading("启停服务命令", level=1)
        for line in ssc.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1d. 部署路径 ──
    dp = report.get("deploy_paths", "")
    if dp:
        doc.add_heading("部署路径", level=1)
        for line in dp.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1e. 服务端口 ──
    sp = report.get("service_ports", "")
    if sp:
        doc.add_heading("服务端口", level=1)
        for line in sp.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1f. 访问方式 ──
    am = report.get("access_methods", "")
    if am:
        doc.add_heading("访问方式", level=1)
        for line in am.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1g. 登录信息 ──
    li = report.get("login_info", "")
    if li:
        doc.add_heading("登录信息", level=1)
        for line in li.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 2. 环境信息 ──
    env = report.get("environment", {})
    if isinstance(env, dict):
        doc.add_heading("环境信息", level=1)
        _add_kv_table(doc, ["项目", "详情"], [[k, str(v)[:200]] for k, v in env.items()])

    # ── 3. 时间线 ──
    doc.add_heading("时间线", level=1)
    doc.add_paragraph(report.get("timeline", "无"))

    # ── 4. 步骤执行结果 ──
    doc.add_heading("步骤执行结果", level=1)
    steps_md = report.get("steps_table", "")
    # Parse markdown table into list
    if steps_md:
        lines = [l.strip() for l in steps_md.split("\n") if l.strip()]
        data_rows = []
        headers = []
        for i, line in enumerate(lines):
            if line.startswith("|") and line.endswith("|"):
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if i == 0:
                    headers = cells
                elif "---" not in line and cells:
                    data_rows.append(cells)
        if headers:
            _add_kv_table(doc, headers, data_rows)

    # ── 5. 关键观察 ──
    obs = report.get("key_observations", [])
    if obs:
        doc.add_heading("关键观察", level=1)
        for o in obs:
            doc.add_paragraph(o, style='List Bullet')

    # ── 6. 部署验证 ──
    doc.add_heading("部署验证", level=1)
    doc.add_paragraph(report.get("verification", "无"))

    # ── 7. 测试记录 ──
    doc.add_heading("测试记录", level=1)
    doc.add_paragraph(report.get("test_results", "无"))

    # ── 8. 问题与处理 ──
    issues = report.get("issues", [])
    if issues:
        doc.add_heading("问题与处理", level=1)
        _add_kv_table(doc, ["严重程度", "问题描述", "处理方式", "状态"],
            [[i.get("severity", ""), i.get("description", "")[:100],
              i.get("resolution", "")[:100], i.get("status", "")] for i in issues])

    # ── 9. 风险评估 ──
    doc.add_heading("风险评估", level=1)
    doc.add_paragraph(report.get("risk_assessment", "无"))

    # ── 10. 改进建议 ──
    recs = report.get("recommendations", [])
    if recs:
        doc.add_heading("改进建议", level=1)
        for i, rec in enumerate(recs, 1):
            doc.add_paragraph(f"{i}. {rec}")

    # ── 11. 总体评估 ──
    doc.add_heading("总体评估", level=1)
    doc.add_paragraph(report.get("overall_assessment", "无"))

    # ── 页脚信息 ──
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"报告由 AIOps 智能运维平台自动生成 | {deployed_at}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
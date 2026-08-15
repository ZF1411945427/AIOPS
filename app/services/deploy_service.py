import json
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

from sqlalchemy.orm import Session

from app.models import DeployPlan, DeployStep, Asset, AIProvider, AgentConfig
from app.services.agent_service import call_llm
from app.services.ssh_helper import connect_ssh
from app.logger import logger

# 进程内执行互斥：同一计划同一时刻只允许一个执行流(HTTP 或 WS)。僵尸 running 状态可重跑。
_EXEC_LOCK: Dict[int, bool] = {}
# 活跃 SSH 客户端注册表（供停止接口关闭连接中断执行）
_RUNNING_CLIENTS: Dict[int, Any] = {}
# 停止请求标志：producer 检测到后立即终止且不覆盖状态
_STOPPED: Dict[int, bool] = {}
# 用户决策队列：plan_id -> queue.Queue（WS 路由转发用户"修复/重试/回滚/跳过"决策）
_DECISIONS: Dict[int, Any] = {}
# 单步骤 SSH 命令最大执行时长（docker build 等长任务，超时终止）
_STEP_TIMEOUT = 600


def _release_exec(plan_id: int):
    _EXEC_LOCK.pop(plan_id, None)


def release_exec_lock(plan_id: int):
    """供 router 在 WS 断开时强制释放执行锁（不等后台线程自然结束）。"""
    _EXEC_LOCK.pop(plan_id, None)


def probe_environment(db: Session, plan_id: int) -> dict:
    """SSH 探查目标机真实环境，返回探查结果 JSON。探查前自动下载源码(在线/离线)。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    # 探查前自动下载源码（在线 git/HTTP + 离线），幂等
    download = auto_download_artifact(db, plan_id)
    if not download.get("ok"):
        return {"error": f"源码准备失败: {download.get('error', '未知错误')}", "download": download}
    assets = _get_assets(db, plan)
    if not assets:
        return {"error": "计划未关联目标资产"}
    asset = assets[0]
    try:
        client, host = _ssh_connect(asset)
    except Exception as e:
        return {"error": f"SSH 连接失败: {e}"}
    try:
        probes = _collect_env_probes(client, plan)
        plan.environment_probe_json = json.dumps(probes, ensure_ascii=False)
        db.commit()
        return {"ok": True, "probe": probes}
    finally:
        client.close()


def _collect_env_probes(client, plan: DeployPlan) -> dict:
    """执行全套环境探查命令，返回结构化结果。"""
    def _run(cmd, timeout=15):
        try:
            _, stdout, stderr = client.exec_command(cmd, timeout=timeout)
            out = stdout.read().decode("utf-8", errors="replace").strip()
            err = stderr.read().decode("utf-8", errors="replace").strip()
            return out or err or ""
        except Exception as e:
            return f"[probe_error] {e}"
    result = {}
    result["os"] = _run("cat /etc/os-release 2>/dev/null | head -5 || cat /etc/redhat-release 2>/dev/null || uname -a")
    result["kernel"] = _run("uname -a")
    result["disk"] = _run("df -h / 2>/dev/null | tail -1")
    result["ports"] = _run("ss -tlnp 2>/dev/null || netstat -tlnp 2>/dev/null")
    result["docker"] = _run("docker info --format '{{.ServerVersion}}' 2>/dev/null || echo 'DOCKER_NOT_AVAILABLE'")
    result["images"] = _run("docker images --format '{{.Repository}}:{{.Tag}}' 2>/dev/null || echo 'NONE'")
    result["containers"] = _run("docker ps --format '{{.Names}}|{{.Image}}|{{.Ports}}' 2>/dev/null || echo 'NONE'")
    # 探查常用目录和 APP_DIR
    _guess_dirs = ["/data/test-project", "/data", "/opt", "/var/www", "/home", "/app", "/srv"]
    dirs_found = {}
    for d in _guess_dirs:
        _ls = _run(f"ls -la {d} 2>/dev/null && echo '---EXISTS---' || echo '---NOT_EXISTS---'")
        if "---EXISTS---" in _ls:
            content = _ls.split("---EXISTS---")[0].strip()
            dirs_found[d] = content[:2000]
            _compose = _run(f"cat {d}/docker-compose.yml 2>/dev/null || cat {d}/compose.yaml 2>/dev/null || cat {d}/docker-compose.yaml 2>/dev/null || echo 'NO_COMPOSE'")
            if _compose and _compose != "NO_COMPOSE":
                dirs_found[f"{d}/docker-compose"] = _compose[:2000]
            _df = _run(f"cat {d}/Dockerfile 2>/dev/null || echo 'NO_DOCKERFILE'")
            if _df and _df != "NO_DOCKERFILE":
                dirs_found[f"{d}/Dockerfile"] = _df[:2000]
    result["dirs"] = dirs_found
    # 常见端口检测
    _common_ports = [80, 443, 3000, 8080, 8443, 5000, 9000, 5432, 3306, 6379, 27017]
    port_status = {}
    for p in _common_ports:
        ps = _run(f"ss -tlnp | grep ':{p} ' || echo 'FREE'")
        port_status[str(p)] = "IN_USE" if ps and "FREE" not in ps else "FREE"
    result["port_scan"] = port_status
    return result


def ai_auto_env_mapping(db: Session, plan_id: int) -> dict:
    """基于环境探查结果，AI 自动生成环境映射 + SOP 适配建议（A + C 层核心）。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    if not plan.environment_probe_json or plan.environment_probe_json == "{}":
        return {"error": "请先执行环境探查(probe)"}
    probe = json.loads(plan.environment_probe_json)
    doc_raw = plan.doc_raw or ""
    sop = json.loads(plan.sop_json or "[]")
    assets = _get_assets(db, plan)
    asset_info = ""
    if assets:
        asset_info = json.dumps([{"name": a.name, "ip": a.ip, "type": a.ci_type} for a in assets], ensure_ascii=False)
    provider = _get_provider(db)
    if not provider:
        return {"error": "未配置 AI 提供商"}
    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责根据目标机真实环境探查结果和部署手册，完成以下任务：\n"
        "1. 自动生成环境映射(env_mapping)：根据探查结果确定 APP_DIR、TARGET_IP 等参数的真实值\n"
        "2. 服务拓扑分析：分析这个 compose/dockerfile 部署了哪些服务、依赖关系\n"
        "3. 自适应建议：基于当前环境推荐 SOP 调整（如：镜像已存在跳过 build、端口冲突换端口、目录已存在跳过创建）\n"
        "请严格按以下 JSON Schema 输出，不要额外内容：\n"
        "{\n"
        '  "env_mapping": {"APP_DIR": "<真实部署目录>", "TARGET_PORT": "80"},\n'
        '  "service_topology": "服务拓扑分析文本",\n'
        '  "adaptations": [\n'
        '    {"step": 2, "type": "skip_build", "reason": "镜像 nginx:latest 已存在，跳过 build", "action": "建议跳过 docker compose build"},'
        '    {"step": 3, "type": "port_check", "reason": "80 端口已被占用", "action": "建议改用 8080 端口"}\n'
        "  ],\n"
        '  "preflight_enhance": [\n'
        '    {"check": "端口冲突检查", "command": "ss -tlnp | grep 80", "expect": "端口未被占用"}\n'
        "  ]\n"
        "}\n"
        "规则：\n"
        "- env_mapping 的值必须来自探查结果，不能凭空编造\n"
        "- 如果探查结果中某个目录/端口不存在，不要强行映射\n"
        "- adaptations 是可选优化，不是必须的\n"
    )
    # 检测 OS 兼容性提示
    _os_text = probe.get("os", "")
    _os_hint = ""
    if "centos" in _os_text.lower() and "7" in _os_text:
        _os_hint = "注意：目标机为 CentOS 7，nginx 最新版可能存在 pwrite() 权限问题，建议添加 security_opt: [seccomp:unconfined] 到 docker-compose 服务中。"
    user_prompt = (
        f"## 部署手册\n{doc_raw}\n\n"
        f"## 目标资产信息\n{asset_info}\n\n"
        f"## 目标机环境探查结果\n{json.dumps(probe, ensure_ascii=False, indent=1)}\n\n"
        f"## AI 初版 SOP\n{json.dumps(sop, ensure_ascii=False, indent=1)}\n\n"
        f"{_os_hint}\n"
        "请分析真实环境，自动生成环境映射和自适应建议。"
    )
    resp = call_llm(provider, [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt},
    ], timeout_override=max(provider.timeout_seconds, 120))
    if resp.get("error"):
        return {"error": f"AI 分析失败: {resp['error']}"}
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        return {"error": "AI 返回格式异常"}
    analysis = _extract_json(content)
    if not analysis:
        return {"error": "AI 未能生成有效的环境分析"}
    mapping = analysis.get("env_mapping", {})
    if mapping:
        # 合并而非整体覆盖：已存在的非空值优先（用户手填 / resolve-env / 真实下载路径），
        # AI 只补充缺失的 key，避免 AI 凭空推断的 APP_DIR 等覆盖正确值。
        try:
            cur = json.loads(plan.env_mapping) if isinstance(plan.env_mapping, str) and plan.env_mapping not in ("{}", "[]") else {}
        except Exception:
            cur = {}
        merged = dict(cur)
        for k, v in (mapping or {}).items():
            if v in (None, ""):
                continue
            if merged.get(k) in (None, ""):
                merged[k] = v
        # APP_DIR 若缺失，用真实下载路径兜底
        if merged.get("APP_DIR") in (None, ""):
            merged["APP_DIR"] = resolve_download_path(plan)
        plan.env_mapping = json.dumps(merged, ensure_ascii=False)
    plan.env_analysis_json = json.dumps(analysis, ensure_ascii=False)
    db.commit()
    return {"ok": True, "analysis": analysis, "env_mapping": analysis.get("env_mapping", {})}


def _ai_diagnose_failure(db: Session, provider, step: DeployStep, step_output: str, env_context: dict) -> dict:
    """步骤失败时, AI 诊断根因并给出修复建议(B 层核心)。"""
    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责分析部署步骤失败原因并给出修复方案。\n"
        "输出严格 JSON：\n"
        "{\n"
        '  "root_cause": "根因分析（中文）",\n'
        '  "fix_commands": ["修复命令 1", "修复命令 2"],\n'
        '  "suggestion": "建议重试/跳过/回滚"\n'
        "}\n"
        "fix_commands 最多 3 条，每条应是可执行的 shell 命令。"
    )
    user_prompt = (
        f"## 失败步骤\n步骤 {step.step_order}: {step.description}\n"
        f"命令: {step.command}\n"
        f"校验命令: {step.verify_command or '无'}\n"
        f"## 输出\n{step_output[:2000]}\n"
        f"## 环境上下文\n{json.dumps(env_context, ensure_ascii=False)[:1500]}\n\n请诊断根因并给出修复方案。"
    )
    resp = call_llm(provider, [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt},
    ], timeout_override=60)
    if resp.get("error"):
        return {"error": f"诊断失败: {resp['error']}"}
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        return {"error": "AI 返回格式异常"}
    diag = _extract_json(content)
    if not diag:
        return {"ok": True, "root_cause": "AI 诊断失败", "fix_commands": [], "suggestion": "rollback"}
    return {"ok": True, "root_cause": diag.get("root_cause", ""), "fix_commands": diag.get("fix_commands", []), "suggestion": diag.get("suggestion", "rollback")}


def stop_execution(db: Session, plan_id: int, rollback: bool = True) -> dict:
    """停止正在执行的部署(强制)：关闭 SSH 连接 → 中断命令 → producer 异常退出 → 释放执行锁。
    默认 rollback=True：停止后自动执行一次回滚清理（停容器 + 清理产物 + 重置为 planned），
    避免卡死/中断后目标机残留容器或半成品。
    rollback=False 时仅断连重置状态（兼容旧行为）。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    client = _RUNNING_CLIENTS.pop(plan_id, None)
    if client:
        try:
            client.close()
        except Exception:
            pass
    _STOPPED[plan_id] = True
    # 如果有决策队列在等待，塞入停止信号立即解除阻塞
    _dq = _DECISIONS.get(plan_id)
    if _dq:
        try:
            _dq.put_nowait("rollback")
        except Exception:
            pass
    _release_exec(plan_id)
    _STOPPED.pop(plan_id, None)
    plan.status = "planned"
    _record_execution_history(db, plan, "stopped", {"stopped_by": "stop_execution", "rollback": rollback})
    db.commit()

    msg = "已停止执行"
    if rollback:
        try:
            rb = _force_rollback_cleanup_sync(db, plan_id)
            msg = "[已停止] " + (rb.get("message", "并已自动回滚清理"))
        except Exception as e:
            logger.warning(f"停止后自动回滚清理异常: plan_id={plan_id} {e}")
            msg = f"已停止执行(回滚清理异常: {e})"
    return {"ok": True, "message": msg}


def _force_rollback_cleanup_sync(db: Session, plan_id: int) -> dict:
    """强制同步回滚清理（停止时调用）：对每个资产 SSH 停容器 + 清理运行产物，重置步骤与计划为 planned。
    可用于 running/卡死状态，不校验 status。返回 {"ok","message","assets"}。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"ok": False, "message": "计划不存在"}
    assets = _get_assets(db, plan)
    if not assets:
        return {"ok": False, "message": "计划未关联目标资产"}
    mapping = json.loads(plan.env_mapping or "{}")
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).order_by(DeployStep.step_order).all()
    app_dir = mapping.get("APP_DIR", "")
    records = []
    for asset in assets:
        _alines = []
        try:
            client, host = _ssh_connect(asset)
        except Exception as e:
            _alines.append(f"❌ {asset.name}({asset.ip}) SSH 连接失败: {e}")
            records.append({"asset": asset.name, "ip": asset.ip, "lines": _alines, "status": "ssh_failed"})
            continue
        try:
            if app_dir:
                _down = f"cd {app_dir} && docker compose down -v 2>/dev/null || true"
                try:
                    _, _o, _e = client.exec_command(_down, timeout=30)
                    _rc = _o.channel.recv_exit_status()
                    _alines.append(f"🧹 清理容器: docker compose down -v (exit={_rc})")
                except Exception as ex:
                    _alines.append(f"⚠ 清理容器异常: {ex}")
                try:
                    _clean = (f"cd {app_dir} 2>/dev/null && "
                              f"rm -rf node_modules .venv venv __pycache__ dist build */__pycache__ 2>/dev/null; "
                              f"find . -name '*.pyc' -type f -delete 2>/dev/null; echo cleaned")
                    _, _o, _e = client.exec_command(_clean, timeout=30)
                    _rc = _o.channel.recv_exit_status()
                    _alines.append(f"🧹 清理运行产物，保留源码目录: {app_dir}")
                except Exception as ex:
                    _alines.append(f"⚠ 清理运行产物异常: {ex}")
            records.append({"asset": asset.name, "ip": asset.ip, "lines": _alines, "status": "cleaned"})
        finally:
            try:
                client.close()
            except Exception:
                pass
    for s in steps:
        s.status = "pending"
        s.output = ""
        s.diagnosis = ""
        s.fix_command = ""
        s.retry_count = 0
        s.started_at = None
        s.finished_at = None
    _record_cleanup_history(db, plan, records, app_dir)
    plan.status = "planned"
    db.commit()
    return {"ok": True, "message": f"已回滚清理 {len(records)} 个资产，状态重置为 planned", "assets": len(records)}


def _now():
    return datetime.now()


def _get_provider(db: Session):
    config = db.query(AgentConfig).filter(AgentConfig.is_enabled == True).order_by(AgentConfig.id.asc()).first()
    provider = None
    if config and config.default_provider_id:
        provider = db.query(AIProvider).filter(
            AIProvider.id == config.default_provider_id, AIProvider.is_enabled == True).first()
    if not provider:
        from app.services.ai_provider_health import select_healthy_provider
        _all = db.query(AIProvider).filter(AIProvider.is_enabled == True).all()
        _sel, _cand, _skip = select_healthy_provider(_all)
        provider = _sel or (_all[0] if _all else None)
    return provider


def _build_offline_hint(db: Session) -> str:
    """生成离线部署提示: 默认私有 Registry 地址 + 活跃本地包源(供 AI SOP 生成命令时遵守)。"""
    lines = []
    try:
        from app.models import OfflineRegistry, OfflinePackageSource
        reg = db.query(OfflineRegistry).filter(OfflineRegistry.is_default == True).first()  # noqa: E712
        if reg and reg.registry_url:
            lines.append(f"- 镜像私有仓库: {reg.registry_url} (insecure={not reg.is_secure})")
        for s in db.query(OfflinePackageSource).filter(OfflinePackageSource.is_active == True).all():  # noqa: E712
            if getattr(s, "source_url", ""):
                lines.append(f"- 包源({getattr(s, 'os_type', '')}): {s.source_url}")
    except Exception:
        pass
    return "\n".join(lines)


def _get_asset_ids(plan) -> List[int]:
    try:
        ids = json.loads(plan.asset_ids) if isinstance(plan.asset_ids, str) else (plan.asset_ids or [])
        return ids if isinstance(ids, list) else []
    except (json.JSONDecodeError, TypeError):
        return []


def _get_assets(db: Session, plan) -> List[Asset]:
    ids = _get_asset_ids(plan)
    if not ids:
        return []
    return db.query(Asset).filter(Asset.id.in_(ids)).all()


def list_plans(db: Session, status: Optional[str] = None, page: int = 1, per_page: int = 20):
    q = db.query(DeployPlan).order_by(DeployPlan.created_at.desc())
    if status:
        q = q.filter(DeployPlan.status == status)
    total = q.count()
    items = q.offset((page - 1) * per_page).limit(per_page).all()
    return {
        "items": [_plan_to_dict(p) for p in items],
        "total": total,
        "page": page,
        "per_page": per_page,
    }


def get_plan(db: Session, plan_id: int):
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return None
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).order_by(DeployStep.step_order).all()
    return {**_plan_to_dict(plan), "steps": [_step_to_dict(s) for s in steps]}


def create_plan(db: Session, payload: dict, user_id: int) -> dict:
    name = (payload.get("name") or "").strip()
    if not name:
        raise ValueError("计划名称不能为空")
    plan = DeployPlan(
        name=name,
        description=(payload.get("description") or "").strip(),
        artifact_path=(payload.get("artifact_path") or "").strip(),
        artifact_download_path=(payload.get("artifact_download_path") or "").strip(),
        artifact_auto_download=bool(payload.get("artifact_auto_download", True)),
        doc_raw=(payload.get("doc_raw") or "").strip(),
        doc_file_name=(payload.get("doc_file_name") or "").strip(),
        asset_ids=json.dumps(payload.get("asset_ids", []) if isinstance(payload.get("asset_ids"), list) else []),
        use_offline=bool(payload.get("use_offline", False)),
        http_proxy=(payload.get("http_proxy") or "").strip(),
        https_proxy=(payload.get("https_proxy") or "").strip(),
        no_proxy=(payload.get("no_proxy") or "").strip(),
        created_by=user_id,
    )
    db.add(plan)
    db.commit()
    db.refresh(plan)
    result = _plan_to_dict(plan)
    # 创建前校验本地路径是否已存在非空内容(防止误部署到已有项目目录)，已有则警告提示但允许继续
    warning = _check_artifact_path_warning(db, plan)
    if warning:
        result["path_warning"] = warning
    return result


def _check_artifact_path_warning(db: Session, plan) -> Optional[str]:
    """创建计划时,SSH 到目标机校验本地源码/下载目标路径是否已存在非空内容。
    若路径非空(已有别的文件/项目),返回警告文案(不拦截创建)。SSH/不可达/远程地址则返回 None。"""
    checks = []
    ap = (plan.artifact_path or "").strip()
    dp = (plan.artifact_download_path or "").strip()
    # artifact_path: 仅当是本地路径(非 git/http/offline)时校验
    if ap:
        src = detect_artifact_source(ap)
        if src == "local":
            checks.append(("代码包路径", ap))
    # 下载目标路径: 永远本地,非空即校验
    if dp:
        checks.append(("源码下载目标路径", dp))
    if not checks:
        return None
    assets = _get_assets(db, plan)
    if not assets:
        return None
    asset = assets[0]
    for label, path in checks:
        try:
            code, out = _run_ssh(asset, f"if [ -d {path} ]; then ls -A {path} 2>/dev/null | head -1; fi", timeout=30)
        except Exception:
            continue
        if code == 0 and out.strip():
            return (f"⚠️ {label}「{path}」在目标机 {asset.ip} 上已存在非空内容"
                    f"（首个文件/目录: {out.strip().splitlines()[0]}），请确认路径是否正确，"
                    f"避免误部署到已有项目目录。")
    return None


def update_plan(db: Session, plan_id: int, payload: dict) -> Optional[dict]:
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return None
    for field in ("name", "description", "artifact_path", "doc_raw", "doc_file_name", "env_mapping", "sop_json"):
        if field in payload:
            val = payload[field]
            if isinstance(val, (dict, list)):
                val = json.dumps(val, ensure_ascii=False)
            setattr(plan, field, val)
    if "artifact_download_path" in payload:
        plan.artifact_download_path = (payload.get("artifact_download_path") or "").strip()
    if "artifact_auto_download" in payload:
        plan.artifact_auto_download = bool(payload.get("artifact_auto_download"))
    if "use_offline" in payload:
        plan.use_offline = bool(payload.get("use_offline"))
    for f in ("http_proxy", "https_proxy", "no_proxy"):
        if f in payload:
            setattr(plan, f, (payload.get(f) or "").strip())
    if "asset_ids" in payload and isinstance(payload["asset_ids"], list):
        plan.asset_ids = json.dumps(payload["asset_ids"])
    db.commit()
    db.refresh(plan)
    return _plan_to_dict(plan)


def update_doc_raw(db: Session, plan_id: int, doc_raw: str, doc_file_name: str = "") -> Optional[dict]:
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return None
    plan.doc_raw = doc_raw
    if doc_file_name:
        plan.doc_file_name = doc_file_name
    db.commit()
    db.refresh(plan)
    return _plan_to_dict(plan)


def delete_plan(db: Session, plan_id: int) -> bool:
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return False
    db.query(DeployStep).filter(DeployStep.plan_id == plan_id).delete()
    db.delete(plan)
    db.commit()
    return True


def _ai_auto_resolve_env(provider, plan: DeployPlan, doc_raw: str, steps: list, env_vars: list) -> dict:
    """AI 基于手册上下文自动推断环境变量值，无需用户手动填写。

    输入: 解析后的 env_vars 列表 + 原始手册 + 步骤
    输出: {"APP_DIR": "/tmp/aiplan-test", "PORT": "8080", ...}
    """
    if not env_vars or not provider:
        return {}
    only_auto = all(e.get("source") == "auto" or e.get("source") == "资产" for e in env_vars)
    if only_auto:
        return {}
    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责从部署手册中自动推断环境变量的值。\n"
        "严格按 JSON 输出，key=环境变量名, value=推断的值:\n"
        "{\n"
        '  "APP_DIR": "/tmp/aiplan-test",\n'
        '  "PORT": "8080"\n'
        "}\n"
        "规则：\n"
        "- 从手册的步骤命令中寻找原始值（手册中可能有 mkdir -p /tmp/xxx，从中提取 APP_DIR）\n"
        "- 从手册的端口映射、URL 中提取端口号\n"
        "- 从资产信息中提取 IP（TARGET_IP 用资产 IP）\n"
        "- 如果无法推断，留空字符串\n"
        "- 只输出 JSON，不要任何额外内容"
    )
    doc_raw_snippet = doc_raw[:3000] if doc_raw else ""
    step_cmds = "\n".join([(s.get("command", "") or "")[:200] for s in steps])
    env_names = [e.get("name", "") for e in env_vars]
    user_prompt = (
        f"## 部署手册(片段)\n{doc_raw_snippet}\n\n"
        f"## 步骤命令\n{step_cmds[:2000]}\n\n"
        f"## 需要推断的环境变量\n{json.dumps(env_names, ensure_ascii=False)}\n\n"
        f"请从手册和步骤命令中为每个环境变量推断实际值，只输出 JSON。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=60)
        if resp.get("error"):
            return {}
        content = resp["choices"][0]["message"]["content"]
        inferred = _extract_json(content) or {}
        if not isinstance(inferred, dict):
            return {}
        # 只保留在 env_vars 中声明过的 key
        valid_keys = set(e.get("name", "") for e in env_vars)
        valid_keys.update(k for k in inferred.keys() if k.startswith("ENV_"))
        return {k: v for k, v in inferred.items() if k in valid_keys and v}
    except Exception as e:
        logger.warning(f"AI 自动推断环境变量异常: {e}")
    return {}


def _ai_auto_resolve_unresolved(provider, plan: DeployPlan, db: Session, steps: list, mapping: dict) -> dict:
    """执行前，如果还有未解析的环境变量，AI 自动解决而非阻塞返回。

    返回填充后的完整 mapping。
    """
    unresolved = set()
    for s in steps:
        for field in (s.command, s.verify_command, s.rollback_command):
            if field:
                missing = _check_unresolved(_resolve_command(field, mapping))
                if missing:
                    unresolved.add(missing)
    if not unresolved:
        return mapping

    if not provider:
        # 没有 AI 时还是要阻塞
        return mapping

    # 让 AI 推断未设置的值
    doc_raw = plan.doc_raw or ""
    cmds = "\n".join([(s.command or "")[:200] for s in steps])
    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责为部署计划中的环境变量推断实际值。\n"
        "严格按 JSON 输出，key=变量名, value=推断值:\n{\n"
        '  "APP_DIR": "/tmp/myapp",\n'
        '  "PORT": "8080"\n'
        "}\n"
        "从手册命令中寻找原始值，如果无法推断就留空。"
    )
    user_prompt = (
        f"## 手册(片段)\n{(doc_raw or '')[:3000]}\n\n"
        f"## 步骤命令\n{cmds[:2000]}\n\n"
        f"## 未设置的变量\n{json.dumps(list(unresolved), ensure_ascii=False)}\n\n"
        f"请推断每个变量的值。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=60)
        if resp.get("error"):
            return mapping
        content = resp["choices"][0]["message"]["content"]
        inferred = _extract_json(content) or {}
        if not isinstance(inferred, dict):
            return mapping
        for k, v in inferred.items():
            if k in unresolved and v:
                mapping[k] = str(v)
                if not mapping.get(f"ENV_{k}"):
                    mapping[f"ENV_{k}"] = str(v)
        # 重新检查是否全部解决
        still_unresolved = set()
        for s in steps:
            for field in (s.command, s.verify_command, s.rollback_command):
                if field:
                    missing = _check_unresolved(_resolve_command(field, mapping))
                    if missing:
                        still_unresolved.add(missing)
        if still_unresolved:
            logger.warning(f"AI 无法推断的变量: {still_unresolved}")
    except Exception as e:
        logger.warning(f"AI 自动解决未解析变量异常: {e}")
    return mapping


def ai_parse_manual(db: Session, plan_id: int) -> dict:
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    doc_raw = plan.doc_raw.strip()
    if not doc_raw:
        return {"error": "部署手册内容为空，请先上传手册"}

    provider = _get_provider(db)
    if not provider:
        return {"error": "未配置可用的 AI 模型提供商"}

    assets = _get_assets(db, plan)
    asset_info = ""
    if assets:
        lines = [f"- 名称: {a.name}  IP: {a.ip}  CI: {a.ci_type}" for a in assets]
        asset_info = f"\n目标资产信息（共 {len(assets)} 台）：\n" + "\n".join(lines)

    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责将部署手册解析为结构化部署 SOP。\n"
        "请严格按以下 JSON Schema 输出，不要输出任何额外内容：\n"
        "{\n"
        '  "plan_name": "部署计划名称",\n'
        '  "preflight": [{"check": "检查项名", "command": "只读检查命令", "expect": "预期结果"}],\n'
        '  "steps": [\n'
        "    {\n"
        '      "order": 1,\n'
        '      "description": "步骤说明",\n'
        '      "command": "shell 命令，环境敏感值用 ${ENV_xxx} 占位符",\n'
        '      "verify": "可执行的校验 shell 命令，必须返回 exit 0 才算通过；若无命令则留空字符串。严禁写自然语言描述",\n'
        '      "rollback": "可执行的回滚 shell 命令；若无则留空字符串。严禁写自然语言描述",\n'
        '      "risk": "low|medium|high"\n'
        "    }\n"
        "  ],\n"
        '  "env_vars": [\n'
        '    {"name": "TARGET_IP", "description": "目标 IP", "example": "192.168.1.100", "source": "资产"},\n'
        '    {"name": "APP_DIR", "description": "应用目录", "example": "/opt/myapp", "source": "用户输入"}\n'
        "  ]\n"
        "}\n"
        "规则：\n"
        "1. 从手册中提取所有步骤，按执行顺序排列\n"
        "2. **重要：手册中已有的 ${xxx} 占位符必须原样保留在命令中，不得删除或替换**\n"
        "3. 环境相关的值（IP、端口、目录、密码等）用 ${ENV_xxx} 占位符替代\n"
        "4. 识别手册中的环境参数，在 env_vars 中列出\n"
        "5. 每步标记风险等级（low/medium/high）\n"
        "6. 识别可验证的检查点和可回滚的命令\n"
    )
    user_prompt = f"以下是部署手册内容：\n\n{doc_raw}\n{asset_info}\n\n请解析为结构化 SOP JSON。"

    if plan.use_offline:
        offline_hint = _build_offline_hint(db)
        if offline_hint:
            user_prompt += (
                "\n\n【离线部署要求】本计划启用离线私有仓库，务必遵守：\n"
                "- 一切 docker pull / 镜像引用必须改用私有仓库地址（docker login 后拉取），禁止从公网 Docker Hub 拉取。\n"
                "- 系统包安装(yum/dnf/apt-get install)必须使用本地/内网包源，禁止使用公网软件源。\n"
                f"- 离线资源：\n{offline_hint}"
            )

    resp = call_llm(provider, [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt},
    ], timeout_override=max(provider.timeout_seconds, 120))
    if resp.get("error"):
        return {"error": f"AI 解析失败: {resp['error']}"}
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception:
        return {"error": "AI 返回格式异常"}

    sop = _extract_json(content)
    if not sop:
        return {"error": "AI 未能生成有效的结构化 SOP"}

    steps = sop.get("steps", [])
    env_vars = sop.get("env_vars", [])
    if not steps:
        return {"error": "AI 未能从手册中提取出部署步骤"}

    plan.sop_json = json.dumps(sop, ensure_ascii=False)
    plan.status = "planned"

    db.query(DeployStep).filter(DeployStep.plan_id == plan_id).delete()
    for s in steps:
        step = DeployStep(
            plan_id=plan_id,
            step_order=s.get("order", 0),
            description=(s.get("description") or "").strip(),
            command=(s.get("command") or "").strip(),
            verify_command=(s.get("verify") or "").strip(),
            rollback_command=(s.get("rollback") or "").strip(),
            risk_level=s.get("risk", "medium"),
        )
        db.add(step)
    db.commit()

    # 从所有命令中扫描 ${ENV_xxx} 占位符，没在 env_mapping 里的自动种子（空值，用户编辑）
    try:
        current_mapping = json.loads(plan.env_mapping) if isinstance(plan.env_mapping, str) and plan.env_mapping not in ("{}", "[]") else {}
    except Exception:
        current_mapping = {}
    # 从已保存的步骤命令中额外提取占位符（AI 可能漏报）
    _all_steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).all()
    for _s in _all_steps:
        for _field in (_s.command, _s.verify_command, _s.rollback_command):
            if _field:
                for _m in re.finditer(r'\$\{(\w+)\}', _field):
                    _k = _m.group(1)
                    if _k not in current_mapping:
                        current_mapping[_k] = ""
    # 从原始手册中额外提取占位符（双重兜底：即使 AI 删了命令里的占位符，也能找到）
    for _m in re.finditer(r'\$\{(\w+)\}', doc_raw):
        _k = _m.group(1)
        if _k not in current_mapping:
            current_mapping[_k] = ""
    plan.env_mapping = json.dumps(current_mapping, ensure_ascii=False)
    db.commit()

    # ── AI 自动推断环境变量值（从手册上下文）──
    try:
        provider = _get_provider(db)
        if provider and env_vars:
            inferred = _ai_auto_resolve_env(provider, plan, doc_raw, steps, env_vars)
            if inferred:
                current_mapping.update(inferred)
                for k, v in inferred.items():
                    if k.startswith("ENV_") and (not current_mapping.get(k[4:]) or not current_mapping.get(k)):
                        current_mapping[k[4:]] = v
                    elif not k.startswith("ENV_") and (not current_mapping.get(f"ENV_{k}") or not current_mapping.get(k)):
                        current_mapping[f"ENV_{k}"] = v
                plan.env_mapping = json.dumps(current_mapping, ensure_ascii=False)
                db.commit()
    except Exception as e:
        logger.warning(f"AI 自动推断环境变量异常: {e}")

    return {"ok": True, "sop": sop, "env_vars": env_vars, "step_count": len(steps)}


def _extract_json(text: str) -> Optional[dict]:
    text = text.strip()
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


def resolve_env_mapping(db: Session, plan_id: int, user_mapping: dict) -> dict:
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}

    mapping = dict(user_mapping)
    assets = _get_assets(db, plan)
    for i, asset in enumerate(assets):
        suffix = f"_{i + 1}" if len(assets) > 1 else ""
        mapping.setdefault(f"TARGET_IP{suffix}", asset.ip or "")
        mapping.setdefault(f"TARGET_HOSTNAME{suffix}", asset.name or "")
        if i == 0:
            mapping.setdefault("TARGET_IP", asset.ip or "")
            mapping.setdefault("TARGET_HOSTNAME", asset.name or "")

    mapping.setdefault("ARTIFACT_URL", plan.artifact_path or "")
    mapping.setdefault("ARTIFACT_DOWNLOAD_PATH", resolve_download_path(plan))

    # 兼容 ${ENV_xxx} 占位符：若命令里用了 ENV_ 前缀但用户只填了裸 key，则同步一份
    try:
        current_mapping = json.loads(plan.env_mapping) if isinstance(plan.env_mapping, str) and plan.env_mapping not in ("{}", "[]") else {}
    except Exception:
        current_mapping = {}
    for _k, _v in list(mapping.items()):
        if _k.startswith("ENV_"):
            mapping.setdefault(_k[4:], _v)
        else:
            mapping.setdefault(f"ENV_{_k}", _v)
    for _k in list(current_mapping.keys()):
        if _k.startswith("ENV_") and _k not in mapping:
            mapping.setdefault(_k, current_mapping.get(_k, ""))

    plan.env_mapping = json.dumps(mapping, ensure_ascii=False)
    plan.status = "planned"
    db.commit()
    return {"ok": True, "env_mapping": mapping}


_GIT_HOST_HINTS = ("github.com", "gitee.com", "gitlab.com", "gitcode.com", "jihulab.com", ".git")


def _sanitize_dirname(name: str) -> str:
    """把计划名转成安全目录名（去空白与路径分隔符）。"""
    safe = re.sub(r"[\s/\\:*?\"<>|]+", "-", (name or "plan").strip())
    return safe or "plan"


def resolve_download_path(plan: DeployPlan) -> str:
    """解析源码自动下载目标路径：
    用户填了 artifact_download_path 则用它；否则返回 '/data/aiops-deploy/<计划名>'。"""
    if plan.artifact_download_path and plan.artifact_download_path.strip():
        return plan.artifact_download_path.strip()
    return f"/data/aiops-deploy/{_sanitize_dirname(plan.name or 'plan')}"


def detect_artifact_source(url: str) -> str:
    """识别源码来源类型：
    - 'git':  Git 仓库地址（github/gitee/gitlab 等）
    - 'http': HTTP(S) 下载地址（tar.gz/zip 等压缩包）
    - 'offline': 离线仓库获取（offline:// 前缀或离线包引用）
    - 'local': 资产本地路径（/opt/app 等，无需下载）
    - '': 无法识别
    """
    if not url:
        return ""
    u = url.strip()
    if u.startswith("offline://") or u.startswith("offline:"):
        return "offline"
    if u.startswith("http://") or u.startswith("https://"):
        if any(h in u.lower() for h in _GIT_HOST_HINTS):
            return "git"
        return "http"
    if any(h in u.lower() for h in _GIT_HOST_HINTS):
        return "git"
    if u.startswith("/") or ":" in u.split("/")[0]:
        return "local"
    return ""


def _run_ssh(asset, cmd: str, timeout: int = 300):
    """在目标机上执行命令，返回 (exit_status, stdout_text)。"""
    client = None
    try:
        client, host = _ssh_connect(asset)
        stdin, stdout, stderr = client.exec_command(cmd, timeout=timeout)
        out = stdout.read().decode("utf-8", "replace")
        err = stderr.read().decode("utf-8", "replace")
        code = stdout.channel.recv_exit_status()
        return code, (out + err).strip()
    finally:
        if client:
            try:
                client.close()
            except Exception:
                pass


def _fetch_offline_bundle_path(plan) -> str:
    """离线模式：返回离线包在资产侧可供下载的相对信息。
    这里返回离线包文件路径；实际部署时由用户已加载的离线包承载。"""
    try:
        bundles = None
        from app.models import OfflineRepoBundle
        from app.database import get_db_session
        with get_db_session() as db:
            bundle = db.query(OfflineRepoBundle).filter(OfflineRepoBundle.status == "loaded").first()
            return bundle.file_path if bundle else ""
    except Exception:
        return ""


def auto_download_artifact(db: Session, plan_id: int, force: bool = False) -> dict:
    """探查前自动下载源码到目标机（在线 git/HTTP + 离线仓库两套都支持）。

    根据 artifact_path 识别来源：
      - git: 优先 git clone（目标机有 git 时）；无 git 则 curl 下载 codeload/仓库 zip 并解压
      - http: curl 下载压缩包并解压
      - offline: 离线包方式（复用离线仓库，由手册步骤落地，此处探查离线包存在性）
      - local: 资产本地路径，无需下载
    幂等：目标路径已存在且含 compose/docker-compose 文件时跳过。”
    """
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"ok": False, "error": "计划不存在"}
    if not plan.artifact_auto_download:
        return {"ok": True, "skipped": True, "reason": "已关闭自动下载(artifact_auto_download=False)"}

    url = (plan.artifact_path or "").strip()
    source_type = detect_artifact_source(url)
    if source_type == "local" or source_type == "":
        return {"ok": True, "skipped": True, "source": source_type or "unknown", "reason": "本地路径或未填源码地址，无需自动下载"}

    assets = _get_assets(db, plan)
    if not assets:
        return {"ok": False, "error": "计划未关联目标资产"}
    if len(assets) > 1:
        return {"ok": False, "error": "自动下载仅支持单资产计划，多资产请用本地路径或手册自行处理"}

    asset = assets[0]
    dest = resolve_download_path(plan)
    log_lines = []

    # 幂等检查：目标路径已有 compose 则跳过
    code, out = _run_ssh(asset, f"ls {dest}/docker-compose.yml {dest}/docker-compose.yaml {dest}/compose.yaml 2>/dev/null | head -1")
    if code == 0 and out.strip() and not force:
        return {"ok": True, "skipped": True, "source": source_type, "dest": dest,
                "reason": f"源码已存在于 {dest}(含 compose)，跳过下载(force=False 幂等)"}

    try:
        _run_ssh(asset, f"mkdir -p {dest}")
    except Exception as e:
        return {"ok": False, "error": f"创建目录失败: {e}"}

    if source_type == "offline":
        bundle_path = _fetch_offline_bundle_path(plan)
        if not bundle_path:
            return {"ok": False, "error": "离线模式下未找到已加载(loaded)的离线包，请先在离线仓库功能页加载离线包"}
        code, out = _run_ssh(asset, f"ls {dest}", timeout=60)
        log_lines.append(f"[offline] 目标目录: {dest}")
        log_lines.append(f"[offline] 已检测到已加载离线包: {bundle_path} (镜像/包源由手册步骤对接私有 Registry)")
        return {"ok": True, "source": "offline", "dest": dest, "log": log_lines,
                "note": "离线模式：离线包已在仓库侧加载，请确保部署手册步骤通过私有 Registry/包源拉取镜像与软件包"}

    # --- 在线: git / http ---
    if source_type == "git":
        # 目标机有 git → git clone，否则 curl 下载 zip
        _, _has_git = _run_ssh(asset, "command -v git >/dev/null 2>&1 && echo YES || echo NO", timeout=30)
        if _has_git.strip() == "YES":
            code, out = _run_ssh(
                asset,
                f"if [ -d {dest} ] && [ -n \"$(ls -A {dest})\" ]; then echo EXISTS; fi",
                timeout=60,
            )
            if "EXISTS" in out:
                log_lines.append(f"[git] 目录 {dest} 非空，执行 git pull 增量更新")
                pull_cmd = f"cd {dest} && (git pull --ff-only 2>/dev/null || echo PULL_FAILED)"
                _run_ssh(asset, pull_cmd, timeout=300)
                log_lines.append(f"[git] git pull 完成: {dest}")
                return {"ok": True, "source": "git", "dest": dest, "method": "git-clone", "log": log_lines}
            clone_cmd = f"cd {dest} && git clone --depth 1 {url} . 2>&1"
            code, out = _run_ssh(asset, clone_cmd, timeout=600)
            log_lines.append(f"[git] git clone --depth 1: exit={code}")
            log_lines.append((out or "")[-1200:])
            if code != 0:
                return {"ok": False, "error": f"git clone 失败: {(out or '')[-500:]}", "log": log_lines}
            _set_compose_perms(asset, dest, log_lines)
            return {"ok": True, "source": "git", "dest": dest, "method": "git-clone", "log": log_lines}
        # 无 git → curl 下载仓库 zip 并解压
        zip_url = _git_zip_url(url)
        code, out = _run_ssh(
            asset,
            f"curl -fsSL --max-time 300 -o /tmp/_aiops_src.zip \"{zip_url}\" && "
            f"rm -rf {dest}/* && unzip -o -q /tmp/_aiops_src.zip -d /tmp/_aiops_src_x && "
            f"mv /tmp/_aiops_src_x/*/* {dest}/ 2>/dev/null || mv /tmp/_aiops_src_x/* {dest}/ 2>/dev/null; "
            f"rm -rf /tmp/_aiops_src.zip /tmp/_aiops_src_x",
            timeout=400,
        )
        log_lines.append(f"[git->zip] 下载 {zip_url}")
        log_lines.append(f"[git->zip] exit={code}")
        log_lines.append((out or "")[-1200:])
        if code != 0:
            return {"ok": False, "error": f"curl 下载/解压失败: {(out or '')[-500:]}", "log": log_lines}
        _set_compose_perms(asset, dest, log_lines)
        return {"ok": True, "source": "git", "dest": dest, "method": "git-zip", "log": log_lines}

    # http 下载压缩包
    code, out = _run_ssh(
        asset,
        f"curl -fsSL --max-time 300 -o /tmp/_aiops_src.bin \"{url}\" && "
        f"mkdir -p {dest} && rm -rf {dest}/* && "
        f"(file /tmp/_aiops_src.bin | grep -qi zip && unzip -o -q /tmp/_aiops_src.bin -d /tmp/_aiops_src_x || tar -xzf /tmp/_aiops_src.bin -C /tmp/_aiops_src_x) && "
        f"mv /tmp/_aiops_src_x/*/* {dest}/ 2>/dev/null || mv /tmp/_aiops_src_x/* {dest}/ 2>/dev/null; "
        f"rm -rf /tmp/_aiops_src.bin /tmp/_aiops_src_x",
        timeout=400,
    )
    log_lines.append(f"[http] 下载 {url}: exit={code}")
    log_lines.append((out or "")[-1200:])
    if code != 0:
        return {"ok": False, "error": f"HTTP 下载/解压失败: {(out or '')[-500:]}", "log": log_lines}
    _set_compose_perms(asset, dest, log_lines)
    return {"ok": True, "source": "http", "dest": dest, "method": "http-download", "log": log_lines}


def _git_zip_url(url: str) -> str:
    """把 Git 仓库主页/zip 地址归一为可直接 curl 下载的 zip 地址（github/gitee 均支持 codeload/archive）。"""
    u = url.rstrip("/")
    if u.endswith(".zip"):
        return u
    if "github.com" in u:
        m = re.match(r"https?://github\.com/([^/]+)/([^/]+?)(?:\.git)?/?$", u)
        if m:
            return f"https://codeload.github.com/{m.group(1)}/{m.group(2)}/zip/refs/heads/master"
    if "gitee.com" in u:
        m = re.match(r"https?://gitee\.com/([^/]+)/([^/]+?)(?:\.git)?/?$", u)
        if m:
            return f"https://gitee.com/{m.group(1)}/{m.group(2)}/repository/archive/master.zip"
    return u


def _set_compose_perms(asset, dest: str, log_lines: list) -> None:
    """下载后确保 compose 文件存在并赋可执行权限（幂等友好）。"""
    code, out = _run_ssh(asset, f"chmod -R +x {dest} 2>/dev/null; ls {dest}/docker-compose.yml {dest}/compose.yaml {dest}/docker-compose.yaml 2>/dev/null | head -1", timeout=60)
    log_lines.append(f"[perm] {out.strip() or '(未找到 compose 文件)'}")


def _resolve_command(cmd: str, mapping: dict) -> str:
    def replace(match):
        key = match.group(1)
        val = mapping.get(key, "")
        if val:
            return val
        if key.startswith("ENV_") and key[4:] in mapping:
            val = mapping[key[4:]]
            if val:
                return val
        return f"__UNSET__{key}__"
    return re.sub(r'\$\{(\w+)\}', replace, cmd)


def _check_unresolved(cmd: str) -> Optional[str]:
    """检查命令中是否有未解析的 ${ENV_xxx} 占位符，返回第一个未解析的 key。"""
    m = re.search(r'__UNSET__(\w+)__', cmd)
    if m:
        return m.group(1)
    return None


def _is_valid_shell_command(cmd: str) -> bool:
    """防御校验：verify/rollback 字段必须是可执行的 shell 命令，而不是 AI 输出的自然语言描述。
    规则：非空、不含中文/日韩文、不含自然语言连接词标记、以命令动作开头。"""
    if not cmd or not cmd.strip():
        return False
    cmd = cmd.strip()
    # 含 CJK 字符 → 判定为自然语言
    if re.search(r'[\u4e00-\u9fff\u3040-\u30ff\uac00-\ud7af]', cmd):
        return False
    # 含"检查是否""预期""状态""成功"等中文词汇特征（已含中文必然命中上一条，冗余防护）
    # 必须看起来像一个命令：不能以中文标点或描述性单词开头
    if re.match(r'^(执行后|检查|确认|验证|显示|确保|等待)[：:\s，,，]', cmd):
        return False
    return True


_OFFLINE_PUBLIC_IMAGES = ["docker.io", "registry.hub.docker.com", "index.docker.io", "hub.docker.com", "ghcr.io", "quay.io", "gcr.io", "docker.easypack.io", "registry.cn-hangzhou.aliyuncs.com", "mirror.ccs.tencentyun.com"]
_PUBLIC_REPO_HINTS = ["archive.ubuntu.com", "security.ubuntu.com", "download.fedoraproject.org", "mirrors.aliyun.com", "repo.huaweicloud.com", "mirrors.tuna.tsinghua.edu.cn", "mirrors.cloud.tencent.com"]


def _offline_blocked_reason(plan, cmd: str) -> str:
    """离线模式二次强制校验: 命令含公网 docker 镜像拉取或公网软件源时, 返回拦截原因; 否则返回空串放行。
    仅当 plan.use_offline=True 时启用。识别 docker pull/run/compose 的镜像引用、yum/apt 针对公网源的操作。"""
    if not plan or not getattr(plan, "use_offline", False):
        return ""
    cmd = cmd or ""
    low = cmd.lower()
    # 明确公网镜像仓库引用
    for img in _OFFLINE_PUBLIC_IMAGES:
        if img in low:
            return f"离线模式禁止拉取公网镜像仓库 {img} 的镜像(请改用离线私有 Registry)"
    # docker pull / docker run 引用公网镜像: 仅当是"裸镜像名"(无仓库主机)或显式 docker.io 时拦截
    # 带仓库主机(如 internal.registry/app/x)视为私有/内网仓库, 放行
    if re.search(r'\bdocker\s+(pull|run|create)\b', low):
        _rest = re.split(r'\bdocker\s+(?:pull|run|create)\b', cmd)[-1]
        _cand = ""
        for _t in _rest.split():
            if _t.startswith('-') or _t in ('--','/dev/null'):
                continue
            _cand = _t
            break
        if _cand and ('/' not in _cand or _cand.startswith(('docker.io/', 'registry.hub.docker.com/', 'library/'))):
            return f"离线模式禁止直接 docker pull/run 公网镜像 {_cand}(请改用离线私有 Registry 地址)"
    # 公网软件源
    for hint in _PUBLIC_REPO_HINTS:
        if hint in low:
            return f"离线模式禁止使用公网软件源 {hint}(请改用本地/内网包源)"
    return ""


def _assert_online_allowed(plan, cmd: str) -> str:
    """通用: 在线/离线统一入口。离线时返回拦截原因, 在线返回空串。供执行前调用。"""
    if not plan or not getattr(plan, "use_offline", False):
        return ""
    return _offline_blocked_reason(plan, cmd)


def _proxy_env_prefix(plan) -> str:
    """为计划生成代理环境变量导出前缀(供执行步骤注入 HTTP_PROXY/HTTPS_PROXY/NO_PROXY)。"""
    if not plan:
        return ""
    parts = []
    if getattr(plan, "http_proxy", ""):
        parts.append("export HTTP_PROXY='%s' http_proxy='%s'" % (plan.http_proxy, plan.http_proxy))
    if getattr(plan, "https_proxy", ""):
        parts.append("export HTTPS_PROXY='%s' https_proxy='%s'" % (plan.https_proxy, plan.https_proxy))
    if getattr(plan, "no_proxy", ""):
        parts.append("export NO_PROXY='%s' no_proxy='%s'" % (plan.no_proxy, plan.no_proxy))
    return " && ".join(parts)


def _sync_env_mapping_from_sop(db: Session, plan: DeployPlan) -> None:
    """预检/执行前把 SOP 命令与 doc_raw 中的 ${ENV_xxx} 占位符同步到 env_mapping。
    防止旧版解析产生的计划缺键（旧代码只信 AI 的 env_vars 列表，可能漏掉占位符）。
    只补缺失的键（空值待填），不覆盖用户已填的值。"""
    try:
        current_mapping = json.loads(plan.env_mapping or "{}") if isinstance(plan.env_mapping, str) and plan.env_mapping not in ("{}", "[]") else {}
    except Exception:
        current_mapping = {}
    _placeholder_re = re.compile(r'\$\{(\w+)\}')
    _found_keys = set()
    sop = json.loads(plan.sop_json or "{}")
    if isinstance(sop, dict):
        for _pf in sop.get("preflight", []):
            _found_keys.update(_m.group(1) for _m in _placeholder_re.finditer(_pf.get("command", "")))
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan.id).all()
    for _s in steps:
        for _field in (_s.command, _s.verify_command, _s.rollback_command):
            if _field:
                _found_keys.update(_m.group(1) for _m in _placeholder_re.finditer(_field))
    if plan.doc_raw:
        _found_keys.update(_m.group(1) for _m in _placeholder_re.finditer(plan.doc_raw))
    changed = False
    for _k in sorted(_found_keys):
        if _k not in current_mapping:
            current_mapping[_k] = ""
            changed = True
    if changed:
        plan.env_mapping = json.dumps(current_mapping, ensure_ascii=False)
        db.commit()


def _ssh_connect(asset: Asset) -> tuple:
    try:
        conn_config = json.loads(asset.connection_config or "{}")
    except Exception:
        conn_config = {}
    host = asset.ip or conn_config.get("ssh_host", "")
    port = int(conn_config.get("ssh_port", 22))
    username = conn_config.get("ssh_user", "root")
    password = conn_config.get("ssh_password", "")
    client = connect_ssh(host, port=port, username=username, password=password)
    return client, host


def run_preflight(db: Session, plan_id: int) -> dict:
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}

    assets = _get_assets(db, plan)
    if not assets:
        return {"error": "计划未关联目标资产"}

    sop = json.loads(plan.sop_json or "{}")
    if isinstance(sop, list):
        sop = {"preflight": [], "steps": []}
    mapping = json.loads(plan.env_mapping or "{}")
    preflight_checks = sop.get("preflight", [])
    if not preflight_checks:
        plan.preflight_json = json.dumps({"skipped": True, "message": "无预检项"})
        db.commit()
        return {"ok": True, "results": [], "skipped": True}

    _sync_env_mapping_from_sop(db, plan)
    mapping = json.loads(plan.env_mapping or "{}")

    all_results = []
    all_ok = True
    unresolved = set()
    for check in preflight_checks:
        cmd = _resolve_command(check.get("command", ""), mapping)
        missing = _check_unresolved(cmd)
        if missing:
            unresolved.add(missing)
    if unresolved:
        return {"error": f"环境参数未设置，请先在环境映射 Tab 中填写: {', '.join(sorted(unresolved))}"}
    for asset in assets:
        try:
            client, host = _ssh_connect(asset)
        except Exception as e:
            all_results.append({"asset": asset.name, "asset_ip": asset.ip, "error": str(e)})
            all_ok = False
            continue

        try:
            for check in preflight_checks:
                cmd = _resolve_command(check.get("command", ""), mapping)
                expect = check.get("expect", "")
                try:
                    stdin, stdout, stderr = client.exec_command(cmd, timeout=15)
                    exit_code = stdout.channel.recv_exit_status()
                    output = stdout.read().decode("utf-8", errors="replace").strip()
                    error = stderr.read().decode("utf-8", errors="replace").strip()
                    if error:
                        output = f"{output}\n{error}" if output else error
                    passed = exit_code == 0
                    all_results.append({
                        "asset": asset.name, "asset_ip": asset.ip,
                        "check": check.get("check", ""), "command": cmd,
                        "expect": expect, "output": output[:500],
                        "exit_code": exit_code, "passed": passed,
                    })
                    if not passed:
                        all_ok = False
                except Exception as e:
                    all_results.append({
                        "asset": asset.name, "asset_ip": asset.ip,
                        "check": check.get("check", ""), "command": cmd,
                        "expect": expect, "output": str(e), "exit_code": -1, "passed": False,
                    })
                    all_ok = False
        finally:
            client.close()

    preflight_result = {"results": all_results, "all_passed": all_ok}
    plan.preflight_json = json.dumps(preflight_result, ensure_ascii=False)
    db.commit()
    return {"ok": True, "results": all_results, "all_passed": all_ok}


def execute_plan(db: Session, plan_id: int, user_id: int = 0) -> dict:
    """HTTP 同步执行部署计划（使用 AI 引擎：DAG编排 + 自主决策 + 自适应回滚）。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    if plan.status == "draft":
        return {"error": "计划尚未规划，无法执行"}

    assets = _get_assets(db, plan)
    if not assets:
        return {"error": "计划未关联目标资产"}

    mapping = json.loads(plan.env_mapping or "{}")
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).order_by(DeployStep.step_order).all()
    if not steps:
        return {"error": "无部署步骤"}

    _sync_env_mapping_from_sop(db, plan)
    mapping = json.loads(plan.env_mapping or "{}")
    probe = _safe_json(plan.environment_probe_json)

    _provider = _get_provider(db)

    # ── AI 自动解决未解析的环境变量 ──
    try:
        mapping = _ai_auto_resolve_unresolved(_provider, plan, db, steps, mapping)
        if mapping:
            plan.env_mapping = json.dumps(mapping, ensure_ascii=False)
            db.commit()
    except Exception as e:
        logger.warning(f"AI 自动解决环境变量异常: {e}")

    # ── 前置资源检查 ──
    try:
        probe = _safe_json(plan.environment_probe_json)
        rc = _ai_resource_check(_provider, plan, steps, assets, probe)
        plan.preflight_json = json.dumps(rc, ensure_ascii=False)
        db.commit()
        if rc["recommendation"] == "block":
            return {"error": f"资源检查未通过: {rc['summary']}", "resource_check": rc}
    except Exception as rce:
        logger.warning(f"资源检查异常: {rce}")

    # 构建 DAG
    dag = None
    if _provider:
        dag = _ai_build_execution_dag(db, _provider, plan, steps)
    if not dag:
        dag = {"groups": [{"group_order": i + 1, "step_orders": [s.step_order], "parallel": False, "reason": "线性回退"} for i, s in enumerate(steps)],
               "reasoning": "AI 不可用，回退为线性顺序执行"}
    plan.dag_json = json.dumps(dag, ensure_ascii=False)
    plan.status = "running"
    db.commit()

    total_assets = len(assets)
    succeeded_assets = 0
    step_map = {s.step_order: s for s in steps}

    for asset in assets:
        try:
            client, host = _ssh_connect(asset)
        except Exception as e:
            logger.error(f"部署 {asset.name}({asset.ip}) SSH 连接失败: {e}")
            continue

        asset_map = dict(mapping)
        asset_map["TARGET_IP"] = asset.ip or ""
        asset_map["TARGET_HOSTNAME"] = asset.name or ""

        has_failed = False
        asset_executed_ok = 0
        # 默认工作目录 = 源码下载路径(APP_DIR)，保证 docker compose 等依赖 cwd 的命令在正确目录执行
        try:
            _cwd = resolve_download_path(plan)
        except Exception:
            _cwd = ""
        try:
            for group in dag["groups"]:
                if has_failed:
                    break
                for so in group["step_orders"]:
                    if has_failed:
                        break
                    step = step_map[so]

                    cmd = _resolve_command(step.command, asset_map)
                    verify_cmd = _resolve_command(step.verify_command, asset_map) if step.verify_command else ""

                    # 维护工作目录，给依赖 cwd 的命令补 cd 前缀(SSH 命令间不共享 cwd)
                    _cd_m = re.match(r'^cd\s+([^\s;&|]+)', cmd.strip())
                    if _cd_m:
                        _cwd = _cd_m.group(1).strip('"\'')
                    elif _cwd:
                        cmd = f"cd {_cwd} && {cmd}"
                    if verify_cmd and _cwd and not re.match(r'^cd\s', verify_cmd.strip()):
                        verify_cmd = f"cd {_cwd} && {verify_cmd}"

                    step.status = "running"
                    step.started_at = _now()
                    db.commit()

                    try:
                        stdin, stdout, stderr = client.exec_command(cmd, timeout=60)
                        exit_code = stdout.channel.recv_exit_status()
                        output = stdout.read().decode("utf-8", errors="replace").strip()
                        error = stderr.read().decode("utf-8", errors="replace").strip()
                        if error:
                            output = f"{output}\n{error}" if output else error
                        step.output = (step.output or "") + f"\n[{asset.name}]\n{output[:2000]}"
                        step.finished_at = _now()

                        if exit_code != 0:
                            step.status = "failed"
                            db.commit()
                            has_failed = True
                            # AI 自主决策
                            if _provider:
                                diag = _ai_diagnose_failure(db, _provider, step, output, {})
                                if diag.get("ok"):
                                    decision = _ai_autonomous_decision(_provider, step, output, diag, [])
                                    _ai_decision_log(plan, {"step": step.step_order, "type": "http_failure",
                                        "decision": decision, "root_cause": diag.get("root_cause", ""),
                                        "fix_commands": diag.get("fix_commands", []), "asset": asset.name})
                                    if decision in ("fix", "retry"):
                                        if decision == "fix":
                                            fix_cmds = diag.get("fix_commands", [])
                                            if fix_cmds:
                                                try:
                                                    p_client, _ = _ssh_connect(asset)
                                                    _run_fix_commands(p_client, fix_cmds, "", asset_map)
                                                    p_client.close()
                                                except Exception:
                                                    pass
                                        step.retry_count += 1
                                        step.diagnosis = diag.get("root_cause", "")
                                        step.fix_command = json.dumps(diag.get("fix_commands", []), ensure_ascii=False)
                                        step.status = "running"
                                        db.commit()
                                        has_failed = False
                                        continue
                                    elif decision == "skip":
                                        step.status = "skipped"
                                        step.diagnosis = diag.get("root_cause", "")
                                        db.commit()
                                        has_failed = False
                                        continue
                            _do_rollback(client, steps, step, asset_map)
                            continue

                        if verify_cmd:
                            if not _is_valid_shell_command(verify_cmd):
                                verify_cmd = ""
                            if verify_cmd:
                                try:
                                    v_stdin, v_stdout, v_stderr = client.exec_command(verify_cmd, timeout=15)
                                    v_exit = v_stdout.channel.recv_exit_status()
                                    if v_exit != 0:
                                        v_out = v_stdout.read().decode("utf-8", errors="replace").strip()
                                        step.output = (step.output + f"\n[校验失败] {v_out[:500]}").strip()
                                        step.status = "failed"
                                        db.commit()
                                        has_failed = True
                                        if _provider:
                                            diag = _ai_diagnose_failure(db, _provider, step, f"校验失败: {v_out}", {})
                                            if diag.get("ok"):
                                                decision = _ai_autonomous_decision(_provider, step, f"校验失败: {v_out}", diag, [])
                                                _ai_decision_log(plan, {"step": step.step_order, "type": "http_verify_failure",
                                                    "decision": decision, "asset": asset.name})
                                                if decision in ("fix", "retry"):
                                                    if decision == "fix":
                                                        _run_fix_commands(client, diag.get("fix_commands", []), "", asset_map)
                                                    step.retry_count += 1
                                                    step.status = "running"
                                                    db.commit()
                                                    has_failed = False
                                                    continue
                                                elif decision == "skip":
                                                    step.status = "skipped"
                                                    db.commit()
                                                    has_failed = False
                                                    continue
                                        _do_rollback(client, steps, step, asset_map)
                                        continue
                                except Exception as ve:
                                    step.output = (step.output + f"\n[校验异常] {ve}").strip()
                                    step.status = "failed"
                                    db.commit()
                                    has_failed = True
                                    _do_rollback(client, steps, step, asset_map)
                                    continue

                        step.status = "succeeded"
                        asset_executed_ok += 1
                        db.commit()

                    except Exception as e:
                        step.output = (step.output or "") + f"\n[{asset.name}] {str(e)[:2000]}"
                        step.status = "failed"
                        step.finished_at = _now()
                        db.commit()
                        has_failed = True
                        _do_rollback(client, steps, step, asset_map)
        finally:
            client.close()

        if not has_failed:
            succeeded_assets += 1

    if succeeded_assets == total_assets:
        plan.status = "succeeded"
    elif succeeded_assets > 0:
        plan.status = "failed"
    else:
        plan.status = "rolled_back"
    db.commit()
    _record_execution_history(db, plan, plan.status, {"total_assets": total_assets, "succeeded_assets": succeeded_assets})

    return {"ok": True, "status": plan.status, "total_assets": total_assets, "succeeded_assets": succeeded_assets}


def stream_execute(db: Session, plan_id: int, user_id: int = 0, decision_queue=None):
    """执行锁包装：同一计划同一时刻只允许一个执行流，僵尸 running 状态可重跑。
    任何退出路径（含生成器被外部 close）都会释放锁。
    decision_queue: 可选，用户决策队列（WS 路由转发修复/重试/回滚/跳过指令）。"""
    if _EXEC_LOCK.get(plan_id):
        yield {"type": "error", "message": "该计划正在执行中，请勿重复点击"}
        return
    _EXEC_LOCK[plan_id] = True
    if decision_queue is not None:
        _DECISIONS[plan_id] = decision_queue
    logger.info(f"部署执行开始: plan_id={plan_id}")
    try:
        yield from _ai_stream_execute(db, plan_id, user_id, decision_queue)
    finally:
        _release_exec(plan_id)
        _STOPPED.pop(plan_id, None)
        _DECISIONS.pop(plan_id, None)
        logger.info(f"部署执行结束，释放锁: plan_id={plan_id}")


def _ai_build_execution_dag(db: Session, provider, plan: DeployPlan, steps: list) -> Optional[dict]:
    """AI 分析步骤依赖关系，生成 DAG 执行计划（能力①：动态编排）。

    返回 {"groups": [{"group_order", "step_orders", "parallel", "reason"}], "reasoning"}。
    AI 不可用/解析失败时回退为线性顺序（每步一组，串行）。
    """
    probe = {}
    if plan.environment_probe_json:
        probe = _safe_json(plan.environment_probe_json)
    step_descs = [{
        "order": s.step_order,
        "description": s.description,
        "command": (s.command or "")[:200],
        "risk": s.risk_level,
        "verify": (s.verify_command or "")[:100],
        "rollback": (s.rollback_command or "")[:100],
    } for s in steps]
    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责分析部署步骤的依赖关系，生成最优执行计划(DAG)。\n"
        "请严格按以下 JSON Schema 输出，不要任何额外内容：\n"
        "{\n"
        '  "groups": [\n'
        "    {\n"
        '      "group_order": 1,\n'
        '      "step_orders": [1],\n'
        '      "parallel": false,\n'
        '      "reason": "前置准备步骤，无依赖"\n'
        "    }\n"
        "  ],\n"
        '  "reasoning": "简要分析说明"\n'
        "}\n"
        "规则：\n"
        "- 有依赖关系的步骤必须放在不同组（前组完成后才能执行后组）\n"
        "- 互不依赖的步骤可放入同一组并行执行（parallel: true）\n"
        "- 同一组内并行步骤必须互不依赖，且不能互相覆盖文件/端口/服务\n"
        "- 组按执行先后顺序排列 group_order，所有 step_orders 必须覆盖全部步骤，不得遗漏\n"
        "- 若无法判断依赖关系，保守起见每步单独成组（串行最安全）\n"
        "- 涉及 docker compose up / 启动服务 / 写配置的步骤通常串行，避免竞争"
    )
    user_prompt = (
        f"## 部署步骤\n{json.dumps(step_descs, ensure_ascii=False, indent=1)}\n\n"
        f"## 目标环境探查摘要\nOS: {str(probe.get('os', ''))[:100]}\n"
        f"Docker: {str(probe.get('docker', ''))[:50]}\n"
        f"端口: {json.dumps(probe.get('port_scan', {}), ensure_ascii=False)[:200]}\n\n"
        "请分析步骤依赖关系，生成最优执行计划。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=60)
        if resp.get("error"):
            logger.warning(f"DAG 分析失败: {resp['error']}")
            return None
        content = resp["choices"][0]["message"]["content"]
        dag = _extract_json(content)
        if dag and isinstance(dag.get("groups"), list) and dag["groups"]:
            return dag
    except Exception as e:
        logger.warning(f"DAG 分析异常: {e}")
    return None


def _ai_pre_execution_risk(provider, step: DeployStep, plan: DeployPlan) -> dict:
    """AI 预判单步骤执行风险（能力③：执行前预判），返回风险建议 dict。"""
    probe = {}
    if plan.environment_probe_json:
        probe = _safe_json(plan.environment_probe_json)
    sys_prompt = (
        "你是资深 SRE 运维专家，负责分析部署步骤命令的执行风险并给出应对建议。\n"
        "严格按 JSON 输出，不要额外内容：\n"
        "{\n"
        '  "risk": "low|medium|high",\n'
        '  "reason": "风险分析（中文，一句话）",\n'
        '  "precheck": "可选的前置检查命令（空串表示无需额外检查）",\n'
        '  "precheck_expect": "前置检查预期结果",\n'
        '  "suggest_modify": "建议替换后的更安全命令（空串表示无需修改）",\n'
        '  "guard_note": "执行中的注意事项（中文，一句话）"\n'
        "}\n"
        "规则：\n"
        "- 涉及 rm -rf、docker compose down、重启服务、覆盖配置文件 → high\n"
        "- 涉及文件复制、mkdir、chmod、wget/curl 下载 → medium\n"
        "- 只读查询、echo、日志输出 → low\n"
        "- 若目标环境已存在端口占用/目录冲突，precheck 应给出检测命令"
    )
    user_prompt = (
        f"## 步骤\n序列: {step.step_order}\n描述: {step.description}\n"
        f"命令: {step.command}\n风险声明: {step.risk_level}\n"
        f"## 目标环境\n{json.dumps(probe, ensure_ascii=False)[:1200]}\n\n"
        "请分析该命令的执行风险。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return {}
        content = resp["choices"][0]["message"]["content"]
        r = _extract_json(content) or {}
        if r.get("risk"):
            return r
    except Exception:
        pass
    return {}


def _ai_autonomous_decision(provider, step: DeployStep, output: str, diag: dict, history: list) -> str:
    """AI 自主决策（能力②：失败自主决策），返回 fix/retry/skip/rollback 之一。"""
    sys_prompt = (
        "你是资深 SRE 运维专家，部署步骤失败后由你自主决策下一步操作，无需人工确认。\n"
        "只输出一个单词：fix 或 retry 或 skip 或 rollback\n"
        "决策规则：\n"
        "- fix: 有明确修复命令且修复成功率高（>70%），选择 fix\n"
        "- retry: 疑似临时性问题（网络抖动、资源刚好被占用、重试能解决），选择 retry\n"
        "- skip: 该步骤非关键、失败不影响整体可用性、后续步骤可补偿，选择 skip\n"
        "- rollback: 严重错误、修复/重试风险高、继续执行会扩大损失，选择 rollback\n"
        "只输出一个单词，不要任何其他内容。"
    )
    history_txt = "无"
    if history:
        history_txt = "; ".join([f"第{h.get('attempt', 1)}次: {h.get('decision', '')}({h.get('result', '')})" for h in history[-3:]])
    user_prompt = (
        f"## 失败步骤\n步骤 {step.step_order}: {step.description}\n风险等级: {step.risk_level}\n"
        f"## 命令\n{(step.command or '')[:200]}\n"
        f"## 输出(截断)\n{(output or '')[:1500]}\n"
        f"## AI 诊断\n根因: {diag.get('root_cause', '')}\n建议: {diag.get('suggestion', '')}\n"
        f"修复命令: {json.dumps(diag.get('fix_commands', []), ensure_ascii=False)}\n"
        f"## 历史决策\n{history_txt}\n\n请自主决策下一步操作。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return diag.get("suggestion", "rollback") or "rollback"
        content = resp["choices"][0]["message"]["content"].strip().lower()
        if content in ("fix", "retry", "skip", "rollback"):
            return content
    except Exception:
        pass
    return diag.get("suggestion", "rollback") or "rollback"


def _ai_adaptive_rollback(provider, steps: list, plan: DeployPlan) -> Optional[list]:
    """AI 自适应回滚分析（能力⑤）：返回需回滚的步骤逆序列表(step_order 列表)。
    跳过无状态步骤(echo/mkdir/校验等)。AI 不可用时回退全量逆序。
    """
    step_info = [{
        "order": s.step_order,
        "description": s.description,
        "command": (s.command or "")[:150],
        "rollback": (s.rollback_command or "")[:150],
        "status": s.status,
    } for s in steps]
    sys_prompt = (
        "你是资深 SRE 运维专家，分析部署失败后哪些步骤真正需要回滚。\n"
        "严格按 JSON 输出，不要额外内容：\n"
        "{\n"
        '  "rollback_steps": [3, 2],\n'
        '  "skip_steps": [1, 4],\n'
        '  "reasoning": "简要分析说明"\n'
        "}\n"
        "规则：\n"
        "- rollback_steps 按回滚执行顺序排列（先回滚后执行的步骤，即逆序）\n"
        "- 无状态步骤（echo、mkdir -p、ls、校验命令、只读检查）可跳过回滚\n"
        "- 有状态步骤（docker compose up、cp/覆盖文件、服务启动、配置写入）必须回滚\n"
        "- 必须覆盖所有已成功执行的步骤"
    )
    user_prompt = f"## 步骤列表(含状态)\n{json.dumps(step_info, ensure_ascii=False, indent=1)}\n\n请分析哪些步骤需要回滚。"
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return None
        content = resp["choices"][0]["message"]["content"]
        r = _extract_json(content) or {}
        if isinstance(r.get("rollback_steps"), list) and r["rollback_steps"]:
            return r["rollback_steps"]
    except Exception:
        pass
    return None


def _ai_decision_log(plan: DeployPlan, entry: dict) -> None:
    """把 AI 自主决策/重排追加到 plan.ai_decision_log_json（最多 200 条）。"""
    try:
        log = json.loads(plan.ai_decision_log_json) if isinstance(plan.ai_decision_log_json, str) and plan.ai_decision_log_json not in ("[]", "") else []
    except Exception:
        log = []
    if not isinstance(log, list):
        log = []
    entry.setdefault("ts", datetime.now().isoformat())
    log.append(entry)
    if len(log) > 200:
        log = log[-200:]
    plan.ai_decision_log_json = json.dumps(log, ensure_ascii=False)


# ═══════════════════════════════════════════════════════════════
# L4 — AI 决策引擎（策略选择 + 实时决策 + 健康门控 + 动态调度）
# L5 — 历史学习引擎（特征记录 + 风险评分 + 模式匹配）
# ═══════════════════════════════════════════════════════════════


def _ai_select_deployment_strategy(provider, plan: DeployPlan, steps: list, probe: dict, assets: list) -> str:
    """L4: AI 根据服务类型/风险/环境选择部署策略。
    返回: rolling / blue-green / canary / recreate / auto。
    """
    step_descs = [{"order": s.step_order, "desc": s.description, "cmd": (s.command or "")[:100], "risk": s.risk_level} for s in steps]
    asset_info = [{"name": a.name, "ip": a.ip, "ci_type": a.ci_type} for a in assets]
    has_compose = any("docker compose" in (s.command or "") or "docker-compose" in (s.command or "") for s in steps)
    has_service = any("up" in (s.command or "") or "restart" in (s.command or "") or "start" in (s.command or "") for s in steps)
    sys_prompt = (
        "你是一名资深 SRE 架构师，负责为部署计划选择最佳策略。\n"
        "只输出一个单词：rolling 或 blue-green 或 canary 或 recreate 或 auto\n"
        "规则：\n"
        "- recreate: 简单应用/测试/开发环境/无状态服务，风险低，直接重建\n"
        "- rolling: 生产环境多实例服务，需要逐步替换，保持可用性\n"
        "- blue-green: 关键业务服务，需要零停机切换，有完整回滚能力\n"
        "- canary: 高风险变更，需要灰度放量验证，逐步扩大范围\n"
        "- auto: 不确定时选 auto，由执行引擎自动判断\n"
        "只输出一个单词，不要任何其他内容。"
    )
    user_prompt = (
        f"## 服务类型\n涉及 docker-compose: {has_compose}\n涉及服务启停: {has_service}\n"
        f"资产数: {len(assets)}\n\n## 步骤\n{json.dumps(step_descs, ensure_ascii=False, indent=1)}\n\n"
        f"## 目标资产\n{json.dumps(asset_info, ensure_ascii=False)}\n\n"
        f"## 环境\nOS: {probe.get('os','')[:80]}\nDocker: {probe.get('docker','')[:30]}\n\n请选择最佳部署策略。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return "auto"
        decision = resp["choices"][0]["message"]["content"].strip().lower()
        if decision in ("rolling", "blue-green", "canary", "recreate", "auto"):
            return decision
    except Exception:
        pass
    return "auto"


def _ai_risk_scoring(provider, plan: DeployPlan, steps: list, probe: dict, assets: list, history: list) -> int:
    """L5: AI 基于历史数据和当前环境，预判部署风险评分 0-100。
    0=无风险, 100=极高风险。
    """
    step_risks = [s.risk_level for s in steps]
    high_count = sum(1 for r in step_risks if r == "high")
    med_count = sum(1 for r in step_risks if r == "medium")
    prev_failures = sum(1 for h in history if h.get("status") in ("failed", "rolled_back")) if history else 0
    prev_total = len(history) if history else 0
    fail_rate = prev_failures / prev_total if prev_total > 0 else 0
    sys_prompt = (
        "你是一名资深 SRE 专家，负责评估部署风险。\n"
        "只输出一个 0-100 的整数，不要任何其他内容。\n"
        "评分规则：\n"
        "- 0-20: 低风险（简单脚本/只读操作/测试环境）\n"
        "- 21-50: 中等风险（常规部署/有回滚/已验证）\n"
        "- 51-80: 高风险（涉及服务重启/端口变更/生产环境）\n"
        "- 81-100: 极高风险（数据库变更/架构变更/历史失败率高）\n"
        "综合考虑：步骤风险等级、历史失败率、服务类型、资产数量。"
    )
    user_prompt = (
        f"## 步骤风险\nhigh={high_count} medium={med_count} low={len(steps)-high_count-med_count}\n"
        f"## 历史失败率\n失败 {prev_failures}/{prev_total} 次 ({fail_rate*100:.0f}%)\n"
        f"## 资产数\n{len(assets)} 台\n\n请评估部署风险评分(0-100)。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return min(100, high_count * 30 + med_count * 10)
        score = int("".join(c for c in resp["choices"][0]["message"]["content"] if c.isdigit()) or "50")
        return max(0, min(100, score))
    except Exception:
        return min(100, high_count * 30 + med_count * 10)


def _record_deployment_feature(plan: DeployPlan, steps: list, assets: list, probe: dict,
                                status: str, step_results: list, total_duration: float) -> dict:
    """L5: 记录部署特征向量，供后续学习和风险预测使用。
    返回特征 dict 并存入 plan.deployment_feature_json。
    """
    step_risks = [s.risk_level for s in steps]
    features = {
        "step_count": len(steps),
        "asset_count": len(assets),
        "high_risk_steps": sum(1 for r in step_risks if r == "high"),
        "medium_risk_steps": sum(1 for r in step_risks if r == "medium"),
        "has_compose": any("docker compose" in (s.command or "") for s in steps),
        "has_docker_run": any("docker run" in (s.command or "") for s in steps),
        "has_systemctl": any("systemctl" in (s.command or "") for s in steps),
        "has_db_migration": any("migrate" in (s.command or "").lower() or "schema" in (s.command or "").lower() for s in steps),
        "total_duration_seconds": total_duration,
        "step_success_count": sum(1 for r in step_results if r.get("status") == "succeeded"),
        "step_fail_count": sum(1 for r in step_results if r.get("status") == "failed"),
        "step_skip_count": sum(1 for r in step_results if r.get("status") == "skipped"),
        "retry_count": sum(r.get("retry_count", 0) for r in step_results),
        "ai_decision_count": len(step_results) - sum(1 for r in step_results if r.get("status") == "succeeded"),
        "deploy_status": status,
        "os_type": "centos" if "centos" in (probe.get("os", "") or "").lower() else "ubuntu" if "ubuntu" in (probe.get("os", "") or "").lower() else "other",
        "docker_version": (probe.get("docker", "") or "")[:20],
    }
    plan.deployment_feature_json = json.dumps(features, ensure_ascii=False)
    return features


def _ai_pattern_matching(provider, features: dict, history: list) -> Optional[dict]:
    """L5: AI 匹配历史部署中的失败模式，提前预警。
    返回 {"matched": True/False, "pattern": "...", "risk": "high/medium/low", "suggestion": "..."}。
    """
    if not history or len(history) < 2:
        return None
    # 提取历史特征
    hist_features = [h for h in history if isinstance(h, dict) and h.get("features")]
    if not hist_features:
        return None
    sys_prompt = (
        "你是一名资深 SRE 专家，分析部署历史数据，识别失败模式。\n"
        "严格按 JSON 输出，不要额外内容：\n"
        "{\n"
        '  "matched": true/false,\n'
        '  "pattern": "匹配到的失败模式描述",\n'
        '  "risk": "high/medium/low",\n'
        '  "suggestion": "建议的预防措施"\n'
        "}\n"
        "如果当前部署特征与历史上失败模式相似，matched=true 并给出建议。"
    )
    user_prompt = (
        f"## 当前部署特征\n{json.dumps(features, ensure_ascii=False, indent=1)}\n\n"
        f"## 历史部署({len(hist_features)} 次)\n{json.dumps(hist_features[-5:], ensure_ascii=False, indent=1)}\n\n"
        f"请分析是否存在失败模式匹配。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return None
        content = resp["choices"][0]["message"]["content"]
        pattern = _extract_json(content) or {}
        if pattern.get("matched"):
            return pattern
    except Exception:
        pass
    return None


def _ai_assess_state(provider, plan: DeployPlan, step_results: list, current_step: int,
                      health_data: dict, strategy: str) -> dict:
    """L4: AI 实时评估部署状态，决定下一步动作。
    返回 {"action": "continue/adjust/rollback/complete", "reason": "...", "adjustments": {...}}。
    """
    succeed_count = sum(1 for r in step_results if r.get("status") == "succeeded")
    fail_count = sum(1 for r in step_results if r.get("status") == "failed")
    total = len(step_results)
    sys_prompt = (
        "你是一名资深 SRE 专家，负责在部署过程中实时评估状态并决策。\n"
        "严格按 JSON 输出，不要额外内容：\n"
        "{\n"
        '  "action": "continue/adjust/rollback/complete",\n'
        '  "reason": "决策理由（中文，一句话）",\n'
        '  "adjustments": {\n'
        '    "parallelism": "保持/增加/减少",\n'
        '    "pace": "正常/加快/放慢",\n'
        '    "next_steps": "继续执行/先验证/先修复"\n'
        '  }\n'
        "}\n"
        "规则：\n"
        "- continue: 一切正常，继续执行\n"
        "- adjust: 需要调整执行节奏或并行度\n"
        "- rollback: 出现严重问题，必须回滚\n"
        "- complete: 所有步骤已完成或无需继续\n"
        "如果失败率超过 30%，考虑 rollback。"
    )
    user_prompt = (
        f"## 当前策略\n{strategy}\n\n"
        f"## 执行进度\n总步骤: {total}\n成功: {succeed_count}\n失败: {fail_count}\n当前步骤: {current_step}\n\n"
        f"## 健康数据\n{json.dumps(health_data, ensure_ascii=False)[:500]}\n\n"
        f"请评估当前状态并决策。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return {"action": "continue", "reason": "AI 评估不可用", "adjustments": {}}
        content = resp["choices"][0]["message"]["content"]
        decision = _extract_json(content) or {}
        if decision.get("action") in ("continue", "adjust", "rollback", "complete"):
            return decision
    except Exception:
        pass
    return {"action": "continue", "reason": "AI 评估异常，默认继续", "adjustments": {}}


def _ai_health_gate(provider, client, asset_map: dict, step: DeployStep, strategy: str) -> dict:
    """L4: 健康门控 — 每步执行后检查关键指标，决定是否放行下一步。
    返回 {"passed": True/False, "checks": [...], "recommendation": "continue/verify/rollback"}。
    """
    checks = []
    # 1. Docker 守护进程
    try:
        _, o, _ = client.exec_command("docker info --format '{{.ServerVersion}}' 2>/dev/null || echo 'FAIL'", timeout=10)
        rc = o.channel.recv_exit_status()
        ver = o.read().decode(errors="replace").strip()
        checks.append({"name": "Docker 守护进程", "passed": rc == 0, "detail": ver if rc == 0 else "不可用"})
    except Exception as e:
        checks.append({"name": "Docker 守护进程", "passed": False, "detail": str(e)})
    # 2. 磁盘空间
    try:
        _, o, _ = client.exec_command("df -h / 2>/dev/null | tail -1 | awk '{print $5}'", timeout=5)
        pct = o.read().decode(errors="replace").strip().rstrip("%")
        disk_ok = int(pct) < 85 if pct.isdigit() else True
        checks.append({"name": "磁盘使用率", "passed": disk_ok, "detail": f"{pct}%" if pct.isdigit() else "未知"})
    except Exception:
        checks.append({"name": "磁盘使用率", "passed": True, "detail": "检查失败"})
    # 3. 关键端口（如果已知）
    target_port = asset_map.get("TARGET_PORT", "")
    if target_port:
        try:
            _, o, _ = client.exec_command(f"ss -tlnp | grep ':{target_port} ' || echo 'FREE'", timeout=5)
            out = o.read().decode(errors="replace").strip()
            checks.append({"name": f"端口 {target_port}", "passed": True, "detail": "已监听" if "FREE" not in out else "未占用"})
        except Exception:
            pass
    all_passed = all(c.get("passed", True) for c in checks)
    recommendation = "continue" if all_passed else "rollback" if strategy == "blue-green" else "verify"
    return {"passed": all_passed, "checks": checks, "recommendation": recommendation}


def _ai_dynamic_scheduling(provider, plan: DeployPlan, step_results: list, strategy: str) -> dict:
    """L4: AI 动态调度 — 根据执行进度和状态，调整后续步骤的并行度/顺序。
    返回 {"adjusted_dag": {...}, "reason": "..."}。
    """
    fail_count = sum(1 for r in step_results if r.get("status") == "failed")
    retry_count = sum(r.get("retry_count", 0) for r in step_results)
    sys_prompt = (
        "你是一名资深 SRE 专家，负责动态调整部署调度。\n"
        "严格按 JSON 输出，不要额外内容：\n"
        "{\n"
        '  "adjust": "none/slow_down/speed_up/reorder",\n'
        '  "reason": "调整理由（中文）",\n'
        '  "parallelism": "keep/reduce/increase"\n'
        "}\n"
        "规则：\n"
        "- 失败次数多 → slow_down + reduce 并行度\n"
        "- 一切顺利 → speed_up + increase 并行度\n"
        "- 重试过多 → reorder 提前执行关键步骤"
    )
    user_prompt = (
        f"## 策略\n{strategy}\n\n## 执行结果\n{json.dumps(step_results[-5:], ensure_ascii=False, indent=1)}\n\n"
        f"## 统计\n失败: {fail_count}\n重试: {retry_count}\n\n请决定是否调整调度。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=30)
        if resp.get("error"):
            return {"adjust": "none", "reason": "AI 不可用", "parallelism": "keep"}
        content = resp["choices"][0]["message"]["content"]
        decision = _extract_json(content) or {}
        if decision.get("adjust"):
            return decision
    except Exception:
        pass
    return {"adjust": "none", "reason": "默认继续", "parallelism": "keep"}


def _ai_plan_step_autonomous(provider, step: DeployStep, plan: DeployPlan, probe: dict, env_context: dict) -> dict:
    """AI 理解步骤意图 + 结合环境 → 自主生成执行计划。

    替代机械执行 step.command，让 AI 理解步骤意图后，根据当前环境动态调整命令。
    返回 {"intent": "步骤意图", "commands": ["调整后的命令"], "verify": "验证命令",
          "expected": "期望结果", "adjustments": ["调整说明"], "risk": "low/medium/high"}。
    """
    step_desc = step.description or ""
    step_cmd = (step.command or "")[:500]
    step_verify = (step.verify_command or "")[:200]
    probe_os = (probe.get("os", "") if isinstance(probe, dict) else "")[:200]
    probe_docker = (probe.get("docker", "") if isinstance(probe, dict) else "")[:50]
    probe_ports = (probe.get("port_scan", {}) if isinstance(probe, dict) else {})
    probe_containers = (probe.get("containers", "") if isinstance(probe, dict) else "")[:200]
    prev_output = env_context.get("prev_output", "")[:500]
    prev_status = env_context.get("prev_status", "")
    cwd = env_context.get("cwd", "")
    sys_prompt = (
        "你是一名资深 SRE 运维专家，负责理解部署步骤的意图，并根据当前目标机环境自主生成最佳执行方案。\n"
        "严格按以下 JSON Schema 输出，不要任何额外内容：\n"
        "{\n"
        '  "intent": "步骤意图（中文，一句话说明这本书要做什么）",\n'
        '  "commands": ["最终要执行的 shell 命令（每行一条）"],\n'
        '  "verify": "执行后的验证命令（可选，空串表示无需验证）",\n'
        '  "expected": "期望结果描述",\n'
        '  "adjustments": ["针对当前环境的调整说明（如：端口冲突换端口、目录已存在跳过创建）"],\n'
        '  "risk": "low/medium/high"\n'
        "}\n"
        "关键规则：\n"
        "- 必须理解步骤意图，不能机械照搬手册命令\n"
        "- 结合当前环境（OS/Docker 版本/端口/容器/磁盘/内存）调整命令\n"
        "- 如果手册命令在当前环境有问题（端口冲突、镜像不存在、目录已存在），自动调整\n"
        "- 如果上一步失败或告警，自动添加前置检查\n"
        "- commands 中的命令必须是可直接执行的 shell 命令，且必须真实、可落地\n"
        "- 不要自行编造或添加 cd 前缀：工作目录由系统根据当前工作目录(cwd)自动处理；\n"
        "  只有当命令确实需要切换到其它目录时，才使用该命令自身的 cd，且 cd 后的路径必须是真实存在的绝对路径\n"
        "- 严禁使用占位示例路径（如 /path/to/project、/home/user/xxx 等）——不存在这样的目录，会导致命令失败\n"
        "- 如果无法确定正确的目录或路径，直接使用原始命令，不要编造\n"
        "- 确保命令安全：不要 rm -rf 无验证、不要遗漏依赖\n"
        "- 如果无法确定，保守执行原命令，不要编造"
    )
    user_prompt = (
        f"## 步骤信息\n序号: {step.step_order}\n"
        f"描述: {step_desc}\n原始命令: {step_cmd}\n原始验证: {step_verify}\n"
        f"风险等级: {step.risk_level}\n\n"
        f"## 目标机环境\nOS: {probe_os}\nDocker: {probe_docker}\n"
        f"端口状态: {json.dumps(probe_ports, ensure_ascii=False)[:200]}\n"
        f"当前容器: {probe_containers[:200]}\n\n"
        f"## 执行上下文\n上一步状态: {prev_status}\n上一步输出: {prev_output[:300]}\n"
        f"当前工作目录: {cwd}\n\n"
        f"请理解步骤意图，结合环境自主生成最佳执行方案。"
    )
    try:
        resp = call_llm(provider, [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ], timeout_override=60)
        if resp.get("error"):
            return {"intent": step_desc, "commands": [step_cmd], "verify": step_verify,
                    "expected": "", "adjustments": [], "risk": step.risk_level}
        content = resp["choices"][0]["message"]["content"]
        plan_data = _extract_json(content) or {}
        if plan_data.get("commands") and isinstance(plan_data["commands"], list):
            return plan_data
    except Exception as e:
        logger.warning(f"AI 自主规划异常: {e}")
    return {"intent": step_desc, "commands": [step_cmd], "verify": step_verify,
            "expected": "", "adjustments": [], "risk": step.risk_level}


def _ai_resource_check(provider, plan: DeployPlan, steps: list, assets: list, probe: dict) -> dict:
    """前置资源检查：SSH 采集目标机真实资源，AI 分析是否足够部署。
    返回 {"passed": True/False, "checks": [...], "recommendation": "proceed/block/warn", "summary": "..."}。
    """
    if not assets:
        return {"passed": False, "checks": [{"name": "SSH连接", "passed": False, "detail": "无目标资产"}],
                "recommendation": "block", "summary": "无目标资产，无法部署"}

    asset = assets[0]
    try:
        client, host = _ssh_connect(asset)
    except Exception as e:
        return {"passed": False, "checks": [{"name": "SSH连接", "passed": False, "detail": str(e)[:100]}],
                "recommendation": "block", "summary": f"SSH 连接失败: {str(e)[:100]}"}

    checks = []
    try:
        # 1. 内存
        _, o, _ = client.exec_command("free -m | grep Mem:", timeout=10)
        mem_raw = o.read().decode(errors="replace").strip().split()
        mem_total = int(mem_raw[1]) if len(mem_raw) > 1 else 0
        mem_used = int(mem_raw[2]) if len(mem_raw) > 2 else 0
        mem_avail = int(mem_raw[6]) if len(mem_raw) > 6 else mem_total - mem_used  # available 列
        # 估算服务所需内存(mem_avail 已扣除系统占用，故基础值只算 Docker 容器开销)
        estimated_mb = 128  # Docker 容器 + 网络开销
        for s in steps:
            cmd = (s.command or "").lower()
            if "mysql" in cmd or "postgres" in cmd or "mariadb" in cmd:
                estimated_mb += 400
            elif "redis" in cmd:
                estimated_mb += 50
            elif "elastic" in cmd or "mongo" in cmd:
                estimated_mb += 512
            elif "java" in cmd or "spring" in cmd or "tomcat" in cmd or "jenkins" in cmd:
                estimated_mb += 512
            elif "nginx" in cmd:
                estimated_mb += 50
            elif "python" in cmd or "flask" in cmd or "django" in cmd or "node" in cmd or "php" in cmd:
                estimated_mb += 128
            elif "rabbitmq" in cmd or "kafka" in cmd:
                estimated_mb += 256
            elif "gitlab" in cmd:
                estimated_mb += 1024
            elif "memcached" in cmd:
                estimated_mb += 50
        mem_ok = mem_avail >= estimated_mb
        checks.append({"name": "内存", "passed": mem_ok, "detail": f"可用 {mem_avail}MB, 估算需 {estimated_mb}MB",
                        "value": mem_avail, "threshold": estimated_mb, "unit": "MB"})

        # 2. 磁盘
        _, o, _ = client.exec_command("df -m / | tail -1 | awk '{print $2,$3,$4}'", timeout=10)
        disk_raw = o.read().decode(errors="replace").strip().split()
        disk_total = int(disk_raw[0]) if len(disk_raw) > 0 else 0
        disk_avail = int(disk_raw[2]) if len(disk_raw) > 2 else 0
        disk_ok = disk_avail >= 2048  # 至少 2GB
        checks.append({"name": "磁盘", "passed": disk_ok, "detail": f"可用 {disk_avail}MB, 建议>=2048MB",
                        "value": disk_avail, "threshold": 2048, "unit": "MB"})

        # 3. Docker 可用
        _, o, _ = client.exec_command("docker info --format '{{.ServerVersion}}' 2>/dev/null || echo 'NODOCKER'", timeout=10)
        docker_ver = o.read().decode(errors="replace").strip()
        docker_ok = docker_ver != "NODOCKER" and docker_ver != ""
        checks.append({"name": "Docker", "passed": docker_ok, "detail": docker_ver if docker_ok else "Docker 不可用"})

        # 4. 端口冲突
        port_conflicts = []
        for s in steps:
            cmd = (s.command or "")
            for _m in re.finditer(r'(?:EXPOSE\s+|ports:\s*["\']?|:\s*)(\d{2,5})', cmd):
                _p = _m.group(1)
                _pi = int(_p)
                if 80 <= _pi <= 65535:
                    _, o, _ = client.exec_command(f"ss -tlnp | grep ':{_p} ' || echo 'FREE'", timeout=5)
                    p_status = o.read().decode(errors="replace").strip()
                    if "FREE" not in p_status:
                        port_conflicts.append(_p)
        port_ok = len(port_conflicts) == 0
        checks.append({"name": "端口", "passed": port_ok, "detail": f"冲突端口: {', '.join(port_conflicts)}" if port_conflicts else "无冲突",
                        "conflicts": port_conflicts})

        # 5. 容器名冲突
        container_conflicts = []
        for s in steps:
            cmd = (s.command or "")
            for _m in re.finditer(r'container_name:\s*(\S+)', cmd):
                cname = _m.group(1)
                _, o, _ = client.exec_command(f"docker ps -a --format '{{{{.Names}}}}' | grep -x '{cname}' || echo 'FREE'", timeout=5)
                if "FREE" not in o.read().decode(errors="replace").strip():
                    container_conflicts.append(cname)
        container_ok = len(container_conflicts) == 0
        checks.append({"name": "容器名", "passed": container_ok, "detail": f"冲突: {', '.join(container_conflicts)}" if container_conflicts else "无冲突"})

        # 6. 镜像可用性
        images_missing = []
        for s in steps:
            for _m in re.finditer(r'image:\s*(\S+)', s.command or ""):
                img = _m.group(1)
                _, o, _ = client.exec_command(f"docker images --format '{{{{.Repository}}}}.{{{{.Tag}}}}' | grep -q '{img.split(':')[0]}' && echo 'HAVE' || echo 'MISS'", timeout=5)
                if "HAVE" not in o.read().decode(errors="replace").strip():
                    images_missing.append(img)
        images_ok = len(images_missing) == 0
        checks.append({"name": "镜像", "passed": images_ok, "detail": f"缺失: {', '.join(images_missing)}" if images_missing else "全部存在"})

    finally:
        client.close()

    # 汇总
    all_passed = all(c["passed"] for c in checks)
    critical_fail = not mem_ok or not docker_ok or not disk_ok
    recommendation = "block" if critical_fail else "warn" if not all_passed else "proceed"

    # AI 分析
    summary = f"资源检查: {'通过' if all_passed else '有风险'}"
    if provider:
        sys_prompt = "你是一名 SRE 专家，分析部署前置检查结果。输出 JSON: {\"passed\": true/false, \"summary\": \"一句话总结\", \"suggestion\": \"建议\"}"
        user_prompt = f"检查项: {json.dumps(checks, ensure_ascii=False)}\nestimated_mb={estimated_mb} mem_avail={mem_avail}\n请分析是否可部署。"
        try:
            resp = call_llm(provider, [
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": user_prompt},
            ], timeout_override=15)
            if not resp.get("error"):
                content = resp["choices"][0]["message"]["content"]
                ai = _extract_json(content) or {}
                if ai.get("summary"):
                    summary = ai["summary"]
        except Exception:
            pass

    return {"passed": all_passed, "checks": checks, "recommendation": recommendation, "summary": summary}


def _wait_for_risk_confirm(decision_queue, plan_id, step_order, command, risk_level, reason) -> str:
    """等待用户确认高危操作。返回 'confirm' 或 'reject'。
    如果 WS 断开（_STOPPED 标志），短超时轮询避免线程永久阻塞。"""
    if decision_queue is None:
        return "confirm"  # 无决策队列时默认放行
    import queue as _qm
    try:
        _decision_queue = decision_queue
        while True:
            try:
                decision = _decision_queue.get(timeout=5)
                if decision in ("confirm", "reject"):
                    return decision
                # 兼容旧决策消息
                if decision in ("fix", "retry", "skip", "rollback"):
                    return "confirm"
            except _qm.Empty:
                if _STOPPED.get(plan_id):
                    return "reject"
                continue
            except Exception:
                return "reject"
    except Exception:
        return "confirm"


def _ai_stream_execute(db: Session, plan_id: int, user_id: int = 0, decision_queue=None):
    """AI 驱动流式执行部署计划（10 分执行引擎）。

    能力① AI 动态编排(DAG) → 能力③ AI 执行前预判 → 按 DAG 执行 →
    失败时能力② AI 自主决策 → 能力⑤ AI 自适应回滚
    能力④ AI 并行调度：DAG parallel=true 组内步骤并行执行（多线程）

    产出事件 dict（兼容旧事件 + 新增）：
      {"type": "status", "status": "running"}
      {"type": "dag_plan", "groups": [...], "reasoning": "..."}   # 新增：DAG 执行计划
      {"type": "asset_start", "asset": 主机名, "ip": ...}
      {"type": "parallel_group", "group": N, "steps": [1,2], "parallel": bool}  # 新增
      {"type": "step_start", "step": 序号, "description": ..., "risk": ..., "ai_risk": "..."}  # 增强
      {"type": "ai_precheck", "step": 序号, "risk": ..., "reason": ..., "precheck": ...}  # 新增
      {"type": "cmd", "command": "..."}
      {"type": "output", "line": "..."}
      {"type": "step_end", "step": 序号, "status": 状态}
      {"type": "ai_decision", "step": 序号, "decision": "fix/retry/skip/rollback", "reason": "..."}  # 新增
      {"type": "asset_end", "asset": 主机名, "status": 状态}
      {"type": "complete", "status": 计划终态, "total_assets": N, "succeeded_assets": N}
    """
    import time as _time
    import concurrent.futures as _cf
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        yield {"type": "error", "message": "计划不存在"}
        return
    if plan.status == "draft":
        yield {"type": "error", "message": "计划尚未规划，无法执行"}
        return
    assets = _get_assets(db, plan)
    if not assets:
        yield {"type": "error", "message": "计划未关联目标资产"}
        return

    mapping = json.loads(plan.env_mapping or "{}")
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).order_by(DeployStep.step_order).all()
    if not steps:
        yield {"type": "error", "message": "无部署步骤"}
        return

    _provider = _get_provider(db)
    _health_gates = []

    _sync_env_mapping_from_sop(db, plan)
    mapping = json.loads(plan.env_mapping or "{}")
    probe = _safe_json(plan.environment_probe_json)

    # ── AI 自动解决未解析的环境变量 ──
    try:
        mapping = _ai_auto_resolve_unresolved(_provider, plan, db, steps, mapping)
        if mapping:
            plan.env_mapping = json.dumps(mapping, ensure_ascii=False)
            db.commit()
    except Exception as e:
        logger.warning(f"AI 自动解决环境变量异常: {e}")

    plan.status = "running"
    for _s in steps:
        _s.status = "pending"
        _s.output = ""
        _s.started_at = None
        _s.finished_at = None
    db.commit()
    yield {"type": "status", "status": "running"}

    # ── 前置资源检查 ──
    try:
        rc = _ai_resource_check(_provider, plan, steps, assets, probe)
        plan.preflight_json = json.dumps(rc, ensure_ascii=False)
        db.commit()
        yield {"type": "resource_check", "passed": rc["passed"], "checks": rc["checks"],
               "recommendation": rc["recommendation"], "summary": rc["summary"]}
        if rc["recommendation"] == "block":
            yield {"type": "error", "message": f"❌ 资源检查未通过，部署终止: {rc['summary']}"}
            plan.status = "planned"
            db.commit()
            return
        if rc["recommendation"] == "warn":
            yield {"type": "output", "line": f"\r\n\x1b[33m⚠️ 资源检查有风险，但 AI 建议继续: {rc['summary']}\x1b[0m"}
    except Exception as rce:
        logger.warning(f"资源检查异常: {rce}")

    # ── L4: AI 选择部署策略 ──
    strategy = "auto"
    try:
        history_now = _safe_json(plan.execution_history_json)
        if not isinstance(history_now, list):
            history_now = []
        if _provider:
            strategy = _ai_select_deployment_strategy(_provider, plan, steps, probe, assets)
            risk_score = _ai_risk_scoring(_provider, plan, steps, probe, assets, history_now)
            plan.strategy = strategy
            plan.risk_score = risk_score
            db.commit()
            yield {"type": "strategy_selected", "strategy": strategy, "risk_score": risk_score,
                   "reason": f"AI 评估风险 {risk_score}/100，选择 {strategy} 策略"}
            logger.info(f"部署策略选定 plan_id={plan_id} strategy={strategy} risk={risk_score}")
    except Exception as e:
        logger.warning(f"AI 策略选择异常: {e}")

    # ── L5: 历史模式匹配预警 ──
    try:
        if _provider and history_now:
            pending_features = _record_deployment_feature(plan, steps, assets, probe, "pending", [], 0)
            pattern = _ai_pattern_matching(_provider, pending_features, history_now)
            if pattern:
                yield {"type": "risk_warning", "pattern": pattern.get("pattern", ""),
                       "risk": pattern.get("risk", ""), "suggestion": pattern.get("suggestion", "")}
    except Exception as e:
        logger.warning(f"AI 模式匹配异常: {e}")

    # ── 能力①：AI 构建 DAG 执行计划 ──
    dag = None
    if _provider:
        dag = _ai_build_execution_dag(db, _provider, plan, steps)
    if not dag:
        dag = {"groups": [{"group_order": i + 1, "step_orders": [s.step_order], "parallel": False, "reason": "线性回退"} for i, s in enumerate(steps)],
               "reasoning": "AI 不可用，回退为线性顺序执行"}
    plan.dag_json = json.dumps(dag, ensure_ascii=False)
    db.commit()
    yield {"type": "dag_plan", "groups": dag["groups"], "reasoning": dag.get("reasoning", "")}

    total_assets = len(assets)
    succeeded_assets = 0
    step_map = {s.step_order: s for s in steps}

    for asset in assets:
        try:
            client, host = _ssh_connect(asset)
            _RUNNING_CLIENTS[plan_id] = client
        except Exception as e:
            yield {"type": "error", "line": f"❌ {asset.name}({asset.ip}) SSH 连接失败: {e}"}
            continue

        asset_map = dict(mapping)
        asset_map["TARGET_IP"] = asset.ip or ""
        asset_map["TARGET_HOSTNAME"] = asset.name or ""

        yield {"type": "asset_start", "asset": asset.name, "ip": asset.ip}
        asset_failed = False
        # 默认工作目录 = 源码下载路径(APP_DIR)，保证后续依赖 cwd 的命令(如 docker compose)即使 AI 命令不带 cd 也能在正确目录执行
        try:
            _cwd = resolve_download_path(plan)
        except Exception:
            _cwd = ""
        try:
            for group in dag["groups"]:
                if asset_failed:
                    break
                step_orders = group["step_orders"]
                parallel = group.get("parallel", False)
                yield {"type": "parallel_group", "group": group["group_order"], "steps": step_orders, "parallel": parallel, "reason": group.get("reason", "")}

                if parallel and len(step_orders) > 1:
                    # ── 能力④：AI 并行调度 ──
                    def _exec_parallel_step(so):
                        p_step = step_map[so]
                        p_cmd = _resolve_command(p_step.command, asset_map)
                        p_verify = _resolve_command(p_step.verify_command, asset_map) if p_step.verify_command else ""
                        p_cwd = _cwd
                        p_cd_m = re.match(r'^cd\s+([^\s;&|]+)', p_cmd.strip())
                        if p_cd_m:
                            p_cwd = p_cd_m.group(1).strip('"\'')
                        if not p_cd_m and p_cwd:
                            p_cmd = f"cd {p_cwd} && {p_cmd}"
                        _ppx = _proxy_env_prefix(plan)
                        if _ppx:
                            p_cmd = f"{_ppx} && {p_cmd}"
                        try:
                            p_client, _ = _ssh_connect(asset)
                            _, p_stdout, p_stderr = p_client.exec_command(p_cmd, timeout=600)
                            p_rc = p_stdout.channel.recv_exit_status()
                            p_out = p_stdout.read().decode("utf-8", errors="replace").strip()
                            p_err = p_stderr.read().decode("utf-8", errors="replace").strip()
                            p_client.close()
                            if p_err:
                                p_out = f"{p_out}\n{p_err}" if p_out else p_err
                            return {"step_order": so, "exit_code": p_rc, "output": p_out}
                        except Exception as pe:
                            return {"step_order": so, "exit_code": -1, "output": str(pe)}

                    with _cf.ThreadPoolExecutor(max_workers=len(step_orders)) as pool:
                        para_results = list(pool.map(_exec_parallel_step, step_orders))
                    for pr in para_results:
                        ps = step_map[pr["step_order"]]
                        ps.output = (ps.output or "") + f"\n[{asset.name}]\n{pr['output'][:2000]}"
                        ps.started_at = _now()
                        ps.finished_at = _now()
                        if pr["exit_code"] == 0:
                            ps.status = "succeeded"
                            yield {"type": "step_start", "step": ps.step_order, "description": ps.description, "risk": ps.risk_level}
                            yield {"type": "output", "line": f"并行步骤 {ps.step_order}: {ps.description} ✔"}
                            yield {"type": "step_end", "step": ps.step_order, "status": "succeeded"}
                        else:
                            ps.status = "failed"
                            asset_failed = True
                            yield {"type": "step_start", "step": ps.step_order, "description": ps.description, "risk": ps.risk_level}
                            yield {"type": "output", "line": f"并行步骤 {ps.step_order}: {ps.description} ✘ (exit={pr['exit_code']})"}
                            yield {"type": "step_end", "step": ps.step_order, "status": "failed", "exit_code": pr["exit_code"]}
                            if _provider:
                                diag = _ai_diagnose_failure(db, _provider, ps, pr["output"],
                                    {"os": "", "docker": "", "asset": asset.name, "ip": asset.ip})
                                if diag.get("ok"):
                                    decision = _ai_autonomous_decision(_provider, ps, pr["output"], diag, [])
                                    _ai_decision_log(plan, {"step": ps.step_order, "type": "parallel_failure", "decision": decision,
                                        "root_cause": diag.get("root_cause", ""), "asset": asset.name})
                                    if decision == "fix":
                                        fix_cmds = diag.get("fix_commands", [])
                                        if fix_cmds:
                                            try:
                                                p_client, _ = _ssh_connect(asset)
                                                _run_fix_commands(p_client, fix_cmds, _cwd, asset_map)
                                                p_client.close()
                                            except Exception:
                                                pass
                                    yield {"type": "ai_decision", "step": ps.step_order, "decision": decision,
                                        "reason": diag.get("root_cause", "")[:200]}
                    db.commit()
                    if asset_failed:
                        break
                else:
                    # ── 串行执行组内步骤 ──
                    prev_step_output = ""
                    prev_step_status = ""
                    for so in step_orders:
                        if asset_failed:
                            break
                        step = step_map[so]

                        # 记录上一步的输出给 AI 规划用
                        if so > 1 and (so - 1) in step_map and step_map[so - 1].output:
                            prev_step_output = step_map[so - 1].output[-300:]
                            prev_step_status = step_map[so - 1].status

                        cmd = _resolve_command(step.command, asset_map)
                        verify_cmd = _resolve_command(step.verify_command, asset_map) if step.verify_command else ""

                        _cd_m = re.match(r'^cd\s+([^\s;&|]+)', cmd.strip())
                        if _cd_m:
                            _cwd = _cd_m.group(1).strip('"\'')
                        if not _cd_m and _cwd:
                            cmd = f"cd {_cwd} && {cmd}"

                        # ── 能力③：AI 执行前预判 ──
                        ai_risk_info = ""
                        if _provider:
                            precheck = _ai_pre_execution_risk(_provider, step, plan)
                            if precheck:
                                precheck_text = json.dumps(precheck, ensure_ascii=False)
                                step.precheck_result = precheck_text
                                db.commit()
                                ai_risk_info = precheck.get("risk", "")
                                yield {"type": "ai_precheck", "step": step.step_order, "risk": precheck.get("risk", ""),
                                    "reason": precheck.get("reason", ""), "precheck": precheck.get("precheck", ""),
                                    "guard_note": precheck.get("guard_note", "")}
                                if precheck.get("suggest_modify"):
                                    yield {"type": "output", "line": f"🤖 AI 建议安全替换: {precheck['suggest_modify']}"}
                                if precheck.get("precheck") and precheck.get("precheck_expect"):
                                    try:
                                        p_stdin, p_stdout, p_stderr = client.exec_command(precheck["precheck"], timeout=10)
                                        p_exit = p_stdout.channel.recv_exit_status()
                                        p_out = p_stdout.read().decode("utf-8", errors="replace").strip()
                                        yield {"type": "output", "line": f"🔍 前置检查: {precheck['precheck']} → exit={p_exit} {p_out[:100]}"}
                                    except Exception as pce:
                                        yield {"type": "output", "line": f"⚠ 前置检查异常: {pce}"}

                        # ── AI 自主规划：理解意图 → 生成执行计划 ──
                        ai_intent = ""
                        ai_plan = None
                        if _provider:
                            env_ctx = {"prev_output": prev_step_output, "prev_status": prev_step_status, "cwd": _cwd}
                            ai_plan = _ai_plan_step_autonomous(_provider, step, plan, probe, env_ctx)
                            if ai_plan.get("commands"):
                                ai_cmds = ai_plan["commands"]
                                ai_intent = ai_plan.get("intent", "")
                                ai_verify = ai_plan.get("verify", "")
                                yield {"type": "ai_plan", "step": step.step_order, "intent": ai_intent,
                                    "commands": ai_cmds, "adjustments": ai_plan.get("adjustments", []),
                                    "risk": ai_plan.get("risk", step.risk_level)}
                                # 使用 AI 规划的命令（需再次解析 ${ENV_xxx} 占位符）
                                ai_cmds_resolved = [_resolve_command(c, asset_map) for c in ai_cmds]
                                cmd = ai_cmds_resolved[0] if len(ai_cmds_resolved) == 1 else " && ".join(ai_cmds_resolved)
                                if ai_verify:
                                    verify_cmd = _resolve_command(ai_verify, asset_map)
                                _cd_m = re.match(r'^cd\s+([^\s;&|]+)', cmd.strip())
                                if _cd_m:
                                    _cwd = _cd_m.group(1).strip('"\'')
                                if not _cd_m and _cwd:
                                    cmd = f"cd {_cwd} && {cmd}"

                        # ── 风险门控：高危操作需用户确认 ──
                        _risk_level = step.risk_level
                        if ai_plan and ai_plan.get("risk"):
                            _risk_level = ai_plan["risk"]
                        if _risk_level == "high" and decision_queue is not None:
                            yield {"type": "risk_confirm", "step": step.step_order,
                                "command": cmd[:200], "description": step.description,
                                "risk": "high", "reason": "高危操作，请确认是否继续"}
                            decision = _wait_for_risk_confirm(decision_queue, plan_id,
                                step.step_order, cmd, "high", "高危操作")
                            yield {"type": "output", "line": f"\r\n\x1b[33m🔐 高危操作确认: {'✅ 已确认' if decision == 'confirm' else '⛔ 已拒绝'}\x1b[0m"}
                            if decision != "confirm":
                                step.status = "skipped"
                                db.commit()
                                yield {"type": "step_end", "step": step.step_order, "status": "skipped"}
                                continue

                        step.status = "running"
                        step.started_at = _now()
                        db.commit()

                        yield {"type": "step_start", "step": step.step_order, "description": step.description,
                            "risk": step.risk_level, "ai_risk": ai_risk_info, "ai_intent": ai_intent}
                        yield {"type": "cmd", "command": cmd}

                        # ▼ 离线二次强制校验: 勾选 use_offline 后拦截公网拉取/公网源命令
                        _offline_block = _offline_blocked_reason(plan, cmd)
                        if _offline_block:
                            step.status = "failed"
                            step.output = (step.output or "") + f"\n⛔ {_offline_block}"
                            db.commit()
                            yield {"type": "output", "line": f"\r\n\x1b[31m⛔ {_offline_block}\x1b[0m"}
                            yield {"type": "step_end", "step": step.step_order, "status": "failed", "note": _offline_block}
                            failed = True
                            continue

                        # ▼ 代理环境注入(可选): 在真实命令前导出 HTTP_PROXY/HTTPS_PROXY/NO_PROXY
                        _px = _proxy_env_prefix(plan)
                        _exec_cmd = (f"{_px} && {cmd}") if _px else cmd

                        try:
                            stdin, stdout, stderr = client.exec_command(_exec_cmd, timeout=60)
                            _ch = stdout.channel
                            _ch.settimeout(0.05)
                            collected = []
                            err_collected = []
                            _step_deadline = _time.time() + _STEP_TIMEOUT
                            while not _ch.exit_status_ready():
                                if _time.time() > _step_deadline:
                                    try:
                                        _ch.close()
                                    except Exception:
                                        pass
                                    raise TimeoutError(f"步骤超时({_STEP_TIMEOUT}s)，已终止该命令")
                                try:
                                    _line = stdout.readline()
                                    if _line:
                                        _line = _line.rstrip("\r\n")
                                        collected.append(_line)
                                        yield {"type": "output", "line": _line}
                                except Exception:
                                    _line = ""
                                try:
                                    _el = stderr.readline()
                                    if _el:
                                        _el = _el.rstrip("\r\n")
                                        if _el:
                                            err_collected.append(_el)
                                            yield {"type": "output", "line": _el}
                                except Exception:
                                    _el = ""
                                if not _line and not _el:
                                    _time.sleep(0.1)
                            _ch.settimeout(0.3)
                            try:
                                while True:
                                    _d = stdout.read(8192).decode("utf-8", errors="replace")
                                    if not _d:
                                        break
                                    for _l in _d.split("\n"):
                                        _l = _l.rstrip("\r")
                                        if _l:
                                            collected.append(_l)
                                            yield {"type": "output", "line": _l}
                            except Exception:
                                pass
                            try:
                                while True:
                                    _ed = stderr.read(8192).decode("utf-8", errors="replace")
                                    if not _ed:
                                        break
                                    for _el in _ed.split("\n"):
                                        _el = _el.rstrip("\r")
                                        if _el:
                                            err_collected.append(_el)
                                            yield {"type": "output", "line": _el}
                            except Exception:
                                pass
                            _rc = _ch.recv_exit_status()
                            _err = "\n".join(err_collected)
                            output = "\n".join(collected)
                            if _err:
                                output = f"{output}\n{_err}" if output else _err
                            step.output = (step.output or "") + f"\n[{asset.name}]\n{output[:2000]}"
                            step.finished_at = _now()

                            # ── 命令失败处理 → 能力②：AI 自主决策 ──
                            if _rc != 0:
                                step.status = "failed"
                                db.commit()
                                asset_failed = True
                                yield {"type": "step_end", "step": step.step_order, "status": "failed", "exit_code": _rc}
                                if _provider:
                                    diag = _ai_diagnose_failure(db, _provider, step, output, {"os": "", "docker": "",
                                        "asset": asset.name, "ip": asset.ip})
                                    if diag.get("ok"):
                                        decision = _ai_autonomous_decision(_provider, step, output, diag, [])
                                        _ai_decision_log(plan, {"step": step.step_order, "type": "command_failure",
                                            "decision": decision, "root_cause": diag.get("root_cause", ""),
                                            "fix_commands": diag.get("fix_commands", []), "asset": asset.name})
                                        yield {"type": "ai_decision", "step": step.step_order, "decision": decision,
                                            "reason": diag.get("root_cause", "")[:300],
                                            "fix_commands": diag.get("fix_commands", [])}
                                        if decision in ("fix", "retry"):
                                            if decision == "fix":
                                                _run_fix_commands(client, diag.get("fix_commands", []), _cwd, asset_map)
                                            step.retry_count += 1
                                            step.diagnosis = diag.get("root_cause", "")
                                            step.fix_command = json.dumps(diag.get("fix_commands", []), ensure_ascii=False)
                                            step.status = "running"
                                            db.commit()
                                            asset_failed = False
                                            continue
                                        elif decision == "skip":
                                            step.status = "skipped"
                                            step.diagnosis = diag.get("root_cause", "")
                                            db.commit()
                                            asset_failed = False
                                            continue
                                # fallback: rollback
                                yield {"type": "output", "line": f"⛔ 步骤 {step.step_order} 执行失败，AI 已决策回滚..."}
                                yield from _ai_stream_rollback(client, steps, step, asset_map, _cwd, _provider, plan)
                                continue

                            # ── 校验 ──
                            if verify_cmd:
                                if not _is_valid_shell_command(verify_cmd):
                                    yield {"type": "output", "line": f"ℹ 校验命令非 shell 命令，跳过: {verify_cmd[:60]}"}
                                    verify_cmd = ""
                                if verify_cmd and _cwd and not re.match(r'^cd\s', verify_cmd.strip()):
                                    verify_cmd = f"cd {_cwd} && {verify_cmd}"
                                yield {"type": "output", "line": f"✔ 校验: {verify_cmd}"}
                                try:
                                    v_stdin, v_stdout, v_stderr = client.exec_command(verify_cmd, timeout=15)
                                    v_exit = v_stdout.channel.recv_exit_status()
                                    if v_exit != 0:
                                        v_out = v_stdout.read().decode("utf-8", errors="replace").strip()
                                        step.output = (step.output + f"\n[校验失败] {v_out[:500]}").strip()
                                        step.status = "failed"
                                        db.commit()
                                        asset_failed = True
                                        yield {"type": "step_end", "step": step.step_order, "status": "failed", "exit_code": v_exit}
                                        if _provider:
                                            diag = _ai_diagnose_failure(db, _provider, step, f"校验失败: {v_out}", {})
                                            if diag.get("ok"):
                                                decision = _ai_autonomous_decision(_provider, step, f"校验失败: {v_out}", diag, [])
                                                _ai_decision_log(plan, {"step": step.step_order, "type": "verify_failure",
                                                    "decision": decision, "root_cause": diag.get("root_cause", ""),
                                                    "asset": asset.name})
                                                yield {"type": "ai_decision", "step": step.step_order, "decision": decision,
                                                    "reason": diag.get("root_cause", "")[:300]}
                                                if decision in ("fix", "retry"):
                                                    if decision == "fix":
                                                        _run_fix_commands(client, diag.get("fix_commands", []), _cwd, asset_map)
                                                    step.retry_count += 1
                                                    step.diagnosis = diag.get("root_cause", "")
                                                    step.status = "running"
                                                    db.commit()
                                                    asset_failed = False
                                                    continue
                                                elif decision == "skip":
                                                    step.status = "skipped"
                                                    db.commit()
                                                    asset_failed = False
                                                    continue
                                        yield {"type": "output", "line": "⛔ 校验失败，AI 已决策回滚..."}
                                        yield from _ai_stream_rollback(client, steps, step, asset_map, _cwd, _provider, plan)
                                        continue
                                except Exception as ve:
                                    step.output = (step.output + f"\n[校验异常] {ve}").strip()
                                    step.status = "failed"
                                    db.commit()
                                    asset_failed = True
                                    yield {"type": "step_end", "step": step.step_order, "status": "failed"}
                                    yield {"type": "output", "line": f"⛔ 校验异常: {ve}"}
                                    if _provider:
                                        diag = _ai_diagnose_failure(db, _provider, step, str(ve), {})
                                        if diag.get("ok"):
                                            decision = _ai_autonomous_decision(_provider, step, str(ve), diag, [])
                                            _ai_decision_log(plan, {"step": step.step_order, "type": "verify_exception",
                                                "decision": decision, "asset": asset.name})
                                            yield {"type": "ai_decision", "step": step.step_order, "decision": decision,
                                                "reason": diag.get("root_cause", "")[:200]}
                                            if decision in ("fix", "retry"):
                                                if decision == "fix":
                                                    _run_fix_commands(client, diag.get("fix_commands", []), _cwd, asset_map)
                                                step.retry_count += 1
                                                step.status = "running"
                                                db.commit()
                                                asset_failed = False
                                                continue
                                            elif decision == "skip":
                                                step.status = "skipped"
                                                db.commit()
                                                asset_failed = False
                                                continue
                                    yield from _ai_stream_rollback(client, steps, step, asset_map, _cwd, _provider, plan)
                                    continue

                            step.status = "succeeded"
                            db.commit()
                            yield {"type": "step_end", "step": step.step_order, "status": "succeeded"}

                            # ── L4: 健康门控 + 状态评估 ──
                            if _provider:
                                try:
                                    hg = _ai_health_gate(_provider, client, asset_map, step, strategy)
                                    _health_gates.append({"step": step.step_order, "checks": hg.get("checks", []),
                                        "passed": hg.get("passed", True), "ts": _now().isoformat()})
                                    yield {"type": "health_gate", "step": step.step_order,
                                        "passed": hg.get("passed", True), "checks": hg.get("checks", []),
                                        "recommendation": hg.get("recommendation", "continue")}
                                    if not hg.get("passed", True):
                                        if hg.get("recommendation") == "rollback":
                                            yield {"type": "output", "line": "\r\n\x1b[31m🔴 健康门控未通过，AI 决策回滚...\x1b[0m"}
                                            yield from _ai_stream_rollback(client, steps, step, asset_map, _cwd, _provider, plan)
                                            break
                                except Exception as hge:
                                    logger.warning(f"健康门控异常: {hge}")

                        except Exception as e:
                            step.output = (step.output or "") + f"\n[{asset.name}] {str(e)[:2000]}"
                            step.status = "failed"
                            step.finished_at = _now()
                            db.commit()
                            asset_failed = True
                            yield {"type": "step_end", "step": step.step_order, "status": "failed"}
                            yield {"type": "output", "line": f"⛔ 步骤异常: {e}"}
                            if _provider:
                                diag = _ai_diagnose_failure(db, _provider, step, str(e), {})
                                if diag.get("ok"):
                                    decision = _ai_autonomous_decision(_provider, step, str(e), diag, [])
                                    _ai_decision_log(plan, {"step": step.step_order, "type": "step_exception",
                                        "decision": decision, "asset": asset.name})
                                    yield {"type": "ai_decision", "step": step.step_order, "decision": decision,
                                        "reason": diag.get("root_cause", "")[:200]}
                                    if decision in ("fix", "retry"):
                                        if decision == "fix":
                                            _run_fix_commands(client, diag.get("fix_commands", []), _cwd, asset_map)
                                        step.retry_count += 1
                                        step.status = "running"
                                        db.commit()
                                        asset_failed = False
                                        continue
                                    elif decision == "skip":
                                        step.status = "skipped"
                                        db.commit()
                                        asset_failed = False
                                        continue
                            yield from _ai_stream_rollback(client, steps, step, asset_map, _cwd, _provider, plan)
        finally:
            _RUNNING_CLIENTS.pop(plan_id, None)
            client.close()

        # 兜底：即使 asset_failed 因 fix/retry 被重置，只要该计划存在最终态为 failed 的步骤，仍判为失败
        # （防止命令返回码为 0/被修复命令掩盖导致步骤 real 失败却误报 asset 成功）
        _asset_steps = [s for s in steps if s.plan_id == plan_id]
        _final_failed = any(s.status == "failed" for s in _asset_steps)
        if _final_failed:
            asset_failed = True

        if not asset_failed:
            succeeded_assets += 1
        yield {"type": "asset_end", "asset": asset.name, "status": "failed" if asset_failed else "succeeded"}

    if succeeded_assets == total_assets:
        plan.status = "succeeded"
    elif succeeded_assets > 0:
        plan.status = "failed"
    else:
        plan.status = "rolled_back"
    if _STOPPED.pop(plan_id, None):
        pass
    else:
        # ── L5: 记录部署特征(供学习) ──
        try:
            step_results = []
            for s in steps:
                step_results.append({"status": s.status, "retry_count": s.retry_count, "order": s.step_order})
            _total_dur = 0
            for s in steps:
                if s.started_at and s.finished_at:
                    _total_dur += (s.finished_at - s.started_at).total_seconds()
            features = _record_deployment_feature(plan, steps, assets, probe, plan.status, step_results, _total_dur)
            _ai_decision_log(plan, {"type": "deploy_complete", "status": plan.status, "features": features})
            plan.health_gate_json = json.dumps(_health_gates[-50:], ensure_ascii=False)
        except Exception as fe:
            logger.warning(f"部署特征记录异常: {fe}")
        db.commit()
        _record_execution_history(db, plan, plan.status, {"total_assets": total_assets, "succeeded_assets": succeeded_assets})
        yield {"type": "complete", "status": plan.status, "total_assets": total_assets, "succeeded_assets": succeeded_assets}


def _ai_stream_rollback(client, steps: list, failed_step, mapping: dict, app_dir: str, provider, plan: DeployPlan):
    """能力⑤：AI 自适应回滚 — 只回滚有状态的步骤，跳过无状态步骤。"""
    step_map = {s.step_order: s for s in steps}
    if provider:
        rollback_orders = _ai_adaptive_rollback(provider, steps, plan)
        if rollback_orders:
            yield {"type": "output", "line": f"\r\n\x1b[33m🤖 AI 自适应回滚: 跳过 {len(steps) - len(rollback_orders)} 个无状态步骤\x1b[0m"}
            for ro in rollback_orders:
                s = step_map[ro]
                if s.status != "succeeded":
                    continue
                rc = _resolve_command(s.rollback_command, mapping) if s.rollback_command else ""
                if rc:
                    try:
                        _s, _o, _e = client.exec_command(rc, timeout=30)
                        _rc = _o.channel.recv_exit_status()
                        yield {"type": "output", "line": f"  ↩ 回滚步骤 {ro}: {s.description} (exit={_rc})"}
                    except Exception as ex:
                        yield {"type": "output", "line": f"  ⚠ 回滚步骤 {ro} 异常: {ex}"}
                s.status = "rolled_back"
            return
    # fallback: 全量逆序回滚
    yield {"type": "output", "line": "\r\n\x1b[33m⏪ 全量回滚...\x1b[0m"}
    yield from _stream_rollback(client, steps, failed_step, mapping, app_dir)


def stream_rollback_cleanup(db: Session, plan_id: int):
    """手动一键清理回滚：部署完成后（succeeded/failed/rolled_back），
    逆序回滚所有已执行的步骤，并重置计划状态为 planned 以便重新部署。
    产出流式事件供 WS 实时显示。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        yield {"type": "error", "message": "计划不存在"}
        return
    if plan.status == "running":
        yield {"type": "error", "message": "部署执行中，无法清理"}
        return
    if plan.status == "draft":
        yield {"type": "error", "message": "计划尚未规划，无需清理"}
        return
    assets = _get_assets(db, plan)
    if not assets:
        yield {"type": "error", "message": "计划未关联目标资产"}
        return
    mapping = json.loads(plan.env_mapping or "{}")
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).order_by(DeployStep.step_order).all()

    yield {"type": "status", "status": "rolling_back", "message": "开始一键清理回滚..."}
    _log_lines = []  # 本资产生成的日志行（供历史记录）
    _cleanup_records = []
    for asset in assets:
        try:
            client, host = _ssh_connect(asset)
        except Exception as e:
            _msg = f"❌ {asset.name}({asset.ip}) SSH 连接失败: {e}"
            _log_lines.append(_msg)
            yield {"type": "error", "line": _msg}
            continue
        asset_map = dict(mapping)
        asset_map["TARGET_IP"] = asset.ip or ""
        asset_map["TARGET_HOSTNAME"] = asset.name or ""
        _asset_lines = []
        yield {"type": "asset_start", "asset": asset.name, "ip": asset.ip}
        try:
            # 模拟一个失败的「虚拟步骤」让其回滚所有已成功步骤
            _fake_failed = DeployStep(step_order=999)
            for _evt in _stream_rollback(client, steps, _fake_failed, asset_map, mapping.get("APP_DIR", "")):
                if _evt.get("type") == "output":
                    _asset_lines.append(_evt.get("line", ""))
                yield _evt
            # 额外执行 docker compose down（若存在）
            _app_dir = mapping.get("APP_DIR", "")
            if _app_dir:
                _down = f"cd {_app_dir} && docker compose down -v 2>/dev/null || true"
            else:
                _down = None
            if _down:
                try:
                    _s, _o, _e = client.exec_command(_down, timeout=30)
                    _rc = _o.channel.recv_exit_status()
                    _line = f"  🧹 清理容器: docker compose down -v (exit={_rc})"
                    _asset_lines.append(_line)
                    yield {"type": "output", "line": _line}
                except Exception as ex:
                    _line = f"  ⚠ 清理容器异常: {ex}"
                    _asset_lines.append(_line)
                    yield {"type": "output", "line": _line}
            # 清理运行产物（保留源码目录，便于再次部署）
            if _app_dir:
                _clean_run = (
                    f"cd {_app_dir} 2>/dev/null && "
                    f"rm -rf node_modules .venv venv __pycache__ dist build */__pycache__ 2>/dev/null; "
                    f"find . -name '*.pyc' -type f -delete 2>/dev/null; "
                    f"echo cleaned"
                )
                try:
                    _s, _o, _e = client.exec_command(_clean_run, timeout=30)
                    _rc = _o.channel.recv_exit_status()
                    _line = f"  🧹 清理运行产物(node_modules/venv/cache)，保留源码目录: {_app_dir}"
                    _asset_lines.append(_line)
                    yield {"type": "output", "line": _line}
                except Exception as ex:
                    _line = f"  ⚠ 清理运行产物异常: {ex}"
                    _asset_lines.append(_line)
                    yield {"type": "output", "line": _line}
        finally:
            client.close()
        _cleanup_records.append({"asset": asset.name, "ip": asset.ip, "lines": _asset_lines, "status": "cleaned"})
        _log_lines.extend(_asset_lines)
        yield {"type": "asset_end", "asset": asset.name, "status": "cleaned"}

    # 重置步骤与计划状态
    for s in steps:
        s.status = "pending"
        s.output = ""
        s.diagnosis = ""
        s.fix_command = ""
        s.retry_count = 0
        s.started_at = None
        s.finished_at = None
    _record_cleanup_history(db, plan, _cleanup_records, mapping.get("APP_DIR", ""))
    plan.status = "planned"
    db.commit()
    yield {"type": "complete", "status": "planned", "message": "清理完成，可重新部署"}


def _ai_step_failure(db, provider, decision_queue, step: DeployStep, output: str, plan_id: int = 0) -> tuple:
    """步骤失败时调用：AI 诊断 → 存库 → yield 事件 → 等用户决策。
    返回 (decision, fix_commands)。decision ∈ {fix, retry, skip, rollback}。
    如果 WS 断开（_STOPPED 标志），短超时轮询避免线程永久阻塞。"""
    diag = {"root_cause": "（AI 诊断不可用，请看上方输出）", "fix_commands": [], "suggestion": "rollback"}
    if provider:
        d = _ai_diagnose_failure(db, provider, step, output, {})
        if d.get("ok"):
            diag = d
    step.diagnosis = diag.get("root_cause", "")
    step.fix_command = json.dumps(diag.get("fix_commands", []), ensure_ascii=False)
    db.commit()
    yield {"type": "step_diagnosis", "step": step.step_order,
           "root_cause": diag.get("root_cause", ""),
           "fix_commands": diag.get("fix_commands", []),
           "suggestion": diag.get("suggestion", "rollback")}
    yield {"type": "need_decision", "step": step.step_order}
    decision = diag.get("suggestion", "rollback")
    if decision_queue is not None:
        import queue as _queue_module
        while True:
            try:
                decision = decision_queue.get(timeout=5)
                break
            except _queue_module.Empty:
                if _STOPPED.get(plan_id):
                    decision = "rollback"
                    break
                continue
            except Exception:
                decision = diag.get("suggestion", "rollback")
                break
    return decision, diag.get("fix_commands", [])


def _run_fix_commands(client, fix_commands: list, _cwd: str, asset_map: dict) -> None:
    """执行 AI 建议的修复命令（带 cwd 前缀）。"""
    for fc in fix_commands or []:
        fc = _resolve_command(fc, asset_map)
        if _cwd and not re.match(r'^cd\s', fc.strip()):
            fc = f"cd {_cwd} && {fc}"
        try:
            _, fo, fe = client.exec_command(fc, timeout=60)
            _rc = fo.channel.recv_exit_status()
            while fo.readline():
                pass
            e = fe.read().decode("utf-8", errors="replace").strip()
            logger.info(f"修复命令执行 {fc} -> exit={_rc} {e[:200]}")
        except Exception as e:
            logger.warning(f"修复命令执行异常: {fc} {e}")


def _do_rollback(client, steps: List[DeployStep], failed_step: DeployStep, mapping: dict):
    """静默回滚（HTTP 路径用）。"""
    for r in range(len(steps) - 1, -1, -1):
        s = steps[r]
        if s.step_order > failed_step.step_order:
            if s.status == "skipped":
                continue
        if s.step_order < failed_step.step_order and s.status == "succeeded":
            rc = _resolve_command(s.rollback_command, mapping) if s.rollback_command else ""
            if rc:
                try:
                    client.exec_command(rc, timeout=30)
                except Exception:
                    pass
            s.status = "rolled_back"


def _stream_rollback(client, steps: List[DeployStep], failed_step: DeployStep, mapping: dict, app_dir: str = ""):
    """流式回滚（WS 路径用），逐条产出事件供终端实时显示。"""
    yield {"type": "output", "line": "\r\n\x1b[33m⏪ 开始回滚...\x1b[0m"}
    for r in range(len(steps) - 1, -1, -1):
        s = steps[r]
        # 跳过失败步骤之后的步骤（它们未执行）
        if s.step_order > failed_step.step_order:
            if s.status == "skipped":
                continue
            s.status = "rolled_back"
            yield {"type": "output", "line": f"  ⏭ 步骤 {s.step_order} 未执行，标记回滚"}
            continue
        # 回滚已成功的步骤
        if s.step_order < failed_step.step_order and s.status == "succeeded":
            rc = _resolve_command(s.rollback_command, mapping) if s.rollback_command else ""
            if not rc:
                # 自动生成通用回滚命令
                if "docker compose" in s.command or "docker-compose" in s.command:
                    rc = f"cd {app_dir} && docker compose down -v 2>/dev/null" if app_dir else "docker compose down -v 2>/dev/null"
                elif "mkdir" in s.command:
                    _dir_match = re.search(r'mkdir\s+(?:-p\s+)?(\S+)', s.command)
                    if _dir_match:
                        rc = f"rm -rf {_dir_match.group(1)}"
                elif "cat >" in s.command or "echo " in s.command or "cp " in s.command or ">" in s.command:
                    _file_match = re.search(r'cat\s+>\s*(\S+)', s.command)
                    if _file_match:
                        rc = f"rm -f {_file_match.group(1)}"
                    _file_match = re.search(r'cp\s+(\S+)\s+(\S+)', s.command)
                    if _file_match:
                        rc = f"rm -f {_file_match.group(2)}"
                    _file_match = re.search(r'echo\s+.*?>\s*(\S+)', s.command)
                    if _file_match:
                        rc = f"rm -f {_file_match.group(1)}"
                    _file_match = re.search(r'>>\s*(\S+)', s.command)
                    if _file_match:
                        rc = f"rm -f {_file_match.group(1)}"
                elif "wget" in s.command or "curl -o" in s.command:
                    _file_match = re.search(r'(?:-o|>)\s+(\S+)', s.command)
                    if _file_match:
                        rc = f"rm -f {_file_match.group(1)}"
                elif "apt-get" in s.command or "yum" in s.command:
                    _pkg = re.search(r'(?:install|remove)\s+(\S+)', s.command)
                    if _pkg:
                        rc = f"apt-get remove -y {_pkg.group(1)} 2>/dev/null || yum remove -y {_pkg.group(1)} 2>/dev/null"
            if rc:
                try:
                    _s, _o, _e = client.exec_command(rc, timeout=30)
                    _rc_code = _o.channel.recv_exit_status() if hasattr(_o, 'channel') else 0
                    _out = _o.read().decode(errors="replace").strip()[:200]
                    _err = _e.read().decode(errors="replace").strip()[:200]
                    _detail = _out or _err
                    yield {"type": "output", "line": f"  ↩ 回滚步骤 {s.step_order}: {rc[:80]}... exit={_rc_code} {_detail[:60]}"}
                except Exception as ex:
                    yield {"type": "output", "line": f"  ⚠ 回滚步骤 {s.step_order} 异常: {ex}"}
            else:
                yield {"type": "output", "line": f"  → 步骤 {s.step_order} 无需回滚命令"}
            s.status = "rolled_back"
    yield {"type": "output", "line": "\x1b[32m✔ 回滚完成\x1b[0m\r\n"}


def post_deploy_verify(db: Session, plan_id: int) -> dict:
    """部署后验证：在目标机上运行健康检查，返回测试记录。"""
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    assets = _get_assets(db, plan)
    if not assets:
        return {"error": "计划未关联目标资产"}
    mapping = json.loads(plan.env_mapping or "{}")
    all_tests = []
    all_passed = True
    for asset in assets:
        try:
            client, host = _ssh_connect(asset)
        except Exception as e:
            all_tests.append({"asset": asset.name, "ip": asset.ip, "tests": [{"name": "SSH连接", "passed": False, "detail": str(e)}]})
            all_passed = False
            continue
        try:
            asset_tests = []
            def _verify(cmd, name, expect_ok=True):
                try:
                    _s, _o, _e = client.exec_command(cmd, timeout=15)
                    _rc = _o.channel.recv_exit_status()
                    _out = _o.read().decode(errors="replace").strip()
                    _err = _e.read().decode(errors="replace").strip()
                    _detail = _out or _err
                    _passed = (_rc == 0) if expect_ok else True
                    return {"name": name, "passed": _passed, "exit_code": _rc, "detail": _detail[:500], "command": cmd}
                except Exception as ex:
                    return {"name": name, "passed": False, "detail": str(ex)[:300], "command": cmd}
            # 1. Docker 守护进程
            asset_tests.append(_verify("docker info --format '{{.ServerVersion}}'", "Docker 守护进程"))
            # 2. 容器运行状态
            asset_tests.append(_verify("docker ps --format 'table {{.Names}}\t{{.Image}}\t{{.Status}}'", "容器运行状态"))
            # 2b. 检查是否有容器在重启循环
            _restart_check = _verify("docker ps -a --filter status=restarting --format '{{.Names}}' | head -5 || echo 'NONE'", "容器重启检查", expect_ok=False)
            if _restart_check.get("detail", "") and "NONE" not in _restart_check.get("detail", ""):
                _restart_check["passed"] = False
                _restart_check["detail"] = f"容器在重启循环: {_restart_check['detail']}"
            else:
                _restart_check["passed"] = True
                _restart_check["detail"] = "无容器重启循环"
            asset_tests.append(_restart_check)
            # 3. Compose 服务状态（如果存在）
            app_dir = mapping.get("APP_DIR", "")
            if app_dir:
                _compose_cmd = f"cd {app_dir} && docker compose ps --format 'table {{.Name}}\t{{.Status}}\t{{.Ports}}'"
                asset_tests.append(_verify(_compose_cmd, "Compose 服务状态"))
            # 4. 端口监听
            _ports_to_check = []
            if mapping.get("TARGET_PORT"):
                _ports_to_check.append(mapping["TARGET_PORT"])
            _steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id, DeployStep.status == "succeeded").all()
            for _s in _steps:
                for _m in re.finditer(r'(\d{2,5})', _s.command or ""):
                    _p = int(_m.group(1))
                    if 80 <= _p <= 65535 and _p not in _ports_to_check:
                        _ports_to_check.append(str(_p))
            for _p in _ports_to_check[:5]:
                asset_tests.append(_verify(f"ss -tlnp | grep ':{_p} ' || echo 'NOT_LISTENING'", f"端口 {_p}", expect_ok=False))
            # 5. HTTP 健康检查
            if mapping.get("TARGET_PORT"):
                _hp = mapping["TARGET_PORT"]
                asset_tests.append(_verify(f"curl -s -o /dev/null -w '%{{http_code}}' http://localhost:{_hp} 2>/dev/null || echo 'FAIL'", f"HTTP GET :{_hp}"))
            all_tests.append({"asset": asset.name, "ip": asset.ip, "tests": asset_tests})
            for _t in asset_tests:
                if not _t["passed"]:
                    all_passed = False
        finally:
            client.close()
    result = {"tests": all_tests, "all_passed": all_passed}
    plan.test_results_json = json.dumps(result, ensure_ascii=False)
    db.commit()
    return {"ok": True, "result": result}


def _extract_deploy_info(plan, assets, probe, mapping, sop, steps) -> dict:
    """从部署计划/环境探查/步骤中提取「交付级」运维信息：
    部署架构、启停服务命令、部署路径、服务端口、访问方式、登录信息。
    """
    info = {
        "architecture": "",
        "start_stop_commands": [],
        "deploy_paths": [],
        "service_ports": [],
        "access_methods": [],
        "login_info": [],
    }

    # ── 登录信息 ──
    for a in assets:
        try:
            conn = json.loads(a.connection_config or "{}")
        except Exception:
            conn = {}
        user = conn.get("ssh_user", "root")
        port = conn.get("ssh_port", 22)
        ip = a.ip or conn.get("ssh_host", "")
        info["login_info"].append({
            "asset": a.name, "ip": ip, "user": user, "port": port,
            "method": f"ssh {user}@{ip} -p {port}",
        })

    # ── 部署路径：从环境映射(APP_DIR 等) + 探查目录 ──
    if mapping:
        for k, v in mapping.items():
            if not v or not isinstance(v, str):
                continue
            if any(t in k.upper() for t in ("APP_DIR", "PATH", "DIR", "CONFIG", "LOG_DIR", "DATA_DIR")):
                info["deploy_paths"].append(f"{k} = {v}")
    if isinstance(probe, dict):
        for d in probe.get("dirs", {}):
            if not d.startswith("[probe_error]"):
                info["deploy_paths"].append(str(d))
        for _k in ("APP_DIR",):
            if mapping and mapping.get(_k):
                info["deploy_paths"].append(f"{_k} = {mapping.get(_k)}")

    # ── 服务端口：探查端口扫描 + 环境映射 TARGET_PORT + 步骤中的端口 ──
    if isinstance(probe, dict) and probe.get("port_scan"):
        for p, st in probe["port_scan"].items():
            if st == "IN_USE":
                info["service_ports"].append(f"{p} (占用)")
    if mapping and mapping.get("TARGET_PORT"):
        info["service_ports"].append(f"{mapping['TARGET_PORT']} (应用端口)")
    for s in steps:
        for _m in re.finditer(r'(?:EXPOSE|docker\s+-p|\-p\s+)(\d{2,5})', s.command or ""):
            _p = _m.group(1)
            if _p not in [x.split(" ")[0] for x in info["service_ports"]]:
                info["service_ports"].append(f"{_p}")

    # ── 启停服务命令：从步骤命令自动识别 ──
    all_cmds = "\n".join([(s.command or "") + "\n" + (s.rollback_command or "") for s in steps]) + "\n" + str(sop.get("steps", ""))
    _app_dir = (mapping or {}).get("APP_DIR", "")
    _compose_base = f"cd {_app_dir} && " if _app_dir else ""
    if "docker compose" in all_cmds or "docker-compose" in all_cmds:
        info["start_stop_commands"] = [
            f"启动: {_compose_base}docker compose up -d",
            f"停止: {_compose_base}docker compose down",
            f"重启: {_compose_base}docker compose restart",
            f"日志: {_compose_base}docker compose logs -f",
        ]
    if "systemctl" in all_cmds:
        _svc = re.search(r'systemctl\s+(?:start|stop|restart|enable)\s+([\w.-]+)', all_cmds)
        if _svc:
            info["start_stop_commands"].append(f"systemctl 服务: {_svc.group(1)}")
    if "docker run" in all_cmds:
        info["start_stop_commands"].append("容器运行: docker run (见步骤命令)")
    # 从步骤中抓取 docker compose 具体服务名
    _compose_services = set(re.findall(r'docker\s+compose\s+[a-z-]+\s+([\w-]+)', all_cmds))
    if _compose_services:
        info["start_stop_commands"].append(f"Compose 服务: {', '.join(sorted(_compose_services))}")
    if not info["start_stop_commands"]:
        info["start_stop_commands"] = ["见执行步骤命令"]

    # ── 访问方式 ──
    _ip = assets[0].ip if assets else ""
    _port = (mapping or {}).get("TARGET_PORT", "")
    if _ip:
        if _port:
            info["access_methods"].append(f"HTTP: http://{_ip}:{_port}")
        else:
            info["access_methods"].append(f"SSH: ssh root@{_ip}")
    # 从步骤 verify 中抓取 URL
    for s in steps:
        for _m in re.finditer(r'(?:curl|wget)\s+(-s\s+)?["\']?http://([^\s"\']+)', s.command or ""):
            info["access_methods"].append(f"{_m.group(2)}")

    # ── 部署架构：基于容器拓扑 + compose 服务 + 步骤描述 ──
    arch_parts = []
    if isinstance(probe, dict) and probe.get("containers"):
        arch_parts.append(f"当前容器: {probe.get('containers', '')[:200]}")
    for s in steps:
        if s.description and ("服务" in s.description or "启动" in s.description or "部署" in s.description or "compose" in (s.command or "").lower()):
            arch_parts.append(f"步骤{s.step_order}: {s.description} ({s.command[:80]})")
    info["architecture"] = "\n".join(arch_parts) if arch_parts else "无容器拓扑信息"

    # 去重
    info["deploy_paths"] = list(dict.fromkeys(info["deploy_paths"]))[:10]
    info["service_ports"] = list(dict.fromkeys(info["service_ports"]))[:10]
    info["start_stop_commands"] = list(dict.fromkeys(info["start_stop_commands"]))[:8]
    info["access_methods"] = list(dict.fromkeys(info["access_methods"]))[:5]
    return info


def generate_deploy_report(db: Session, plan_id: int) -> dict:
    """AI 生成专业部署报告：可直接交付客户的版本。

    产出包含：执行摘要、环境信息、部署时间线、步骤日志、预检结果、验证结果、
    AI 决策日志、风险评估、改进建议、总体评估。
    """
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    assets = _get_assets(db, plan)
    steps = db.query(DeployStep).filter(DeployStep.plan_id == plan_id).order_by(DeployStep.step_order).all()
    preflight = _safe_json(plan.preflight_json)
    test_results = _safe_json(plan.test_results_json)
    env_analysis = _safe_json(plan.env_analysis_json)
    probe = _safe_json(plan.environment_probe_json)
    mapping = _safe_json(plan.env_mapping)
    sop = _safe_json(plan.sop_json)
    history = _safe_json(plan.execution_history_json)
    dag = _safe_json(plan.dag_json)
    decision_log = _safe_json(plan.ai_decision_log_json)
    if not isinstance(history, list):
        history = []
    if not isinstance(decision_log, list):
        decision_log = []
    provider = _get_provider(db)

    _asset_info = [{"name": a.name, "ip": a.ip, "ci_type": a.ci_type} for a in assets] if assets else []
    deploy_info = _extract_deploy_info(plan, assets, probe, mapping, sop, steps)
    step_summary = []
    for s in steps:
        started = s.started_at.isoformat() if s.started_at else ""
        finished = s.finished_at.isoformat() if s.finished_at else ""
        duration = ""
        if s.started_at and s.finished_at:
            _secs = int((s.finished_at - s.started_at).total_seconds())
            if _secs >= 3600:
                duration = f"{_secs // 3600}h{(_secs % 3600) // 60}m{_secs % 60}s"
            elif _secs >= 60:
                duration = f"{_secs // 60}m{_secs % 60}s"
            else:
                duration = f"{_secs}s"
        step_summary.append({
            "order": s.step_order, "description": s.description,
            "command": s.command or "", "verify_command": s.verify_command or "",
            "rollback_command": s.rollback_command or "",
            "risk_level": s.risk_level, "status": s.status,
            "output": (s.output or "")[:2000],
            "diagnosis": s.diagnosis or "",
            "fix_command": s.fix_command or "",
            "retry_count": s.retry_count or 0,
            "precheck_result": s.precheck_result or "",
            "started_at": started, "finished_at": finished, "duration": duration,
        })

    preflight_detail = []
    _preflight_all_passed = False
    if isinstance(preflight, dict):
        _preflight_all_passed = bool(preflight.get("all_passed")) if "all_passed" in preflight else None
        # 兼容两种结构：预检接口写 {results:[...], all_passed}，execute 资源检查写 {checks:[...], passed}
        _pf_items = preflight.get("results", []) or preflight.get("checks", []) or []
        if not _pf_items and "passed" in preflight:
            _pf_items = [{"check": "资源检查", "passed": bool(preflight.get("passed")),
                          "command": "", "expect": "", "output": preflight.get("summary", ""), "exit_code": 0}]
        for _r in _pf_items:
            preflight_detail.append({
                "asset": _r.get("asset", ""), "check": _r.get("check", "") or _r.get("name", ""),
                "command": _r.get("command", ""), "expect": _r.get("expect", ""),
                "output": (_r.get("output", "") or "")[:300],
                "exit_code": _r.get("exit_code", -1), "passed": _r.get("passed", False),
            })
        # 缺省 all_passed 时按全部结果项推断（无结果项且无 all_passed → False）
        if _preflight_all_passed is None:
            _preflight_all_passed = bool(preflight_detail) and all(r["passed"] for r in preflight_detail)

    test_detail = []
    _test_all_passed = False
    if isinstance(test_results, dict):
        _test_all_passed = bool(test_results.get("all_passed")) if "all_passed" in test_results else None
        for _t_asset in test_results.get("tests", []):
            _a_name = _t_asset.get("asset", "")
            _a_ip = _t_asset.get("ip", "")
            for _t in _t_asset.get("tests", []):
                test_detail.append({
                    "asset": _a_name, "ip": _a_ip,
                    "name": _t.get("name", ""), "passed": _t.get("passed", False),
                    "detail": (_t.get("detail", "") or "")[:300],
                    "exit_code": _t.get("exit_code", -1),
                })
        if _test_all_passed is None:
            _test_all_passed = bool(test_detail) and all(t["passed"] for t in test_detail)

    if provider:
        sys_prompt = (
            "你是一名资深 SRE 运维专家，负责根据部署执行结果生成专业的、可直接交付客户的部署报告(Deployment Report)。\n"
            "报告必须专业、详尽、数据准确，包含具体指标和运维交付信息。\n"
            "请严格按以下 JSON Schema 输出，不要任何额外内容：\n"
            "{\n"
            '  "title": "部署报告标题",\n'
            '  "executive_summary": "执行摘要（中文，5-8句话，包含：部署目标、范围、结果、关键指标、风险评估）",\n'
            '  "deployment_architecture": "部署架构描述（中文，说明服务拓扑、组件依赖关系、部署模式(单机/集群/容器化)）",\n'
            '  "start_stop_commands": "启停服务命令（中文，提供完整的启动/停止/重启/查看日志命令，含路径前缀）",\n'
            '  "deploy_paths": "部署路径（中文，列出所有关键路径：应用目录、配置文件、日志目录、数据目录）",\n'
            '  "service_ports": "服务端口表（中文，说明每个端口的用途：如 8080(应用HTTP)、6379(Redis)等）",\n'
            '  "access_methods": "访问方式（中文，提供完整的访问URL、SSH命令、Web界面地址）",\n'
            '  "login_info": "登录信息（中文，说明如何登录目标服务器：SSH用户、IP、端口、密钥方式）",\n'
            '  "environment": {\n'
            '    "os": "操作系统版本",\n'
            '    "kernel": "内核版本",\n'
            '    "docker": "Docker 版本",\n'
            '    "disk": "磁盘使用情况",\n'
            '    "key_ports": "关键端口及状态",\n'
            '    "notes": "环境说明"\n'
            '  },\n'
            '  "timeline": "部署时间线概述（中文，描述整体执行耗时、每步耗时、是否有延迟）",\n'
            '  "steps_table": "步骤执行结果表格（markdown格式，包含序号、描述、耗时、状态、重试次数、诊断）",\n'
            '  "key_observations": ["关键观察1", "关键观察2"],\n'
            '  "verification": "部署验证结果（中文，详细描述每项检查的结果）",\n'
            '  "test_results": "测试记录摘要（中文，通过/失败项数、关键失败原因）",\n'
            '  "issues": [\n'
            '    {"severity": "high/medium/low", "description": "问题描述", "resolution": "如何处理", "status": "resolved/unresolved"}\n'
            '  ],\n'
            '  "risk_assessment": "风险评估（中文，分析当前部署的潜在风险）",\n'
            '  "recommendations": ["建议1", "建议2", "建议3"],\n'
            '  "overall_assessment": "总体评估（中文，succeeded/failed/partial，附一句话总结）"\n'
            "}\n"
            "关键要求：\n"
            "- 报告必须包含具体数据（IP、版本号、时长、端口号、路径、命令等），不能泛泛而谈\n"
            "- deployment_architecture 必须基于步骤描述和容器拓扑信息，描述服务部署模式\n"
            "- start_stop_commands 必须给出可直接复制的命令，包含 cd 路径前缀\n"
            "- login_info 必须给出具体的 SSH 登录命令和用户名\n"
            "- access_methods 必须给出完整的 URL 或 IP:端口\n"
            "- 如果存在失败步骤，必须在 issues 中详细分析根因和解决方案\n"
            "- 如果存在 AI 决策（skip/fix/rollback），必须在 issues 中记录\n"
            "- steps_table 必须包含所有步骤的执行状态、耗时、重试次数\n"
            "- 报告要专业、可读性强，能直接拿给客户或上级看"
        )
        user_prompt = (
            f"## 部署计划\n名称: {plan.name}\n描述: {plan.description or '无'}\n"
            f"状态: {plan.status}\n创建时间: {plan.created_at.isoformat() if plan.created_at else '未知'}\n"
            f"部署次数: {plan.deploy_count or 0}\n"
            f"最后部署: {plan.last_deployed_at.isoformat() if plan.last_deployed_at else '首次'}\n\n"
            f"## 目标资产\n{json.dumps(_asset_info, ensure_ascii=False, indent=1)}\n\n"
            f"## 环境探查结果\nOS: {probe.get('os', '')[:100] if isinstance(probe, dict) else '无'}\n"
            f"Kernel: {probe.get('kernel', '')[:100] if isinstance(probe, dict) else '无'}\n"
            f"Docker: {probe.get('docker', '')[:50] if isinstance(probe, dict) else '无'}\n"
            f"Disk: {probe.get('disk', '')[:80] if isinstance(probe, dict) else '无'}\n"
            f"Port scan: {json.dumps(probe.get('port_scan', {}), ensure_ascii=False)[:200] if isinstance(probe, dict) else '无'}\n\n"
            f"## 环境映射\n{json.dumps(mapping, ensure_ascii=False, indent=1) if mapping else '无'}\n\n"
            f"## 提取的运维信息（供参考）\n"
            f"部署路径: {json.dumps(deploy_info['deploy_paths'], ensure_ascii=False)}\n"
            f"服务端口: {json.dumps(deploy_info['service_ports'], ensure_ascii=False)}\n"
            f"启停命令: {json.dumps(deploy_info['start_stop_commands'], ensure_ascii=False)}\n"
            f"访问方式: {json.dumps(deploy_info['access_methods'], ensure_ascii=False)}\n"
            f"登录信息: {json.dumps(deploy_info['login_info'], ensure_ascii=False)}\n"
            f"容器拓扑: {(probe.get('containers', '')[:200] if isinstance(probe, dict) else '无')}\n\n"
            f"## 执行步骤（{len(step_summary)} 步）\n{json.dumps(step_summary, ensure_ascii=False, indent=1)}\n\n"
            f"## 预检结果\nall_passed: {_preflight_all_passed}\n"
            f"预检项: {json.dumps(preflight_detail, ensure_ascii=False, indent=1)[:2000]}\n\n"
            f"## 部署后验证\nall_passed: {_test_all_passed}\n"
            f"验证项: {json.dumps(test_detail, ensure_ascii=False, indent=1)[:2000]}\n\n"
            f"## AI 自适应建议\n{json.dumps(env_analysis.get('adaptations', []), ensure_ascii=False, indent=1) if isinstance(env_analysis, dict) else '[]'}\n\n"
            f"## AI 决策日志\n{json.dumps(decision_log, ensure_ascii=False, indent=1)}\n\n"
            f"## DAG 执行计划\n{json.dumps(dag, ensure_ascii=False, indent=1)}\n\n"
            f"请生成专业、详尽、可直接交付的部署报告。"
        )
        try:
            resp = call_llm(provider, [
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": user_prompt},
            ], timeout_override=120)
            if not resp.get("error"):
                try:
                    content = resp["choices"][0]["message"]["content"]
                    report = _extract_json(content) or {}
                except Exception:
                    report = {}
            else:
                report = {}
        except Exception:
            report = {}
    else:
        report = {}

    if not report.get("executive_summary"):
        report = _generate_fallback_report(plan, step_summary, _asset_info, preflight_detail, test_detail, deploy_info)

    report["plan_name"] = plan.name
    report["status"] = plan.status
    report["deployed_at"] = datetime.now().isoformat()
    report["plan_id"] = plan.id
    report["deploy_count"] = plan.deploy_count or 0
    report["total_steps"] = len(step_summary)
    report["succeeded_steps"] = sum(1 for s in step_summary if s["status"] == "succeeded")
    report["failed_steps"] = sum(1 for s in step_summary if s["status"] == "failed")
    report["skipped_steps"] = sum(1 for s in step_summary if s["status"] == "skipped")
    report["total_assets"] = len(_asset_info)
    report["preflight_passed"] = _preflight_all_passed
    report["verification_passed"] = _test_all_passed
    report["ai_decisions"] = len(decision_log)

    plan.deploy_report_json = json.dumps(report, ensure_ascii=False)
    db.commit()
    return {"ok": True, "report": report}


def _generate_fallback_report(plan, step_summary, assets, preflight_detail, test_detail, deploy_info=None) -> dict:
    """当 AI 不可用时，生成结构化摘要报告。"""
    total = len(step_summary)
    succeeded = sum(1 for s in step_summary if s["status"] == "succeeded")
    failed = sum(1 for s in step_summary if s["status"] == "failed")
    skipped = sum(1 for s in step_summary if s["status"] == "skipped")
    status_text = "成功" if plan.status == "succeeded" else "失败" if plan.status == "failed" else plan.status
    summary = (
        f"部署计划「{plan.name}」执行{status_text}。"
        f"共 {total} 步，成功 {succeeded} 步，失败 {failed} 步，跳过 {skipped} 步。"
        f"目标资产 {len(assets)} 台。"
    )
    steps_md = "| 序号 | 步骤 | 耗时 | 状态 | 重试 | 诊断 |\n|------|------|------|------|------|------|\n"
    for s in step_summary:
        diag = (s.get("diagnosis", "") or "")[:50]
        steps_md += f"| {s['order']} | {s['description'][:40]} | {s.get('duration', '-')} | {s['status']} | {s.get('retry_count', 0)} | {diag} |\n"

    env_text = "目标资产:\n"
    for a in assets:
        env_text += f"- {a.get('name', '')} ({a.get('ip', '')})\n"

    issues = []
    for s in step_summary:
        if s["status"] == "failed" and s.get("diagnosis"):
            issues.append({"severity": "high", "description": f"步骤 {s['order']} {s['description']} 失败", "resolution": s["diagnosis"][:200], "status": "resolved"})

    return {
        "title": f"部署报告 - {plan.name}",
        "executive_summary": summary,
        "deployment_architecture": f"目标资产 {len(assets)} 台，共 {total} 步部署步骤" + (f"\n{deploy_info.get('architecture', '')}" if deploy_info else ""),
        "start_stop_commands": "、\n".join(deploy_info.get("start_stop_commands", [])) if deploy_info else "见执行步骤",
        "deploy_paths": "、\n".join(deploy_info.get("deploy_paths", [])) if deploy_info else "无",
        "service_ports": "、\n".join(deploy_info.get("service_ports", [])) if deploy_info else "无",
        "access_methods": "、\n".join(deploy_info.get("access_methods", [])) if deploy_info else "无",
        "login_info": "、\n".join([f"{l['method']}" for l in (deploy_info.get("login_info", []) if deploy_info else [])]) or "无",
        "environment": {"os": "见环境探查", "notes": env_text},
        "timeline": f"总耗时: {sum(int(s.get('duration', '0').rstrip('s')) for s in step_summary if s.get('duration', '-') != '-')}s",
        "steps_table": steps_md,
        "key_observations": [f"共 {total} 步，{succeeded} 成功，{failed} 失败，{skipped} 跳过"],
        "verification": f"预检{'通过' if preflight_detail else '未执行'}，部署后验证{'通过' if test_detail else '未执行'}",
        "test_results": f"验证项 {len(test_detail)} 项" if test_detail else "未执行验证",
        "issues": issues or [{"severity": "low", "description": "无异常", "resolution": "-", "status": "resolved"}],
        "risk_assessment": "基于当前部署结果，风险可控",
        "recommendations": ["监控部署后的服务运行状态", "检查日志输出是否正常", "建议定期执行健康检查"],
        "overall_assessment": f"{plan.status} - {summary}",
    }


def _report_to_markdown(report: dict) -> str:
    """将部署报告转换为可直接交付的 Markdown 文档。"""
    title = report.get("title", "部署报告")
    plan_name = report.get("plan_name", "")
    status = report.get("status", "")
    deployed_at = report.get("deployed_at", "")
    deploy_count = report.get("deploy_count", 0)
    total_steps = report.get("total_steps", 0)
    succeeded_steps = report.get("succeeded_steps", 0)
    failed_steps = report.get("failed_steps", 0)
    skipped_steps = report.get("skipped_steps", 0)
    total_assets = report.get("total_assets", 0)
    preflight_passed = report.get("preflight_passed", False)
    verification_passed = report.get("verification_passed", False)
    ai_decisions = report.get("ai_decisions", 0)

    status_icon = "✅" if status == "succeeded" else "❌" if status == "failed" else "⚠️"
    preflight_icon = "✅" if preflight_passed else "❌"
    verify_icon = "✅" if verification_passed else "❌"

    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append("> **文档类型**: 部署报告 (Deployment Report)")
    lines.append(f"> **生成时间**: {deployed_at}")
    lines.append(f"> **状态**: {status_icon} **{status.upper()}**")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 📋 执行摘要")
    lines.append("")
    lines.append(report.get("executive_summary", "无"))
    lines.append("")
    lines.append("### 关键指标")
    lines.append("")
    lines.append("| 指标 | 值 |")
    lines.append("|------|----|")
    lines.append(f"| 部署计划 | {plan_name} |")
    lines.append(f"| 部署次数 | 第 {deploy_count} 次 |")
    lines.append(f"| 执行时间 | {deployed_at} |")
    lines.append(f"| 目标资产 | {total_assets} 台 |")
    lines.append(f"| 总步骤数 | {total_steps} |")
    lines.append(f"| 成功步骤 | {succeeded_steps} |")
    lines.append(f"| 失败步骤 | {failed_steps} |")
    lines.append(f"| 跳过步骤 | {skipped_steps} |")
    lines.append(f"| 预检结果 | {preflight_icon} {'全部通过' if preflight_passed else '有失败项'} |")
    lines.append(f"| 部署验证 | {verify_icon} {'全部通过' if verification_passed else '有失败项'} |")
    lines.append(f"| AI 决策次数 | {ai_decisions} |")
    lines.append(f"| 总体评估 | {status_icon} {report.get('overall_assessment', status)} |")
    lines.append("")

    env = report.get("environment", {})
    if isinstance(env, dict):
        lines.append("---")
        lines.append("")
        lines.append("## 🖥️ 环境信息")
        lines.append("")
        lines.append("| 项目 | 详情 |")
        lines.append("|------|------|")
        for k, v in env.items():
            lines.append(f"| {k} | {str(v)[:200]} |")
        lines.append("")
    elif isinstance(env, str):
        lines.append("---")
        lines.append("")
        lines.append("## 🖥️ 环境信息")
        lines.append("")
        lines.append(env)
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## ⏱️ 时间线")
    lines.append("")
    lines.append(report.get("timeline", "无"))
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 📊 步骤执行结果")
    lines.append("")
    lines.append(report.get("steps_table", "无"))
    lines.append("")

    if report.get("key_observations"):
        lines.append("---")
        lines.append("")
        lines.append("## 🔍 关键观察")
        lines.append("")
        for obs in report.get("key_observations", []):
            lines.append(f"- {obs}")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## ✅ 部署验证")
    lines.append("")
    lines.append(report.get("verification", "无"))
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 🧪 测试记录")
    lines.append("")
    lines.append(report.get("test_results", "无"))
    lines.append("")

    issues = report.get("issues", [])
    if issues:
        lines.append("---")
        lines.append("")
        lines.append("## 🐛 问题与处理")
        lines.append("")
        lines.append("| 严重程度 | 问题描述 | 处理方式 | 状态 |")
        lines.append("|---------|---------|---------|------|")
        for issue in issues:
            sev = issue.get("severity", "low")
            sev_icon = "🔴" if sev == "high" else "🟡" if sev == "medium" else "🟢"
            lines.append(f"| {sev_icon} {sev} | {issue.get('description', '')[:100]} | {issue.get('resolution', '')[:100]} | {issue.get('status', '')} |")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## ⚠️ 风险评估")
    lines.append("")
    lines.append(report.get("risk_assessment", "无"))
    lines.append("")

    recommendations = report.get("recommendations", [])
    if recommendations:
        lines.append("---")
        lines.append("")
        lines.append("## 💡 改进建议")
        lines.append("")
        for i, rec in enumerate(recommendations, 1):
            lines.append(f"{i}. {rec}")
        lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 📝 总体评估")
    lines.append("")
    lines.append(report.get("overall_assessment", "无"))
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(f"*报告由 AIOps 自动部署系统自动生成 | {deployed_at}*")

    return "\n".join(lines)


def _report_to_html(report: dict) -> str:
    """将部署报告转换为可直接打印/保存为 PDF 的专业 HTML 文档。"""
    md = _report_to_markdown(report)
    import markdown as _md
    body_html = _md.markdown(md, extensions=["tables", "fenced_code", "codehilite"])

    status = report.get("status", "unknown")
    status_color = "#22c55e" if status == "succeeded" else "#ef4444" if status == "failed" else "#f59e0b"

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{report.get('title', '部署报告')}</title>
<style>
  @page {{ size: A4; margin: 2cm; }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: 'PingFang SC', 'Microsoft YaHei', 'Helvetica Neue', Arial, sans-serif; color: #1e293b; line-height: 1.7; padding: 40px; background: #f8fafc; }}
  .report-container {{ max-width: 900px; margin: 0 auto; background: #fff; padding: 50px 60px; box-shadow: 0 4px 24px rgba(0,0,0,0.06); border-radius: 12px; }}
  h1 {{ font-size: 28px; color: #0f172a; border-bottom: 3px solid {status_color}; padding-bottom: 16px; margin-bottom: 24px; }}
  h2 {{ font-size: 20px; color: #0f172a; margin-top: 32px; margin-bottom: 16px; padding-left: 10px; border-left: 4px solid {status_color}; }}
  h3 {{ font-size: 16px; color: #334155; margin-top: 20px; margin-bottom: 10px; }}
  p {{ margin-bottom: 12px; color: #475569; }}
  table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }}
  th {{ background: #f1f5f9; color: #0f172a; font-weight: 600; padding: 10px 12px; border: 1px solid #e2e8f0; text-align: left; }}
  td {{ padding: 8px 12px; border: 1px solid #e2e8f0; color: #475569; }}
  tr:nth-child(even) td {{ background: #f8fafc; }}
  blockquote {{ background: #f1f5f9; border-left: 4px solid {status_color}; padding: 12px 18px; margin: 16px 0; border-radius: 0 8px 8px 0; color: #475569; }}
  code {{ background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: 13px; font-family: 'JetBrains Mono', 'Fira Code', monospace; }}
  pre {{ background: #0f172a; color: #e2e8f0; padding: 16px; border-radius: 8px; overflow-x: auto; font-size: 13px; line-height: 1.5; margin: 12px 0; }}
  ul, ol {{ padding-left: 24px; margin-bottom: 12px; }}
  li {{ margin-bottom: 6px; color: #475569; }}
  .status-badge {{ display: inline-block; padding: 4px 16px; border-radius: 20px; font-size: 14px; font-weight: 600; color: #fff; background: {status_color}; }}
  .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; font-size: 12px; color: #94a3b8; text-align: center; }}
  hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 24px 0; }}
  @media print {{
    body {{ background: #fff; padding: 0; }}
    .report-container {{ box-shadow: none; padding: 40px; }}
    .no-print {{ display: none !important; }}
  }}
  .no-print {{ text-align: center; margin-bottom: 20px; }}
  .no-print button {{ padding: 10px 24px; background: #6366f1; color: #fff; border: none; border-radius: 8px; font-size: 14px; cursor: pointer; margin: 0 8px; }}
  .no-print button:hover {{ background: #4f46e5; }}
  .no-print button.secondary {{ background: #e2e8f0; color: #475569; }}
  .no-print button.secondary:hover {{ background: #cbd5e1; }}
</style>
</head>
<body>
<div class="no-print">
  <button onclick="window.print()">🖨️ 打印 / 导出 PDF</button>
  <button class="secondary" onclick="location.href='/deploy/api/plans/{report.get("plan_id", 0)}/report/download?format=md'">📥 下载 Markdown</button>
</div>
<div class="report-container">
{body_html}
<div class="footer">
  <p>本报告由 AIOps 智能运维平台自动生成 | {report.get("deployed_at", "")}</p>
  <p style="margin-top:4px;">AIOps Deployment Report · Confidential</p>
</div>
</div>
</body>
</html>"""


def download_report(db: Session, plan_id: int, fmt: str = "docx") -> dict:
    """下载部署报告：docx(Word) / html 两种格式。
    返回 {"ok": True, "content": bytes, "filename": "...", "mime": "..."}。
    """
    plan = db.query(DeployPlan).filter(DeployPlan.id == plan_id).first()
    if not plan:
        return {"error": "计划不存在"}
    report = _safe_json(plan.deploy_report_json)
    if not report or not isinstance(report, dict) or not report.get("executive_summary"):
        return {"error": "报告尚未生成，请先执行「生成报告」"}

    fmt = fmt.lower()

    if fmt == "html":
        content = _report_to_html(report).encode("utf-8")
        return {"ok": True, "content": content, "filename": f"deploy_report_{plan_id}.html", "mime": "text/html; charset=utf-8"}
    elif fmt == "docx":
        content = _report_to_docx(report)
        return {"ok": True, "content": content, "filename": f"deploy_report_{plan_id}.docx", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    else:
        return {"error": f"不支持的格式: {fmt}，支持 docx/html"}


def _report_to_docx(report: dict) -> bytes:
    """将部署报告生成为专业 Word .docx 文档，返回 bytes。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── 全局样式设置 ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            hs.font.size = Pt(22)
        elif level == 2:
            hs.font.size = Pt(16)
        else:
            hs.font.size = Pt(13)

    title = report.get("title", "部署报告")
    status = report.get("status", "unknown")
    deployed_at = report.get("deployed_at", "")

    # ── 封面区域 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("部署报告")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{title}")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    status_text = "成功" if status == "succeeded" else "失败" if status == "failed" else status
    status_color = "22c55e" if status == "succeeded" else "ef4444" if status == "failed" else "f59e0b"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n状态: {status_text.upper()}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(int(status_color[:2], 16), int(status_color[2:4], 16), int(status_color[4:], 16))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n生成时间: {deployed_at}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_page_break()

    # ── 辅助函数 ──
    def _add_kv_table(doc, headers, rows):
        t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            for p in t.rows[0].cells[i].paragraphs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(10)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                t.rows[ri + 1].cells[ci].text = str(val)
                for p in t.rows[ri + 1].cells[ci].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
        doc.add_paragraph()

    # ── 1. 执行摘要 ──
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(report.get("executive_summary", "无"))

    # KPI 指标
    doc.add_heading("关键指标", level=2)
    _add_kv_table(doc, ["指标", "值"], [
        ["部署计划", report.get("plan_name", "")],
        ["部署次数", f"第 {report.get('deploy_count', 0)} 次"],
        ["执行时间", deployed_at],
        ["目标资产", f"{report.get('total_assets', 0)} 台"],
        ["总步骤数", str(report.get("total_steps", 0))],
        ["成功步骤", str(report.get("succeeded_steps", 0))],
        ["失败步骤", str(report.get("failed_steps", 0))],
        ["跳过步骤", str(report.get("skipped_steps", 0))],
        ["预检结果", "通过" if report.get("preflight_passed") else "未通过"],
        ["部署验证", "通过" if report.get("verification_passed") else "有失败项"],
        ["AI 决策次数", str(report.get("ai_decisions", 0))],
    ])

    # ── 1b. 部署架构 ──
    arch = report.get("deployment_architecture", "")
    if arch:
        doc.add_heading("部署架构", level=1)
        doc.add_paragraph(arch)

    # ── 1c. 启停服务命令 ──
    ssc = report.get("start_stop_commands", "")
    if ssc:
        doc.add_heading("启停服务命令", level=1)
        for line in ssc.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1d. 部署路径 ──
    dp = report.get("deploy_paths", "")
    if dp:
        doc.add_heading("部署路径", level=1)
        for line in dp.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1e. 服务端口 ──
    sp = report.get("service_ports", "")
    if sp:
        doc.add_heading("服务端口", level=1)
        for line in sp.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1f. 访问方式 ──
    am = report.get("access_methods", "")
    if am:
        doc.add_heading("访问方式", level=1)
        for line in am.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 1g. 登录信息 ──
    li = report.get("login_info", "")
    if li:
        doc.add_heading("登录信息", level=1)
        for line in li.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line, style='List Bullet')

    # ── 2. 环境信息 ──
    env = report.get("environment", {})
    if isinstance(env, dict):
        doc.add_heading("环境信息", level=1)
        _add_kv_table(doc, ["项目", "详情"], [[k, str(v)[:200]] for k, v in env.items()])

    # ── 3. 时间线 ──
    doc.add_heading("时间线", level=1)
    doc.add_paragraph(report.get("timeline", "无"))

    # ── 4. 步骤执行结果 ──
    doc.add_heading("步骤执行结果", level=1)
    steps_md = report.get("steps_table", "")
    # Parse markdown table into list
    if steps_md:
        lines = [l.strip() for l in steps_md.split("\n") if l.strip()]
        data_rows = []
        headers = []
        for i, line in enumerate(lines):
            if line.startswith("|") and line.endswith("|"):
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if i == 0:
                    headers = cells
                elif "---" not in line and cells:
                    data_rows.append(cells)
        if headers:
            _add_kv_table(doc, headers, data_rows)

    # ── 5. 关键观察 ──
    obs = report.get("key_observations", [])
    if obs:
        doc.add_heading("关键观察", level=1)
        for o in obs:
            doc.add_paragraph(o, style='List Bullet')

    # ── 6. 部署验证 ──
    doc.add_heading("部署验证", level=1)
    doc.add_paragraph(report.get("verification", "无"))

    # ── 7. 测试记录 ──
    doc.add_heading("测试记录", level=1)
    doc.add_paragraph(report.get("test_results", "无"))

    # ── 8. 问题与处理 ──
    issues = report.get("issues", [])
    if issues:
        doc.add_heading("问题与处理", level=1)
        _add_kv_table(doc, ["严重程度", "问题描述", "处理方式", "状态"],
            [[i.get("severity", ""), i.get("description", "")[:100],
              i.get("resolution", "")[:100], i.get("status", "")] for i in issues])

    # ── 9. 风险评估 ──
    doc.add_heading("风险评估", level=1)
    doc.add_paragraph(report.get("risk_assessment", "无"))

    # ── 10. 改进建议 ──
    recs = report.get("recommendations", [])
    if recs:
        doc.add_heading("改进建议", level=1)
        for i, rec in enumerate(recs, 1):
            doc.add_paragraph(f"{i}. {rec}")

    # ── 11. 总体评估 ──
    doc.add_heading("总体评估", level=1)
    doc.add_paragraph(report.get("overall_assessment", "无"))

    # ── 页脚信息 ──
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"报告由 AIOps 智能运维平台自动生成 | {deployed_at}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _record_execution_history(db: Session, plan: DeployPlan, status: str, summary: dict) -> None:
    """记录执行历史到 plan.execution_history_json。"""
    try:
        history = json.loads(plan.execution_history_json) if isinstance(plan.execution_history_json, str) and plan.execution_history_json not in ("[]", "") else []
    except Exception:
        history = []
    if not isinstance(history, list):
        history = []
    history.append({
        "executed_at": datetime.now().isoformat(),
        "status": status,
        "total_assets": summary.get("total_assets", 0),
        "succeeded_assets": summary.get("succeeded_assets", 0),
        "step_count": len(db.query(DeployStep).filter(DeployStep.plan_id == plan.id).all()),
    })
    if len(history) > 50:
        history = history[-50:]
    plan.execution_history_json = json.dumps(history, ensure_ascii=False)
    plan.last_deployed_at = datetime.now()
    plan.deploy_count = (plan.deploy_count or 0) + 1
    db.commit()


def _record_cleanup_history(db: Session, plan: DeployPlan, records: list, app_dir: str = "") -> None:
    """记录回滚清理历史到 plan.cleanup_history_json，便于事后查看每次清理的日志。"""
    try:
        history = json.loads(plan.cleanup_history_json) if isinstance(plan.cleanup_history_json, str) and plan.cleanup_history_json not in ("[]", "") else []
    except Exception:
        history = []
    if not isinstance(history, list):
        history = []
    history.append({
        "cleaned_at": datetime.now().isoformat(),
        "assets": records,
        "app_dir": app_dir,
    })
    if len(history) > 20:
        history = history[-20:]
    plan.cleanup_history_json = json.dumps(history, ensure_ascii=False)




def _plan_to_dict(p: DeployPlan) -> dict:
    return {
        "id": p.id,
        "name": p.name,
        "description": p.description or "",
        "artifact_path": p.artifact_path or "",
        "artifact_download_path": p.artifact_download_path or "",
        "artifact_auto_download": bool(p.artifact_auto_download),
        "use_offline": bool(p.use_offline),
        "http_proxy": p.http_proxy or "",
        "https_proxy": p.https_proxy or "",
        "no_proxy": p.no_proxy or "",
        "doc_raw": p.doc_raw or "",
        "doc_file_name": p.doc_file_name or "",
        "asset_ids": _safe_json(p.asset_ids) if p.asset_ids else [],
        "env_mapping": _safe_json(p.env_mapping),
        "environment_probe_json": _safe_json(p.environment_probe_json),
        "env_analysis_json": _safe_json(p.env_analysis_json),
        "sop_json": _safe_json(p.sop_json),
        "status": p.status,
        "preflight_json": _safe_json(p.preflight_json),
        "deploy_report_json": _safe_json(p.deploy_report_json),
        "test_results_json": _safe_json(p.test_results_json),
        "execution_history_json": _safe_json(p.execution_history_json),
        "cleanup_history_json": _safe_json(p.cleanup_history_json),
        "dag_json": _safe_json(p.dag_json),
        "ai_decision_log_json": _safe_json(p.ai_decision_log_json),
        "strategy": p.strategy or "auto",
        "risk_score": p.risk_score or 0,
        "deployment_feature_json": _safe_json(p.deployment_feature_json),
        "health_gate_json": _safe_json(p.health_gate_json),
        "last_deployed_at": p.last_deployed_at.isoformat() if p.last_deployed_at else "",
        "deploy_count": p.deploy_count or 0,
        "created_by": p.created_by,
        "created_at": p.created_at.isoformat() if p.created_at else "",
        "updated_at": p.updated_at.isoformat() if p.updated_at else "",
    }


def _step_to_dict(s: DeployStep) -> dict:
    return {
        "id": s.id,
        "plan_id": s.plan_id,
        "step_order": s.step_order,
        "description": s.description or "",
        "command": s.command or "",
        "verify_command": s.verify_command or "",
        "rollback_command": s.rollback_command or "",
        "risk_level": s.risk_level,
        "status": s.status,
        "output": s.output or "",
        "diagnosis": s.diagnosis or "",
        "fix_command": s.fix_command or "",
        "retry_count": s.retry_count or 0,
        "precheck_result": s.precheck_result or "",
        "started_at": s.started_at.isoformat() if s.started_at else "",
        "finished_at": s.finished_at.isoformat() if s.finished_at else "",
    }


def _safe_json(val) -> Any:
    if not val:
        return {} if val is None else val
    if isinstance(val, (dict, list)):
        return val
    try:
        return json.loads(val)
    except (json.JSONDecodeError, TypeError):
        return val
"""知识库 RAG 服务：文档解析 / 切片 / TF-IDF 向量化 / 语义检索.

设计原则：
- 零新依赖：TF-IDF 用 numpy（已有），不依赖 sentence-transformers/langchain
- 双模式 Embedding：TF-IDF 本地（默认）→ 升级后可接 LLM Provider 的 /embeddings API
- 向量存储：embedding 存 JSON 字符串（兼容 SQLite），升级 pgvector 后改 vector 类型
- 检索：内存余弦相似度，适合 <1 万切片；超过后建议升级 pgvector + IVFFlat
"""
import json
import math
import os
import re
import threading
from collections import Counter
from datetime import datetime
from typing import Dict, List, Optional, Tuple

from sqlalchemy.orm import Session

from app.models import KbDocument, KbChunk


# ─── 文档解析 ───────────────────────────────────────────────────

def parse_document(file_path: str, file_ext: str) -> str:
    """解析文档返回纯文本. 支持 txt/md（原生）/ pdf（pypdf 可选）/ docx（python-docx 可选）."""
    ext = (file_ext or "").lower().lstrip(".")
    try:
        if ext in ("txt", "md", "markdown", ""):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        if ext == "pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return f"[解析失败：未安装 pypdf，无法解析 PDF。请运行 pip install pypdf 后重新上传]"
            reader = PdfReader(file_path)
            texts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                texts.append(t)
            return "\n".join(texts)
        if ext == "docx":
            try:
                import docx
            except ImportError:
                return f"[解析失败：未安装 python-docx，无法解析 Word。请运行 pip install python-docx 后重新上传]"
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        # 未知扩展名按文本读
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[解析异常: {e}]"


# ─── 文本切片 ───────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    """按字符长度切片，保留重叠. 中文友好（按字符而非单词）.

    优先按段落边界切分，避免截断句子；段落过长时按 chunk_size 硬切。
    """
    if not text or not text.strip():
        return []
    # 先按双换行（段落）切分
    paragraphs = re.split(r'\n\s*\n', text)
    chunks: List[str] = []
    buf = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # 段落本身超长：按 chunk_size 硬切
        if len(para) > chunk_size:
            if buf:
                chunks.append(buf)
                buf = ""
            for i in range(0, len(para), chunk_size - overlap):
                chunks.append(para[i:i + chunk_size])
            continue
        # 累积到 buf，超长则切出
        if len(buf) + len(para) + 1 > chunk_size and buf:
            chunks.append(buf)
            # 保留 overlap 字符作为下一块开头
            buf = buf[-overlap:] + "\n" + para if overlap else para
        else:
            buf = (buf + "\n" + para) if buf else para
    if buf:
        chunks.append(buf)
    # 过滤过短切片（<20 字符无检索价值）
    return [c for c in chunks if len(c) >= 20]


# ─── TF-IDF 向量化 ──────────────────────────────────────────────

_TOKEN_RE = re.compile(r'[\u4e00-\u9fa5]|[a-zA-Z0-9]+')


def tokenize(text: str) -> List[str]:
    """中英文混合分词：中文按字、英文/数字按词. 不依赖 jieba，零外部依赖."""
    if not text:
        return []
    return _TOKEN_RE.findall(text.lower())


# 全局 IDF 缓存：{cache_key: idf_map}
# cache_key 用 "tfidf:v1"，文档变更时清空
_IDF_CACHE: Dict[str, dict] = {}
_IDF_CACHE_LOCK = threading.Lock()
_IDF_CACHE_VERSION = "tfidf:v1"


def _build_idf_map(db: Session) -> dict:
    """从所有已索引 chunk 的文本构建全局 IDF 映射. 首次查询时构建，之后复用."""
    with _IDF_CACHE_LOCK:
        if _IDF_CACHE_VERSION in _IDF_CACHE:
            return _IDF_CACHE[_IDF_CACHE_VERSION]
    # 加载所有 chunk 文本（只取 content 列，避免内存爆炸）
    chunks = db.query(KbChunk.content).filter(KbChunk.embedding_mode == "tfidf").all()
    if not chunks:
        idf = {}
    else:
        N = len(chunks)
        df: Counter = Counter()
        for (content,) in chunks:
            tokens = set(tokenize(content))
            for w in tokens:
                df[w] += 1
        # IDF = log((N+1)/(df+1)) + 1 （平滑，避免 df=N 时 idf=0）
        idf = {w: math.log((N + 1) / (c + 1)) + 1 for w, c in df.items()}
    with _IDF_CACHE_LOCK:
        _IDF_CACHE[_IDF_CACHE_VERSION] = idf
    return idf


def invalidate_idf_cache():
    """文档/切片变更后清空 IDF 缓存，下次查询重建."""
    with _IDF_CACHE_LOCK:
        _IDF_CACHE.pop(_IDF_CACHE_VERSION, None)


def build_tfidf_vector(tokens: List[str], idf_map: dict) -> Dict[str, float]:
    """构建 TF-IDF 稀疏向量 {word: weight}."""
    if not tokens:
        return {}
    tf = Counter(tokens)
    total = sum(tf.values()) or 1
    vec: Dict[str, float] = {}
    for word, count in tf.items():
        idf = idf_map.get(word)
        if idf is None:
            # 查询词不在语料中：用默认 IDF（视为稀有词），避免零向量
            idf = 1.0
        vec[word] = (count / total) * idf
    return vec


def cosine_similarity(v1: Dict[str, float], v2: Dict[str, float]) -> float:
    """稀疏向量余弦相似度."""
    if not v1 or not v2:
        return 0.0
    # 取较小向量遍历，提升性能
    if len(v1) > len(v2):
        v1, v2 = v2, v1
    dot = sum(wv * v2.get(w, 0.0) for w, wv in v1.items() if w in v2)
    norm1 = math.sqrt(sum(v * v for v in v1.values()))
    norm2 = math.sqrt(sum(v * v for v in v2.values()))
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return dot / (norm1 * norm2)


# ─── 文档索引流程 ───────────────────────────────────────────────

def index_document(db: Session, document_id: int) -> Tuple[bool, str]:
    """对文档执行：切片 → TF-IDF 向量化 → 入库. 返回 (success, message)."""
    doc = db.query(KbDocument).filter(KbDocument.id == document_id).first()
    if not doc:
        return False, f"文档 id={document_id} 不存在"
    if not doc.content or not doc.content.strip():
        doc.status = "failed"
        db.commit()
        return False, "文档内容为空"
    # 已索引先清旧切片（支持重新索引）
    db.query(KbChunk).filter(KbChunk.document_id == document_id).delete()
    db.commit()
    # 切片
    chunks = chunk_text(doc.content)
    if not chunks:
        doc.status = "failed"
        db.commit()
        return False, "切片后无有效内容（文档过短）"
    # 构建 IDF（基于现有语料 + 本文档；为简单起见，先入库切片再重建 IDF 缓存）
    # 此处用临时 IDF：仅本文档 + 已有 chunks
    # 更精确做法：先入库无 embedding 的切片，再统一计算 IDF，再回填 embedding
    # Phase 1 简化：先入库无 embedding 切片，然后调 recompute_all_embeddings
    for idx, content in enumerate(chunks):
        chunk = KbChunk(
            document_id=document_id,
            chunk_index=idx,
            content=content,
            embedding="",  # 占位，后续 recompute 填充
            embedding_mode="tfidf",
            token_count=len(tokenize(content)),
            tags=doc.tags or "",
            asset_type=doc.asset_type or "",
            severity=doc.severity or "warning",
        )
        db.add(chunk)
    doc.chunk_count = len(chunks)
    doc.status = "indexed"
    db.commit()
    # 重算所有切片的 embedding（包含本次新增），并清 IDF 缓存
    recompute_all_embeddings(db)
    invalidate_idf_cache()
    return True, f"已索引 {len(chunks)} 个切片"


def recompute_all_embeddings(db: Session):
    """重算所有 tfidf 切片的 embedding 向量. 用于文档变更后统一刷新."""
    # 先取所有切片文本构建 IDF
    chunks = db.query(KbChunk).filter(KbChunk.embedding_mode == "tfidf").all()
    if not chunks:
        return
    N = len(chunks)
    df: Counter = Counter()
    token_lists = []
    for c in chunks:
        toks = tokenize(c.content)
        token_lists.append(toks)
        for w in set(toks):
            df[w] += 1
    idf = {w: math.log((N + 1) / (cnt + 1)) + 1 for w, cnt in df.items()}
    # 回填 embedding
    for c, toks in zip(chunks, token_lists):
        vec = build_tfidf_vector(toks, idf)
        c.embedding = json.dumps(vec, ensure_ascii=False)
    db.commit()


# ─── 语义检索 ───────────────────────────────────────────────────

def vector_search(
    db: Session,
    query: str,
    top_k: int = 5,
    asset_type: Optional[str] = None,
    severity: Optional[str] = None,
    tags: Optional[str] = None,
) -> List[dict]:
    """语义检索：query 向量化 → 余弦相似度 Top-K → 返回结果.

    返回结构: [{document_id, document_title, content, similarity, tags, asset_type, severity}, ...]
    """
    if not query or not query.strip():
        return []
    # IDF 缓存（首次构建）
    idf_map = _build_idf_map(db)
    query_tokens = tokenize(query)
    query_vec = build_tfidf_vector(query_tokens, idf_map)
    if not query_vec:
        return []
    # 加载切片（带过滤）
    q = db.query(KbChunk).filter(KbChunk.embedding_mode == "tfidf")
    if asset_type:
        q = q.filter(KbChunk.asset_type == asset_type)
    if severity:
        q = q.filter(KbChunk.severity == severity)
    if tags:
        q = q.filter(KbChunk.tags.ilike(f"%{tags}%"))
    chunks = q.all()
    if not chunks:
        return []
    # 计算相似度并排序
    scored = []
    for c in chunks:
        if not c.embedding:
            continue
        try:
            chunk_vec = json.loads(c.embedding)
        except (json.JSONDecodeError, ValueError):
            continue
        sim = cosine_similarity(query_vec, chunk_vec)
        # 相似度阈值过滤：低于 0.05 视为无关，避免噪音结果
        if sim >= 0.05:
            scored.append((sim, c))
    scored.sort(key=lambda x: -x[0])
    # 取 Top-K，关联文档标题
    doc_ids = {c.document_id for _, c in scored[:top_k]}
    docs = {d.id: d for d in db.query(KbDocument).filter(KbDocument.id.in_(doc_ids)).all()} if doc_ids else {}
    results = []
    for sim, c in scored[:top_k]:
        doc = docs.get(c.document_id)
        results.append({
            "chunk_id": c.id,
            "document_id": c.document_id,
            "document_title": doc.title if doc else "(文档已删除)",
            "source_type": doc.source_type if doc else "",
            "content": c.content,
            "similarity": round(sim, 4),
            "tags": c.tags or "",
            "asset_type": c.asset_type or "",
            "severity": c.severity or "",
        })
    return results


# ─── 文档 CRUD ──────────────────────────────────────────────────

def list_documents(db: Session, search: str = "", source_type: str = "") -> List[KbDocument]:
    q = db.query(KbDocument)
    if search:
        q = q.filter(KbDocument.title.ilike(f"%{search}%"))
    if source_type:
        q = q.filter(KbDocument.source_type == source_type)
    return q.order_by(KbDocument.created_at.desc()).all()


def get_document(db: Session, doc_id: int) -> Optional[KbDocument]:
    return db.query(KbDocument).filter(KbDocument.id == doc_id).first()


def create_document(db: Session, data: dict) -> KbDocument:
    doc = KbDocument(**data)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


def update_document(db: Session, doc_id: int, data: dict) -> Optional[KbDocument]:
    doc = get_document(db, doc_id)
    if not doc:
        return None
    for k, v in data.items():
        setattr(doc, k, v)
    doc.updated_at = datetime.now()
    db.commit()
    db.refresh(doc)
    return doc


def delete_document(db: Session, doc_id: int):
    db.query(KbChunk).filter(KbChunk.document_id == doc_id).delete()
    db.query(KbDocument).filter(KbDocument.id == doc_id).delete()
    db.commit()
    invalidate_idf_cache()


def list_chunks(db: Session, doc_id: int) -> List[KbChunk]:
    return db.query(KbChunk).filter(KbChunk.document_id == doc_id).order_by(KbChunk.chunk_index).all()


# ─── 知识库自动沉淀：告警/故障单归档 ────────────────────────────

def archive_alert_case(db: Session, alert_id: int, title: str, content: str, tags: str = "", asset_type: str = "") -> KbDocument:
    """告警解决后自动归档为知识文档（source_type=alert_case）."""
    doc = KbDocument(
        title=f"[告警归档] {title}"[:256],
        source_type="alert_case",
        content=content,
        tags=tags,
        asset_type=asset_type,
        status="pending",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    # 自动索引
    index_document(db, doc.id)
    return doc


def archive_incident_case(db: Session, incident_id: int, title: str, content: str, tags: str = "") -> KbDocument:
    """故障单关闭后自动归档为知识文档（source_type=incident_case）."""
    doc = KbDocument(
        title=f"[故障归档] {title}"[:256],
        source_type="incident_case",
        content=content,
        tags=tags,
        status="pending",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    index_document(db, doc.id)
    return doc
